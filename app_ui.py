import streamlit as st
import os
import json
import hashlib
import shutil
import time
import tempfile
import uuid
import io
import openpyxl
from datetime import datetime as _dt
from google import genai
from exportar import (
    render_panel_exportacion,
    generar_html_plan_clase, generar_html_sintesis, generar_html_abp,
    generar_html_evaluaciones, generar_html_ficha, generar_html_pdc,
)
from db_pria import (
    init_db, crear_sesion, get_sesiones, get_sesion,
    guardar_micro_objetivos, get_micro_objetivos,
    marcar_objetivo, marcar_multiples,
    get_deuda_academica, get_dependencias_bloqueadas, get_resumen_deuda,
    guardar_plan_buffer, get_planes_buffer,
    minutos_para_fin_clase,
    # Auth
    crear_usuario, verificar_login, get_all_usuarios,
    actualizar_password, toggle_usuario_activo, eliminar_usuario,
    get_usuario_by_email,
    # Sistema Diario
    guardar_horario_docente, get_horario_dia, get_all_hojas,
    guardar_eventos_calendario, get_eventos_fecha, get_eventos_rango,
    guardar_actividades_cronograma, get_actividades_fecha,
    guardar_comisiones, get_comisiones_docente, get_all_comisiones,
    marcar_bloque_diario, cerrar_bloque, get_logs_dia, get_or_create_sesion_diaria,
    get_objetivos_semana_materia, reset_dia_docente, reabrir_bloque,
    guardar_vigilancias, get_vigilancias,
)
from parser_archivos import (
    parse_horarios, parse_calendario,
    parse_cronograma, parse_comisiones,
)

# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURACIÓN DE PANTALLA
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PRIA v5.4 — Método Palma-Ribera",
    layout="wide",
    page_icon="🦉"
)

# ─────────────────────────────────────────────────────────────────────────────
# CSS — PRIA Design System v2  ·  Las Palmas Green  ·  Depth + Hierarchy
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');

/* ══════════════════════════════════════════════════════════
   TOKENS — AID Dark System  ·  Elite Professional SaaS
══════════════════════════════════════════════════════════ */
:root {
    /* AID Dark */
    --bg:       #0D0F12;
    --surface:  #16191D;
    --surface2: #1E2128;
    --border:   #2A2D35;
    --border2:  #3A3D48;
    /* Accents */
    --blue:    #00C2FF;
    --emerald: #34C759;
    --amber:   #FF9500;
    --crimson: #FF3B30;
    /* Text */
    --text-1: #E8EAF0;
    --text-2: #9DA3B4;
    --text-3: #6B7185;
    /* Legacy aliases kept for any remaining references */
    --g900: #1B5E20; --g700: #2E7D32; --g500: #43A047;
    --g300: #81C784; --g100: #C8E6C9; --g50: #E8F5E9;
    --white: #FFFFFF;
    /* Shadows */
    --shadow-xs: 0 1px 3px rgba(0,0,0,0.3);
    --shadow-sm: 0 2px 8px rgba(0,0,0,0.4);
    --shadow-md: 0 4px 20px rgba(0,0,0,0.5);
    --shadow-lg: 0 8px 40px rgba(0,0,0,0.6);
    --radius-sm: 8px;
    --radius-md: 12px;
    --radius-lg: 18px;
    --transition: all 0.18s cubic-bezier(0.4,0,0.2,1);
}

/* ══════════════════════════════════════════════════════════
   GLOBAL FOUNDATION
══════════════════════════════════════════════════════════ */
html, body, .stApp {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
    background: var(--bg) !important;
    min-height: 100vh;
}
.main .block-container {
    padding-top: 2rem !important;
    max-width: 1200px;
}

/* ══════════════════════════════════════════════════════════
   TYPOGRAPHY
══════════════════════════════════════════════════════════ */
body, p, span, li, div, td, th, caption,
label, small, em, strong, blockquote,
.stMarkdown, .stText, .stCaption,
div[data-testid="stMarkdownContainer"],
div[data-testid="stMarkdownContainer"] *,
div[data-testid="stWidgetLabel"] p,
div[data-testid="stWidgetLabel"] span {
    color: var(--text-1) !important;
}
h1 {
    font-size: 2rem !important; font-weight: 800 !important;
    color: var(--text-1) !important; letter-spacing: -0.5px;
}
h2 {
    color: var(--text-1) !important; font-weight: 700 !important;
    font-size: 1.3rem !important;
    border-left: 3px solid var(--blue);
    padding-left: 0.6rem; margin-top: 1.2rem !important;
}
h3 { color: var(--text-1) !important; font-weight: 600 !important; font-size: 1.1rem !important; }
h4, h5, h6 { color: var(--text-2) !important; font-weight: 600 !important; }

/* ══════════════════════════════════════════════════════════
   TABS
══════════════════════════════════════════════════════════ */
div[data-testid="stTabs"] [role="tablist"] {
    gap: 4px !important;
    background: var(--surface) !important;
    border-radius: var(--radius-md) !important;
    padding: 5px !important;
    border: 1px solid var(--border) !important;
}
button[data-baseweb="tab"] {
    border-radius: var(--radius-sm) !important;
    padding: 0.45rem 1.1rem !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    transition: var(--transition) !important;
    border: none !important;
    color: var(--text-2) !important;
    background: transparent !important;
}
button[data-baseweb="tab"] p,
button[data-baseweb="tab"] span { color: inherit !important; font-weight: inherit !important; }
button[data-baseweb="tab"][aria-selected="true"] {
    background: var(--surface2) !important;
    color: var(--text-1) !important;
    box-shadow: var(--shadow-sm) !important;
}
div[data-testid="stTabs"] [role="tabpanel"] { padding-top: 1.5rem !important; }

/* ══════════════════════════════════════════════════════════
   BUTTONS
══════════════════════════════════════════════════════════ */
.stButton > button {
    background: var(--surface2) !important;
    color: var(--text-1) !important;
    font-weight: 600 !important;
    font-size: 0.88rem !important;
    border: 1px solid var(--border2) !important;
    border-radius: var(--radius-sm) !important;
    padding: 0.55rem 1.4rem !important;
    height: auto !important;
    min-height: 2.6rem !important;
    transition: var(--transition) !important;
    letter-spacing: 0.01em;
}
.stButton > button:hover {
    background: var(--border) !important;
    border-color: var(--blue) !important;
    color: var(--blue) !important;
    transform: translateY(-1px) !important;
}
.stButton > button[kind="primary"] {
    background: var(--blue) !important;
    color: #000 !important;
    border: none !important;
    font-weight: 700 !important;
}
.stButton > button[kind="primary"]:hover {
    background: #33CFFF !important;
    color: #000 !important;
    border: none !important;
    transform: translateY(-1px) !important;
}

/* ══════════════════════════════════════════════════════════
   INPUTS & SELECTS
══════════════════════════════════════════════════════════ */
input[type="text"], input[type="number"], textarea {
    background-color: var(--surface) !important;
    color: var(--text-1) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-sm) !important;
    transition: var(--transition) !important;
}
input[type="text"]:focus, input[type="number"]:focus, textarea:focus {
    border-color: var(--blue) !important;
    box-shadow: 0 0 0 2px rgba(0,194,255,0.15) !important;
    outline: none !important;
}
div[data-baseweb="select"] > div,
div[data-baseweb="select"] span {
    background-color: var(--surface) !important;
    color: var(--text-1) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-sm) !important;
}
div[data-baseweb="select"] > div:hover { border-color: var(--blue) !important; }
ul[data-baseweb="menu"] {
    background-color: var(--surface2) !important;
    border-radius: var(--radius-sm) !important;
    box-shadow: var(--shadow-md) !important;
    border: 1px solid var(--border) !important;
}
ul[data-baseweb="menu"] li, ul[data-baseweb="menu"] li span { color: var(--text-1) !important; }
ul[data-baseweb="menu"] li:hover { background-color: var(--border) !important; }

/* ══════════════════════════════════════════════════════════
   EXPANDERS  (card style)
══════════════════════════════════════════════════════════ */
div[data-testid="stExpander"] {
    background: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-md) !important;
    margin-bottom: 0.6rem !important;
    overflow: hidden !important;
    border-left: 3px solid var(--border2) !important;
    transition: var(--transition) !important;
}
details summary p, details summary span,
div[data-testid="stExpander"] summary p {
    color: var(--text-2) !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
}
div[data-testid="stVerticalBlockBorderWrapper"] {
    background: var(--surface) !important;
    border: 1px solid var(--border) !important;
    border-radius: var(--radius-md) !important;
}

/* ══════════════════════════════════════════════════════════
   ALERTS
══════════════════════════════════════════════════════════ */
div[data-testid="stAlert"] {
    border-radius: var(--radius-sm) !important;
    border-width: 0 !important;
    border-left-width: 3px !important;
    border-left-style: solid !important;
}
div[data-testid="stAlert"][data-baseweb="notification"][kind="info"]    { border-left-color: var(--blue) !important;    background: rgba(0,194,255,0.07) !important; }
div[data-testid="stAlert"][data-baseweb="notification"][kind="success"] { border-left-color: var(--emerald) !important; background: rgba(52,199,89,0.07) !important; }
div[data-testid="stAlert"][data-baseweb="notification"][kind="warning"] { border-left-color: var(--amber) !important;   background: rgba(255,149,0,0.07) !important; }
div[data-testid="stAlert"][data-baseweb="notification"][kind="error"]   { border-left-color: var(--crimson) !important; background: rgba(255,59,48,0.07) !important; }

/* ══════════════════════════════════════════════════════════
   METRICS
══════════════════════════════════════════════════════════ */
div[data-testid="stMetric"] {
    background: var(--surface) !important;
    border-radius: var(--radius-md) !important;
    padding: 1rem 1.2rem !important;
    border: 1px solid var(--border) !important;
    border-top: 2px solid var(--blue) !important;
}
div[data-testid="stMetricValue"] {
    font-size: 1.8rem !important; font-weight: 800 !important; color: var(--text-1) !important;
}
div[data-testid="stMetricLabel"] {
    font-weight: 500 !important; color: var(--text-2) !important;
    font-size: 0.78rem !important; text-transform: uppercase !important; letter-spacing: 0.06em !important;
}
div[data-testid="stMetricDelta"] { font-weight: 700 !important; }

/* ══════════════════════════════════════════════════════════
   DATAFRAME
══════════════════════════════════════════════════════════ */
.stDataFrame {
    border-radius: var(--radius-md) !important;
    border: 1px solid var(--border) !important;
    overflow: hidden !important;
}
.stDataFrame th {
    background: var(--surface2) !important;
    font-weight: 700 !important;
    font-size: 0.8rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.04em !important;
    color: var(--text-2) !important;
}
.stDataFrame tr:hover td { background: var(--surface2) !important; }

/* ══════════════════════════════════════════════════════════
   PROGRESS BAR
══════════════════════════════════════════════════════════ */
div[data-testid="stProgressBar"] > div {
    background: linear-gradient(90deg, var(--blue), var(--emerald)) !important;
    border-radius: 99px !important;
}
div[data-testid="stProgressBar"] {
    background: var(--border) !important;
    border-radius: 99px !important;
    height: 6px !important;
}

/* ══════════════════════════════════════════════════════════
   RADIO & CHECKBOX
══════════════════════════════════════════════════════════ */
.stRadio label, .stRadio label p,
.stRadio div[role="radiogroup"] label,
.stRadio div[role="radiogroup"] label p,
.stCheckbox label, .stCheckbox label p {
    color: var(--text-1) !important;
    font-weight: 500 !important;
}

/* ══════════════════════════════════════════════════════════
   DIVIDER
══════════════════════════════════════════════════════════ */
hr {
    border: none !important;
    height: 1px !important;
    background: var(--border) !important;
    margin: 1.2rem 0 !important;
}

/* ══════════════════════════════════════════════════════════
   SIDEBAR
══════════════════════════════════════════════════════════ */
section[data-testid="stSidebar"] {
    background: var(--surface) !important;
    border-right: 1px solid var(--border) !important;
}
section[data-testid="stSidebar"] * { color: var(--text-1) !important; }

/* ══════════════════════════════════════════════════════════
   SPINNER
══════════════════════════════════════════════════════════ */
div[data-testid="stSpinner"] > div {
    border-top-color: var(--blue) !important;
}

/* ══════════════════════════════════════════════════════════
   TABLA DIAGNÓSTICOS
══════════════════════════════════════════════════════════ */
.diag-table {
    width: 100%; border-collapse: separate; border-spacing: 0;
    margin: 10px 0 18px 0; border-radius: var(--radius-md); overflow: hidden;
}
.diag-table td {
    border: 1px solid var(--border); padding: 10px 14px; text-align: center;
    background: var(--surface); color: var(--text-1) !important;
    font-weight: 600; font-size: 0.88rem; line-height: 1.5;
}
.diag-table tr:first-child td { background: var(--surface2) !important; font-weight: 800; color: var(--text-2) !important; }
.diag-table tr:hover td { background: var(--surface2) !important; }

/* ══════════════════════════════════════════════════════════
   FILE UPLOADER
══════════════════════════════════════════════════════════ */
div[data-testid="stFileUploader"] {
    background: var(--surface) !important;
    border: 1.5px dashed var(--border2) !important;
    border-radius: var(--radius-md) !important;
    padding: 0.5rem !important;
}
div[data-testid="stFileUploader"]:hover { border-color: var(--blue) !important; }

/* ══════════════════════════════════════════════════════════
   NUMBER INPUT ARROWS
══════════════════════════════════════════════════════════ */
div[data-testid="stNumberInput"] button {
    background: var(--surface2) !important;
    border-color: var(--border) !important;
    color: var(--text-2) !important;
    min-height: unset !important;
    box-shadow: none !important;
}
div[data-testid="stNumberInput"] button:hover {
    background: var(--border) !important;
    transform: none !important;
}

/* ══════════════════════════════════════════════════════════
   TOAST NOTIFICATIONS
══════════════════════════════════════════════════════════ */
div[data-testid="stToast"] {
    background: var(--surface2) !important;
    border: 1px solid var(--border2) !important;
    border-radius: var(--radius-md) !important;
    box-shadow: var(--shadow-lg) !important;
}
div[data-testid="stToast"] * { color: var(--text-1) !important; }

/* ══════════════════════════════════════════════════════════
   DOWNLOAD BUTTON
══════════════════════════════════════════════════════════ */
div[data-testid="stDownloadButton"] > button {
    background: var(--emerald) !important;
    color: #000 !important;
    font-weight: 700 !important;
    border-radius: var(--radius-sm) !important;
    border: none !important;
}
div[data-testid="stDownloadButton"] > button:hover {
    opacity: 0.88 !important;
    transform: translateY(-1px) !important;
}

/* ══ AID BLOCK CARDS — Elite Professional Density ══════════ */
.aid-card--closed {
    background: rgba(52,199,89,0.04);
    border: 1px solid rgba(52,199,89,0.18);
    border-radius: 8px;
    padding: 6px 14px;
    margin-bottom: 3px;
    display: flex; align-items: center; gap: 10px;
}
.aid-badge {
    display: inline-block; padding: 1px 6px; border-radius: 4px;
    font-size: 0.65rem; font-weight: 700; letter-spacing: 0.07em;
    text-transform: uppercase; vertical-align: middle; flex-shrink: 0;
}
.aid-badge--closed  { background: rgba(52,199,89,0.12);  color: #34C759; }
.aid-badge--active  { background: rgba(0,194,255,0.12);   color: #00C2FF; }
.aid-badge--pending { background: rgba(255,149,0,0.12);   color: #FF9500; }

/* Compact block wrappers — high-density data rows */
div[data-testid="stVerticalBlockBorderWrapper"] {
    padding: 0.35rem 0.75rem !important;
    margin-bottom: 2px !important;
    border-radius: 6px !important;
    box-shadow: none !important;
}
/* Compress inter-element spacing INSIDE block containers */
div[data-testid="stVerticalBlockBorderWrapper"] .stMarkdown {
    margin-bottom: 0 !important; line-height: 1.3 !important;
}
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stCheckbox"] {
    padding-top: 0 !important; padding-bottom: 0 !important;
    min-height: 1.55rem !important;
}
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stCheckbox"] label,
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stCheckbox"] label p {
    font-size: 0.8rem !important; color: var(--text-2) !important;
    line-height: 1.3 !important;
}
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stHorizontalBlock"] {
    gap: 0.35rem !important; align-items: flex-start !important;
    margin-top: 0 !important;
}
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="column"] {
    padding: 0 !important;
}
/* Compact button inside blocks */
div[data-testid="stVerticalBlockBorderWrapper"] .stButton > button {
    padding: 0.25rem 0.5rem !important;
    min-height: 1.85rem !important;
    font-size: 0.78rem !important;
}
/* Compact text input (observations) inside blocks */
div[data-testid="stVerticalBlockBorderWrapper"] input[type="text"] {
    padding: 2px 7px !important;
    height: 1.7rem !important;
    font-size: 0.78rem !important;
    color: var(--text-3) !important;
}
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stTextInput"] {
    margin-top: 2px !important;
}
/* Active block — sibling selector glassmorphism trick */
.aid-active-marker + div + div[data-testid="stVerticalBlockBorderWrapper"],
.aid-active-marker ~ div[data-testid="stVerticalBlockBorderWrapper"]:first-of-type {
    border-color: rgba(0,194,255,0.5) !important;
    background: rgba(0,194,255,0.03) !important;
    box-shadow: inset 0 1px 0 rgba(0,194,255,0.08) !important;
}
/* Time display monospace */
.aid-time {
    font-family: 'SF Mono', 'Fira Code', 'Consolas', monospace;
    font-size: 0.8rem; font-weight: 600; color: #9DA3B4; letter-spacing: 0.02em;
}
/* Objectives inline row */
.aid-obj-row {
    display: flex; align-items: center; gap: 6px;
    padding: 2px 0; border-bottom: 1px solid rgba(255,255,255,0.04);
}
.aid-sync-state {
    display: flex; align-items: center; gap: 6px;
    padding: 5px 8px; border-radius: 6px;
    background: rgba(0,194,255,0.05); border: 1px solid rgba(0,194,255,0.12);
    font-size: 0.75rem; color: #9DA3B4; margin: 4px 0;
}

/* AID Day Header — compact precision instrument */
.aid-day-header {
    background: var(--surface); border: 1px solid var(--border);
    border-radius: 8px; padding: 12px 18px;
    margin-bottom: 14px;
    display: flex; align-items: center; justify-content: space-between;
}
.aid-day-header__left  { display: flex; flex-direction: column; gap: 1px; }
.aid-day-header__title { font-size: 1.05rem; font-weight: 800; color: #E8EAF0; letter-spacing: -0.3px; }
.aid-day-header__sub   { font-size: 0.75rem; font-weight: 500; color: #9DA3B4; }
.aid-day-header__right { display: flex; align-items: center; gap: 20px; }
.aid-day-header__meta  { text-align: right; }
.aid-day-header__cycle { font-size: 0.65rem; font-weight: 700; color: #00C2FF;
                          letter-spacing: 0.1em; text-transform: uppercase; }
.aid-day-header__daycount { font-size: 0.68rem; font-weight: 600; color: #6B7185;
                              letter-spacing: 0.04em; }
.aid-day-header__pct   { font-size: 1.5rem; font-weight: 800; color: #E8EAF0; line-height: 1; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# RUTAS DEL SISTEMA
# ─────────────────────────────────────────────────────────────────────────────
BASE_DIR         = os.path.dirname(os.path.abspath(__file__))
PROMPTS_DIR      = os.path.join(BASE_DIR, "prompts_maestros")
def _secret(key, default=None):
    try:
        v = st.secrets.get(key, None)
        if v is not None:
            return v
    except Exception:
        pass
    return os.environ.get(key, default)

CACHE_DIR        = _secret("CACHE_DIR",  os.path.join(BASE_DIR, "cache_libros"))
LOG_DIR          = _secret("LOG_DIR",    os.path.join(BASE_DIR, "logs"))
SESSION_BASE_DIR = os.path.join(tempfile.gettempdir(), "pria_sessions")
for _d in (CACHE_DIR, LOG_DIR, SESSION_BASE_DIR):
    os.makedirs(_d, exist_ok=True)

GEMINI_MODEL = "gemini-2.0-flash-lite"

# ─────────────────────────────────────────────────────────────────────────────
# INICIALIZACIÓN DE SESIÓN
# ─────────────────────────────────────────────────────────────────────────────
_defaults: dict = {
    "session_id":             None,
    "autenticado":            False,
    "grado_nivel":            "5to primaria",
    "key_index":              0,
    "last_error":             None,   # raw text of last failed generation
    "last_generar_fn":        None,   # prompt filename of last call
    "last_generar_vars":      None,   # variables dict of last call
    "last_generar_json":      False,  # expect_json flag of last call
    "uploaded_tb_bytes":      None,
    "uploaded_tb_name":       None,
    "uploaded_tb_hash":       None,
    "tb_extracted":           None,
    "uploaded_sb_bytes":      None,
    "uploaded_sb_name":       None,
    "uploaded_sb_hash":       None,
    "sb_extracted":           None,
    "uploaded_diag_files":    [],   # list of (nombre, bytes)
    "diagnosticos_texto":     None,
    "diagnosticos_tabla":     [],
    "res_m0a":                None,
    "res_m0b":                None,
    "res_m0c":                None,
    "tema_activo":            "",
    "tema_hash":              "",
    "leccion_index":          0,
    "conceptos_activos":      [],
    "palabras_clave_activas": [],
    "contenido_tema_activo":  "",
    "res_m1a":                None,
    "res_m1a_prev":           None,
    "res_m1b":                None,
    "res_m1c":                None,
    "res_m2a":                None,
    "res_m2b":                None,
    "mostrar_adaptaciones_prev": False,
    "_pptx_cache":            None,
    "teacher_name":           "",
    "school_name":            "",
}
for _k, _v in _defaults.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def limpiar_nombre_archivo(nombre: str) -> str:
    return nombre.replace(".pdf","").replace("_TB","").replace("_SB","").replace("_"," ").strip()

def forzar_lista(valor) -> list:
    if isinstance(valor, list):  return [str(v) for v in valor if v]
    if isinstance(valor, str):   return [v.strip() for v in valor.split(",") if v.strip()]
    return []

def _get_keys() -> list:
    import json as _json
    raw = _secret("GEMINI_API_KEYS", "[]")
    if isinstance(raw, list):
        return [str(k).strip() for k in raw if str(k).strip()]
    try:
        parsed = _json.loads(raw)
        if isinstance(parsed, list):
            keys = [str(k).strip() for k in parsed if str(k).strip()]
        elif isinstance(parsed, str):
            keys = [parsed.strip()] if parsed.strip() else []
        else:
            keys = []
        if keys:
            return keys
    except Exception:
        pass

    # Fallback: single-key env/secret
    single = str(_secret("GEMINI_API_KEY", "") or "").strip()
    if single:
        return [single]

    # Last fallback: raw non-JSON string in GEMINI_API_KEYS
    raw_s = str(raw or "").strip()
    return [raw_s] if raw_s and raw_s != "[]" else []

def _rotate_key():
    keys = _get_keys()
    if keys:
        st.session_state.key_index = (st.session_state.key_index + 1) % len(keys)

def _topic_hash(tema: str) -> str:
    if not tema: return ""
    return hashlib.md5(str(tema).strip().lower().encode()).hexdigest()

# ── Limpieza de sesiones temporales antiguas (se ejecuta al arrancar) ────────
def _cleanup_old_sessions():
    try:
        ahora = time.time()
        for nombre in os.listdir(SESSION_BASE_DIR):
            ruta = os.path.join(SESSION_BASE_DIR, nombre)
            if os.path.isdir(ruta) and ahora - os.path.getmtime(ruta) > 86400:
                shutil.rmtree(ruta, ignore_errors=True)
    except Exception:
        pass

_cleanup_old_sessions()


def _cleanup_old_cache():
    """Elimina entradas de caché global con más de 30 días de antigüedad."""
    try:
        ahora = time.time()
        for nombre in os.listdir(CACHE_DIR):
            ruta = os.path.join(CACHE_DIR, nombre)
            if os.path.isfile(ruta) and ahora - os.path.getmtime(ruta) > 86400 * 30:
                os.remove(ruta)
    except Exception:
        pass

_cleanup_old_cache()


def _log_event(action: str, success: bool, error_msg: str = ""):
    """Añade una línea JSON al log de piloto."""
    import datetime
    entry = {
        "ts":      datetime.datetime.utcnow().isoformat(),
        "session": (st.session_state.session_id or "")[:8],
        "action":  action,
        "ok":      success,
        "err":     error_msg[:300] if error_msg else "",
    }
    try:
        log_path = os.path.join(LOG_DIR, "pilot.log")
        with open(log_path, "a", encoding="utf-8") as fh:
            fh.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception:
        pass


def get_session_temp_dir() -> str:
    """Devuelve (y crea) un directorio temporal único por sesión de usuario."""
    if not st.session_state.session_id:
        st.session_state.session_id = str(uuid.uuid4())
    session_dir = os.path.join(SESSION_BASE_DIR, st.session_state.session_id)
    os.makedirs(session_dir, exist_ok=True)
    return session_dir


# ── Caché global por hash SHA-256 ────────────────────────────────────────────
def _bytes_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()

def _cargar_cache_hash(h: str) -> dict | None:
    path = os.path.join(CACHE_DIR, f"{h}.json")
    if not os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as fh:
            return json.load(fh)
    except Exception:
        return None

def _guardar_cache_hash(h: str, data: dict):
    path = os.path.join(CACHE_DIR, f"{h}.json")
    try:
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(data, fh, ensure_ascii=False, indent=2)
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# LECTURA AUTOMÁTICA DE DIAGNÓSTICOS (PDF / DOCX / TXT)
# ─────────────────────────────────────────────────────────────────────────────
def leer_diagnosticos(archivos_subidos: list) -> tuple:
    """
    archivos_subidos: list of (nombre_archivo, bytes)
    Devuelve (perfil_texto, tabla).
    """
    if not archivos_subidos:
        return None, []

    textos_raw = []
    keys = _get_keys()

    for nombre_archivo, contenido_bytes in archivos_subidos:
        ext       = nombre_archivo.lower().rsplit(".", 1)[-1]
        file_hash = _bytes_hash(contenido_bytes)
        cache_key = f"diag_{file_hash}"

        # Intenta cargar desde caché por hash (evita re-procesar archivos idénticos)
        cached = _cargar_cache_hash(cache_key)
        if cached and "texto" in cached:
            textos_raw.append(f"[{nombre_archivo}]\n{cached['texto']}")
            continue

        if ext == "txt":
            try:
                c = contenido_bytes.decode("utf-8", errors="ignore").strip()
                if c:
                    textos_raw.append(f"[{nombre_archivo}]\n{c}")
                    _guardar_cache_hash(cache_key, {"texto": c})
            except Exception as e:
                textos_raw.append(f"[{nombre_archivo}] Error TXT: {e}")

        elif ext == "docx":
            try:
                import io as _io
                import docx as _docx
                doc = _docx.Document(_io.BytesIO(contenido_bytes))
                p = "\n".join(x.text for x in doc.paragraphs if x.text.strip())
                if p:
                    textos_raw.append(f"[{nombre_archivo}]\n{p}")
                    _guardar_cache_hash(cache_key, {"texto": p})
            except ImportError:
                textos_raw.append(f"[{nombre_archivo}] Instala python-docx.")
            except Exception as e:
                textos_raw.append(f"[{nombre_archivo}] Error DOCX: {e}")

        elif ext == "pdf" and keys:
            temp_path = os.path.join(get_session_temp_dir(), f"diag_{file_hash[:8]}_{nombre_archivo}")
            with open(temp_path, "wb") as fh:
                fh.write(contenido_bytes)
            intentos_pdf = 0
            while intentos_pdf < len(keys):
                idx = st.session_state.key_index % len(keys)
                try:
                    client   = genai.Client(api_key=keys[idx])
                    uploaded = client.files.upload(file=temp_path)
                    resp = client.models.generate_content(
                        model=GEMINI_MODEL,
                        contents=[uploaded,
                            "Extrae todo el texto relevante de este documento de diagnóstico escolar. "
                            "Incluye nombre del estudiante (si aparece), diagnóstico principal, "
                            "necesidades específicas y recomendaciones pedagógicas. "
                            "Responde SOLO con el texto extraído, sin comentarios adicionales."]
                    )
                    client.files.delete(name=uploaded.name)
                    texto = resp.text.strip()
                    textos_raw.append(f"[{nombre_archivo}]\n{texto}")
                    _guardar_cache_hash(cache_key, {"texto": texto})
                    break
                except Exception as e:
                    err = str(e).upper()
                    if any(x in err for x in ["429","QUOTA","EXHAUSTED"]):
                        _rotate_key(); intentos_pdf += 1
                    else:
                        textos_raw.append(f"[{nombre_archivo}] Error PDF: {str(e)[:200]}"); break

    if not textos_raw:
        return None, []

    texto_combinado = "\n\n".join(textos_raw)
    prompt_sintesis = (
        "Eres especialista en necesidades educativas especiales.\n"
        "Documentos de diagnóstico del aula:\n\n"
        f"{texto_combinado}\n\n"
        "Devuelve ÚNICAMENTE un objeto JSON con:\n"
        '  "perfil_texto": string compacto (ej. "TEA (2), TDAH (1)")\n'
        '  "tabla": lista de objetos con "Diagnóstico", "Estudiantes" (int), "Notas pedagógicas"\n'
        "Sin texto adicional ni markdown."
    )
    intentos_sint = 0
    while intentos_sint < len(keys):
        idx = st.session_state.key_index % len(keys)
        try:
            client = genai.Client(api_key=keys[idx])
            resp   = client.models.generate_content(model=GEMINI_MODEL, contents=prompt_sintesis)
            clean  = resp.text.replace("```json","").replace("```","").strip()
            data   = json.loads(clean)
            return data.get("perfil_texto", None), data.get("tabla", [])
        except Exception as e:
            err = str(e).upper()
            if any(x in err for x in ["429","QUOTA","EXHAUSTED"]):
                _rotate_key(); intentos_sint += 1
            else:
                break

    perfil_crudo = "; ".join(t.split("\n")[0].replace("[","").replace("]","") for t in textos_raw[:5])
    return perfil_crudo, []


# ─────────────────────────────────────────────────────────────────────────────
# MOTOR ALPHA — extrae índice + CONTENIDO REAL + PÁGINAS por tema
# FIX: guarda resultado en disco (cache_libros/) para no re-analizar el PDF
# ─────────────────────────────────────────────────────────────────────────────
def analizar_pdf_ocr(pdf_bytes: bytes, filename: str) -> dict:
    """
    Extrae de un PDF (bytes):
      - unidad_real, temas, contenido_temas, paginas_temas
    Guarda y carga desde caché en disco (por filename) para evitar re-análisis.
    """
    VACIO = {"unidad_real":"Sin datos","temas":[],"contenido_temas":{},"paginas_temas":{}}

    # ── Caché global por hash SHA-256 ────────────────────────────────────────
    h = _bytes_hash(pdf_bytes)
    datos_cache = _cargar_cache_hash(h)
    if datos_cache:
        return datos_cache  # Hit de caché: sin llamada a Gemini

    # Escribe bytes a directorio temporal de sesión para subida a Gemini
    ruta_pdf = os.path.join(get_session_temp_dir(), filename)
    with open(ruta_pdf, "wb") as fh:
        fh.write(pdf_bytes)

    ruta_prompt = os.path.join(PROMPTS_DIR, "Motor_Alpha-2.txt")
    if not os.path.exists(ruta_prompt):
        return {**VACIO, "unidad_real":"Falta Motor_Alpha-2.txt"}

    with open(ruta_prompt, "r", encoding="utf-8") as fh:
        prompt_base = fh.read()

    prompt_completo = prompt_base + """

═══════════════════════════════════════════════════════
INSTRUCCIÓN ADICIONAL — CONTENIDO CURRICULAR Y PÁGINAS
═══════════════════════════════════════════════════════
Además del índice, examina el cuerpo de cada tema y extrae:

1. "contenido_temas": dict donde cada clave es el nombre exacto del tema
   y el valor es el texto LITERAL del libro para ese tema (definiciones,
   explicaciones, ejemplos, reglas). Mínimo 80 palabras por tema.
   Si una página es ilegible, escribe "(página ilegible)".

2. "paginas_temas": dict donde cada clave es el nombre exacto del tema
   y el valor es el rango de páginas como aparece en el índice del libro,
   por ejemplo: "pp. 12-15" o "pág. 34".

Incluye ambas claves en el JSON de salida.
Los nombres de los temas deben coincidir EXACTAMENTE con los de "temas".
"""

    keys = _get_keys()
    if not keys:
        return {**VACIO, "unidad_real":"Falta GEMINI_API_KEYS"}

    intentos = 0
    while intentos < len(keys):
        idx    = st.session_state.key_index % len(keys)
        client = genai.Client(api_key=keys[idx])
        try:
            uploaded = client.files.upload(file=ruta_pdf)
            response = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=[uploaded, prompt_completo]
            )
            client.files.delete(name=uploaded.name)
            clean = response.text.replace("```json","").replace("```","").strip()
            data  = json.loads(clean)
            if "contenido_temas" not in data: data["contenido_temas"] = {}
            if "paginas_temas"   not in data: data["paginas_temas"]   = {}
            _guardar_cache_hash(h, data)
            return data
        except json.JSONDecodeError:
            return {**VACIO, "unidad_real":"Error de parsing JSON"}
        except Exception as e:
            err = str(e).upper()
            if any(x in err for x in ["429","RESOURCE_EXHAUSTED","QUOTA"]):
                st.session_state.key_index = (idx+1) % len(keys)
                intentos += 1
            else:
                return {**VACIO, "unidad_real":f"Error Gemini: {str(e)[:150]}"}

    return {**VACIO, "unidad_real":"Todas las llaves agotadas"}


# ─────────────────────────────────────────────────────────────────────────────
# MOTOR GEMINI GENÉRICO
# ─────────────────────────────────────────────────────────────────────────────
def generar_con_gemini(prompt_filename: str, variables: dict, expect_json: bool = False,
                       required_fields: list | None = None):
    # Guarda args para eventual reintento
    st.session_state.last_generar_fn   = prompt_filename
    st.session_state.last_generar_vars = variables
    st.session_state.last_generar_json = expect_json

    ruta = os.path.join(PROMPTS_DIR, prompt_filename)
    if not os.path.exists(ruta):
        msg = f"Prompt no encontrado: {prompt_filename}"
        st.error(f"❌ {msg}")
        _log_event(f"generar:{prompt_filename}", False, msg)
        return None

    with open(ruta,"r",encoding="utf-8") as fh:
        template = fh.read()

    # Inyectar grado_nivel automáticamente si el caller no lo pasó
    if "grado_nivel" not in variables:
        variables = {**variables, "grado_nivel": st.session_state.get("grado_nivel", "5to primaria")}

    prompt_final = template
    for key, value in variables.items():
        prompt_final = prompt_final.replace("{"+key+"}", str(value))

    keys     = _get_keys()
    intentos = 0
    while intentos < len(keys):
        idx    = st.session_state.key_index % len(keys)
        client = genai.Client(api_key=keys[idx])
        try:
            _log_event(f"generar:{prompt_filename}", True)
            response = client.models.generate_content(model=GEMINI_MODEL, contents=prompt_final)
            if expect_json:
                clean = response.text.replace("```json","").replace("```","").strip()
                try:
                    data = json.loads(clean)
                except json.JSONDecodeError as e:
                    raw_snippet = response.text[:600]
                    st.session_state.last_error = response.text
                    _log_event(f"json_error:{prompt_filename}", False, str(e))
                    st.error(
                        f"❌ **JSON inválido** en `{prompt_filename}` — {e}\n\n"
                        f"**Respuesta recibida (primeros 600 caracteres):**\n```\n{raw_snippet}\n```"
                    )
                    return None
                # Validación de campos requeridos
                if required_fields:
                    faltantes = [f for f in required_fields if f not in data]
                    if faltantes:
                        st.warning(f"⚠️ `{prompt_filename}` devolvió JSON sin los campos esperados: "
                                   f"`{'`, `'.join(faltantes)}`. Verifica el resultado antes de continuar.")
                st.session_state.last_error = None
                return data
            st.session_state.last_error = None
            return response.text
        except Exception as e:
            err = str(e).upper()
            if any(x in err for x in ["429","RESOURCE_EXHAUSTED","QUOTA"]):
                st.session_state.key_index = (idx+1) % len(keys)
                intentos += 1
            else:
                _log_event(f"generar:{prompt_filename}", False, str(e))
                st.error(f"❌ Error en **{prompt_filename}**: {str(e)[:300]}")
                return None

    msg = "Todas las llaves API han agotado su cuota."
    _log_event(f"generar:{prompt_filename}", False, msg)
    st.error(f"⚠️ {msg}")
    return None


# ═════════════════════════════════════════════════════════════════════════════
# MÓDULO HORARIO — lee XLSX de cargas horarias (portado desde Nexus Docente)
# ═════════════════════════════════════════════════════════════════════════════
def leer_horario_xlsx(xlsx_bytes: bytes, teacher_name: str) -> tuple:
    """
    Lee el XLSX de horarios y devuelve (schedule_rows, time_headers).
    Portado desde pipeline-planificacion/app/api/schedule/route.ts
    """
    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)

    # Buscar hoja cuyo nombre contenga el nombre del docente
    target_sheet = None
    for sname in wb.sheetnames:
        if teacher_name.upper() in sname.upper():
            target_sheet = wb[sname]
            break

    if target_sheet is None:
        return None, None

    # Convertir hoja a matriz 2D
    data = []
    for row in target_sheet.iter_rows(values_only=True):
        data.append([str(c).strip() if c is not None else "" for c in row])

    days = ["LUNES", "MARTES", "MIERCOLES", "JUEVES", "VIERNES"]

    # Encontrar fila de encabezados que contenga "LUNES"
    header_row_idx = -1
    for i, row in enumerate(data):
        if "LUNES" in row:
            header_row_idx = i
            break

    if header_row_idx == -1:
        return [], []

    headers = data[header_row_idx]
    lunes_idx = headers.index("LUNES")
    time_headers = [h for h in headers[:lunes_idx] if h and h != "None"] or ["HORARIO"]

    day_indices = {d: (headers.index(d) if d in headers else -1) for d in days}

    schedule_map: dict = {}
    skip_words = ["VIGILANCIA", "SNACK", "INGRESO", "NONE"]

    for row in data[header_row_idx + 1:]:
        row_times = []
        is_valid = False
        for c in range(lunes_idx):
            val = str(row[c] if c < len(row) else "").strip()
            row_times.append(val)
            if "-" in val and any(ch.isdigit() for ch in val):
                is_valid = True

        if not is_valid:
            continue

        key = tuple(row_times)
        if key not in schedule_map:
            schedule_map[key] = {"times": row_times, **{d: [] for d in days}}

        for day, col_idx in day_indices.items():
            if col_idx != -1 and col_idx < len(row):
                val = str(row[col_idx] or "").strip()
                if val and not any(w in val.upper() for w in skip_words):
                    schedule_map[key][day].append(val)

    schedule = []
    for entry in schedule_map.values():
        schedule.append({
            "times": entry["times"],
            "LUNES":     " / ".join(entry["LUNES"])     or "Libre",
            "MARTES":    " / ".join(entry["MARTES"])    or "Libre",
            "MIERCOLES": " / ".join(entry["MIERCOLES"]) or "Libre",
            "JUEVES":    " / ".join(entry["JUEVES"])    or "Libre",
            "VIERNES":   " / ".join(entry["VIERNES"])   or "Libre",
        })

    return schedule, time_headers


# ═════════════════════════════════════════════════════════════════════════════
# MÓDULO PDC — genera PDC trimestral con Gemini (portado desde Nexus Docente)
# ═════════════════════════════════════════════════════════════════════════════
def _leer_docx_texto(docx_bytes: bytes) -> str:
    """Extrae texto plano de un archivo .docx."""
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def generar_pdc_trimestral(teacher_name: str, level: str, subject: str,
                           pdc_texto: str, semanas_calendario: int = 12) -> str | None:
    """
    Genera un PDC trimestral adaptado usando Gemini.
    Usa Motor_PDC_Trimestral.txt — igual que el resto de los motores de PRIA.
    """
    max_units = 4 if "matemática" in subject.lower() or "matematica" in subject.lower() else 3

    return generar_con_gemini(
        prompt_filename="Motor_PDC_Trimestral.txt",
        variables={
            "teacher_name":       teacher_name,
            "level":              level,
            "subject":            subject,
            "max_units":          str(max_units),
            "semanas_calendario": str(semanas_calendario),
            "pdc_texto":          pdc_texto[:3000],
        },
        expect_json=False,
    )


# ═════════════════════════════════════════════════════════════════════════════
# INTERFAZ PRINCIPAL
# ═════════════════════════════════════════════════════════════════════════════
st.title("🦉 PRIA — Planificación Neuro-Inclusiva")
st.caption("Método Palma-Ribera · DUA · Inteligencias Múltiples · Motor Gemini")

# ─────────────────────────────────────────────────────────────────────────────
# AUTENTICACIÓN — login por usuario (email + contraseña)
# ─────────────────────────────────────────────────────────────────────────────
if "usuario_email" not in st.session_state:
    st.session_state.usuario_email  = None
    st.session_state.usuario_nombre = None
    st.session_state.usuario_hoja   = None
    st.session_state.usuario_rol    = None

if not st.session_state.get("autenticado"):
    with st.container(border=True):
        st.markdown("### 🔐 Acceso al Sistema PRIA")
        st.caption("Ingresa con tu usuario o correo institucional y contraseña.")
        _email = st.text_input("Usuario o correo:", placeholder="admin  /  nombre@laspalmas.edu",
                               key="login_email")
        _pwd   = st.text_input("Contraseña:", type="password", key="login_pwd")
        if st.button("Ingresar", type="primary", use_container_width=True, key="btn_login"):
            _usuario = verificar_login(_email, _pwd)
            if _usuario:
                st.session_state.autenticado      = True
                st.session_state.usuario_email    = _usuario["email"]
                st.session_state.usuario_nombre   = _usuario["nombre"]
                st.session_state.usuario_hoja     = _usuario["nombre_hoja"]
                st.session_state.usuario_rol      = _usuario["rol"]
                st.session_state.teacher_name     = _usuario["nombre"]
                _log_event("login", True)
                st.rerun()
            else:
                _log_event("login_failed", False, "wrong credentials")
                st.error("❌ Correo o contraseña incorrectos.")
    st.stop()

ss = st.session_state

# ═════════════════════════════════════════════════════════════════════════════
# SIDEBAR — Perfil Docente + Materiales
# ═════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/owl.png", width=80)

    # ── Perfil Docente ────────────────────────────────────────────────────────
    st.markdown("### 👤 Perfil Docente")
    _nombre_display = ss.get("usuario_nombre") or ss.get("teacher_name", "")
    _rol_display    = ss.get("usuario_rol", "docente")
    st.markdown(f"**{_nombre_display}**")
    st.caption(f"{'👑 Administrador' if _rol_display == 'admin' else '🎓 Docente'} · {ss.get('usuario_email','')}")
    if st.button("Cerrar sesión", key="btn_logout", use_container_width=True):
        for _k in ["autenticado","usuario_email","usuario_nombre","usuario_hoja","usuario_rol"]:
            st.session_state.pop(_k, None)
        st.rerun()

    # ── Nivel y Grado ─────────────────────────────────────────────────────────
    st.divider()
    st.markdown("### 🎓 Nivel Educativo")
    _nivel = st.selectbox("Nivel:", ["Primaria", "Secundaria"],
                          index=0 if "primaria" in ss.grado_nivel else 1,
                          key="sel_nivel")
    _max_grado = 6 if _nivel == "Primaria" else 5
    _sufijos = {1: "ro", 2: "do", 3: "ro", 4: "to", 5: "to", 6: "to"}
    _grado_options = [f"{g}{_sufijos.get(g,'mo')} {_nivel.lower()}" for g in range(1, _max_grado + 1)]
    _sel_idx = next((i for i, o in enumerate(_grado_options) if o == ss.grado_nivel),
                    4 if _nivel == "Primaria" else 0)
    _grado = st.selectbox("Grado:", _grado_options, index=_sel_idx, key="sel_grado")
    if _grado != ss.grado_nivel:
        ss.grado_nivel = _grado
        st.rerun()

    # ── Materiales ────────────────────────────────────────────────────────────
    st.divider()
    st.markdown("### 📥 Materiales")

    uploaded_tb = st.file_uploader("📗 Libro de Texto (PDF):", type=["pdf"], key="uploader_tb")
    if uploaded_tb is not None:
        tb_bytes    = uploaded_tb.getvalue()
        tb_hash_now = _bytes_hash(tb_bytes)
        if tb_hash_now != ss.uploaded_tb_hash:
            ss.uploaded_tb_bytes = tb_bytes
            ss.uploaded_tb_name  = uploaded_tb.name
            ss.uploaded_tb_hash  = tb_hash_now
            ss.tb_extracted      = None
            _log_event(f"upload_tb:{uploaded_tb.name}", True)
        if ss.tb_extracted is None:
            if not _cargar_cache_hash(tb_hash_now):
                with st.spinner("👁️ Motor Alpha leyendo el PDF…"):
                    ss.tb_extracted = analizar_pdf_ocr(ss.uploaded_tb_bytes, ss.uploaded_tb_name)
                    _log_event("analizar_tb", "unidad_real" in ss.tb_extracted)
            else:
                ss.tb_extracted = analizar_pdf_ocr(ss.uploaded_tb_bytes, ss.uploaded_tb_name)
        datos_libro = ss.tb_extracted
        cache_label = "📦 caché" if _cargar_cache_hash(ss.uploaded_tb_hash) else "🔍 nuevo"
        st.success(f"📚 {datos_libro.get('unidad_real','…')} *({cache_label})*")
    else:
        datos_libro = {"unidad_real": "", "temas": [], "contenido_temas": {}, "paginas_temas": {}}

    usar_sb = st.radio("📘 ¿Usa Student Book?", ["No", "Sí"], horizontal=True, key="radio_sb")
    datos_sb = {}
    if usar_sb == "Sí":
        uploaded_sb = st.file_uploader("📙 Student Book (PDF):", type=["pdf"], key="uploader_sb")
        if uploaded_sb is not None:
            sb_bytes    = uploaded_sb.getvalue()
            sb_hash_now = _bytes_hash(sb_bytes)
            if sb_hash_now != ss.uploaded_sb_hash:
                ss.uploaded_sb_bytes = sb_bytes
                ss.uploaded_sb_name  = uploaded_sb.name
                ss.uploaded_sb_hash  = sb_hash_now
                ss.sb_extracted      = None
            if ss.sb_extracted is None:
                cache_sb = _cargar_cache_hash(sb_hash_now)
                if not cache_sb:
                    with st.spinner("👁️ Leyendo Student Book…"):
                        ss.sb_extracted = analizar_pdf_ocr(ss.uploaded_sb_bytes, ss.uploaded_sb_name)
                else:
                    ss.sb_extracted = cache_sb
            datos_sb = ss.sb_extracted or {}

    # ── Diagnósticos ──────────────────────────────────────────────────────────
    st.divider()
    st.markdown("### 🩺 Diagnósticos")
    uploaded_diags = st.file_uploader(
        "Archivos (PDF, DOCX, TXT):",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="uploader_diag"
    )
    if uploaded_diags:
        ss.uploaded_diag_files = [(f.name, f.getvalue()) for f in uploaded_diags]

    if ss.uploaded_diag_files and ss.diagnosticos_texto is None:
        with st.spinner("🔍 Leyendo diagnósticos…"):
            p_texto, p_tabla = leer_diagnosticos(ss.uploaded_diag_files)
            ss.diagnosticos_texto = p_texto
            ss.diagnosticos_tabla = p_tabla

    if ss.diagnosticos_texto:
        diag_preview = ss.diagnosticos_texto[:80] + ("…" if len(ss.diagnosticos_texto) > 80 else "")
        st.success(f"✅ {diag_preview}")
        if ss.diagnosticos_tabla:
            st.dataframe(ss.diagnosticos_tabla, use_container_width=True, hide_index=True)
        if st.button("🔄 Releer diagnósticos", use_container_width=True, key="btn_releer_diag"):
            ss.diagnosticos_texto = None
            ss.diagnosticos_tabla = []
            st.rerun()
    elif ss.uploaded_diag_files:
        st.info("Procesando archivos…")
    else:
        st.caption("Sin archivos cargados.")

    DIAGNOSTICOS = ss.diagnosticos_texto or ""

    # ── Estado del Sistema ────────────────────────────────────────────────────
    st.divider()
    st.markdown("### 📊 Estado del Sistema")
    estado = {
        "Síntesis Unidad":  ss.res_m0a is not None,
        "Proyecto ABP":     ss.res_m0b is not None,
        "Plan de Clase":    ss.res_m1a is not None,
        "Diapositivas":     ss.res_m1b is not None,
        "Ficha Gamificada": ss.res_m1c is not None,
        "Pop Quiz":         ss.res_m2a is not None,
        "Guía del Tutor":   ss.res_m2b is not None,
    }
    for nombre, listo in estado.items():
        st.markdown(f"{'✅' if listo else '⏳'} {nombre}")

    if ss.tema_activo:
        st.divider()
        st.markdown(f"**Tema activo:** {ss.tema_activo}")
        st.markdown(f"**Lecciones generadas:** {ss.leccion_index}")

    cache_files = [f for f in os.listdir(CACHE_DIR) if f.endswith(".json")] if os.path.exists(CACHE_DIR) else []
    if cache_files:
        st.divider()
        st.markdown(f"**📦 Caché:** {len(cache_files)} entradas")
        if st.button("🗑️ Vaciar caché", use_container_width=True):
            for cf in cache_files:
                try: os.remove(os.path.join(CACHE_DIR, cf))
                except Exception: pass
            ss.tb_extracted = None
            ss.sb_extracted = None
            st.rerun()

    if ss.last_error:
        st.divider()
        st.warning("⚠️ Última generación falló.")
        if st.button("🔄 Reintentar", use_container_width=True):
            fn = ss.last_generar_fn; vrs = ss.last_generar_vars; xjsn = ss.last_generar_json
            if fn and vrs is not None:
                ss.last_error = None
                generar_con_gemini(fn, vrs, xjsn)
                st.rerun()

    st.divider()
    if st.button("🧹 Reiniciar Todo", use_container_width=True):
        st.cache_data.clear()
        for k, v in _defaults.items():
            ss[k] = v
        st.rerun()
    if st.button("🚪 Cerrar Sesión", use_container_width=True):
        _log_event("logout", True)
        ss.autenticado = False
        st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# DATOS DERIVADOS (accesibles en todas las zonas)
# ─────────────────────────────────────────────────────────────────────────────
temas_reales    = forzar_lista(datos_libro.get("temas", []))
contenido_temas = datos_libro.get("contenido_temas", {})
paginas_temas   = datos_libro.get("paginas_temas",   {})

# ─────────────────────────────────────────────────────────────────────────────
# EMPTY STATE — sin libro cargado
# ─────────────────────────────────────────────────────────────────────────────
# REMOVED: Teachers now land directly in the Daily Tracker (admin pre-loads dependencies)


# ─────────────────────────────────────────────────────────────────────────────
# PANEL DE ADMINISTRACIÓN (solo rol admin)
# ─────────────────────────────────────────────────────────────────────────────
if ss.get("usuario_rol") == "admin":
    with st.expander("⚙️ Panel de Administración", expanded=False):
        adm_tab_arch, adm_tab_users, adm_tab_tracker = st.tabs(["📂 Archivos Fuente", "👥 Gestión de Usuarios", "🌅 Reset Diario"])

        # ── ARCHIVOS FUENTE ───────────────────────────────────────────────────
        with adm_tab_arch:
            st.markdown("Sube los archivos oficiales del colegio. Al subir, los datos anteriores se reemplazan.")
            st.info("💾 Todos los archivos importados aquí se guardan permanentemente en la base de datos. **No necesitas volver a subirlos cada vez**, solo cuando hayan actualizaciones oficiales.", icon="💾")
            col_a1, col_a2 = st.columns(2)

            with col_a1:
                st.markdown("**📅 Horarios Oficiales (.xlsx)**")
                st.info("💡 El nombre de la pestaña en el Excel (ej. RUDDY) debe coincidir con el código del docente para sincronizarse.", icon="💡")
                f_horario = st.file_uploader("Horarios", type=["xlsx"], key="adm_horario",
                                             label_visibility="collapsed")
                if f_horario and st.button("Importar horarios", key="btn_imp_horario"):
                    with st.spinner("Parseando horarios..."):
                        try:
                            registros = parse_horarios(f_horario.read())
                            guardar_horario_docente(registros)
                            st.success(f"✅ {len(registros)} bloques importados para {len(set(r['nombre_hoja'] for r in registros))} docentes.")
                        except Exception as _e:
                            st.error(f"Error: {_e}")

                st.markdown("**📋 Cronograma Semanal Docente (.xlsx)**")
                f_crono = st.file_uploader("Cronograma", type=["xlsx"], key="adm_crono",
                                           label_visibility="collapsed")
                if f_crono and st.button("Importar cronograma", key="btn_imp_crono"):
                    with st.spinner("Parseando cronograma..."):
                        try:
                            registros = parse_cronograma(f_crono.read())
                            guardar_actividades_cronograma(registros)
                            st.success(f"✅ {len(registros)} actividades importadas.")
                        except Exception as _e:
                            st.error(f"Error: {_e}")

            with col_a2:
                st.markdown("**📆 Calendario Interno (.xlsx)**")
                f_cal = st.file_uploader("Calendario", type=["xlsx"], key="adm_cal",
                                         label_visibility="collapsed")
                if f_cal and st.button("Importar calendario", key="btn_imp_cal"):
                    with st.spinner("Parseando calendario..."):
                        try:
                            registros = parse_calendario(f_cal.read())
                            guardar_eventos_calendario(registros)
                            st.success(f"✅ {len(registros)} eventos importados.")
                        except Exception as _e:
                            st.error(f"Error: {_e}")

                st.markdown("**🏛️ Comisiones Docentes (.docx)**")
                f_com = st.file_uploader("Comisiones", type=["docx"], key="adm_com",
                                         label_visibility="collapsed")
                if f_com and st.button("Importar comisiones", key="btn_imp_com"):
                    with st.spinner("Parseando comisiones..."):
                        try:
                            registros = parse_comisiones(f_com.read())
                            guardar_comisiones(registros)
                            st.success(f"✅ {len(registros)} asignaciones importadas.")
                        except Exception as _e:
                            st.error(f"Error: {_e}")

                st.markdown("**👁️ Roles de Vigilancia Recreos (.pdf)**")
                f_vig = st.file_uploader("Vigilancias", type=["pdf"], key="adm_vig",
                                         label_visibility="collapsed")
                if f_vig and st.button("Extraer ubicaciones (AI Gemini)", key="btn_imp_vig"):
                    with st.spinner("Extrayendo Inteligencia..."):
                        try:
                            import tempfile
                            _keys = _get_keys()
                            if not _keys:
                                st.error(
                                    "No hay clave Gemini configurada en el servidor. "
                                    "Configura `GEMINI_API_KEYS` (JSON array) o `GEMINI_API_KEY` en Railway Variables."
                                )
                                st.stop()
                            _client = genai.Client(api_key=_keys[0])
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                                tmp.write(f_vig.read())
                                tmp_path = tmp.name
                            _gfile = _client.files.upload(file=tmp_path)
                            _prompt = "Devuelve un JSON estrictamente con la lista de TODOS los docentes y su lugar de vigilancia. Formato de repuesta SOLAMENTE arreglo: [{\"nombre_hoja\": \"(un solo apellido o primer nombre en MAYUSCULAS del docente)\", \"ubicacion\": \"(zona y turno)\"}]. Identifica a cada docente por su nombre de tabla primario (ej. VANESA o SUSI o CUELLAR)."
                            _resp = _client.models.generate_content(model=GEMINI_MODEL, contents=[_prompt, _gfile])
                            import os; os.remove(tmp_path)
                            _clean = _resp.text.replace("```json","").replace("```","").strip()
                            import json; _data = json.loads(_clean)
                            if isinstance(_data, dict): _data = _data.get("vigilancias", [_data])
                            for entry in _data:
                                n = entry.get("nombre_hoja", "").upper().replace(",", "").replace(".", "")
                                entry["nombre_hoja"] = n.split()[0] if n else "UNKNOWN"
                            guardar_vigilancias(_data)
                            st.success(f"✅ {len(_data)} asignaciones importadas inteligentemente.")
                        except Exception as _e:
                            st.error(f"Error AI: {_e}")

            # ── PLAN SEMANAL DOCENTE — Palma-Ribera Extraction Engine ────────
            st.divider()
            st.markdown("**📋 Plan Semanal Docente (.docx / .txt)**")
            st.caption("La IA infiere objetivos desde el campo CONTENIDO y los descompone en tareas progresivas por bloque horario.")
            _pa_r1c1, _pa_r1c2 = st.columns([5, 2])
            with _pa_r1c1:
                f_plan = st.file_uploader("Plan semanal", type=["docx", "txt"],
                                          key="adm_plan", label_visibility="collapsed")
            with _pa_r1c2:
                _sw_default = max(1, (_dt.now().date() - _dt(2026, 2, 2).date()).days // 7 + 1)
                _plan_semana = st.number_input("Semana escolar N°", min_value=1, max_value=40,
                                               value=_sw_default, key="adm_plan_semana")

            _pa_r2c1, _pa_r2c2 = st.columns([3, 3])
            with _pa_r2c1:
                _plan_grado = st.text_input("Grado / Nivel", key="adm_plan_grado",
                                            placeholder="Ej. 5to Primaria")
            with _pa_r2c2:
                _hojas_plan = get_all_hojas()
                _plan_hoja = st.selectbox(
                    "Sincronizar con el horario de:", _hojas_plan if _hojas_plan else ["— sin horarios —"],
                    key="adm_plan_hoja"
                )

            _can_import = bool(f_plan and _plan_grado.strip())
            if st.button("🧠 Inferir objetivos y tareas progresivas", type="primary",
                         key="btn_imp_plan", disabled=not _can_import):
                with st.spinner("Aplicando Método Palma-Ribera — inferencia de objetivos desde CONTENIDO…"):
                    try:
                        import docx as _dx
                        if f_plan.name.lower().endswith('.docx'):
                            _doc = _dx.Document(io.BytesIO(f_plan.read()))
                            _texts = [p.text.strip() for p in _doc.paragraphs if p.text.strip()]
                            for table in _doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        if cell.text.strip(): _texts.append(cell.text.strip())
                            _plan_txt = "\n".join(_texts)
                        else:
                            _plan_txt = f_plan.read().decode('utf-8', errors='ignore')

                        if not _plan_txt.strip():
                            st.error("El archivo está vacío o no se pudo leer el texto.")
                        else:
                            _plan_prompt = (
                                "Eres un experto en el Método Palma-Ribera para planificación pedagógica "
                                f"de Las Palmas School. Grado: {_plan_grado.strip()}.\n\n"
                                "Se te proporciona un plan semanal docente. Tu tarea es:\n"
                                "1. Identificar TODOS los bloques de clase con su horario y materia.\n"
                                "2. Tomar el campo CONTENIDO de cada bloque como fuente primaria — "
                                "NO busques una sección llamada 'Objetivos'.\n"
                                "3. INFERIR el objetivo de aprendizaje usando un verbo cognitivo de la "
                                "taxonomía de Bloom (representar, analizar, deducir, construir, evaluar…).\n"
                                "4. DESCOMPONER cada contenido en exactamente 3 o 4 tareas progresivas y "
                                "accionables que el docente ejecutará en el aula, ordenadas cronológicamente.\n\n"
                                "CRÍTICO: Debes INFERIR OBLIGATORIAMENTE los objetivos y las tareas para CADA bloque "
                                "incluso si el documento solo contiene un título o tema genérico. Usa tu conocimiento "
                                "pedagógico para crear las tareas automáticamente si no constan de forma explícita.\n\n"
                                f"PLAN SEMANAL:\n{_plan_txt[:6000]}\n\n"
                                "Responde ÚNICAMENTE con un array JSON válido (sin markdown, sin texto extra).\n"
                                "Formato exacto por cada bloque encontrado:\n"
                                '[{"hora_inicio":"07:55","hora_fin":"08:40","materia":"LENGUAJE",'
                                '"objetivo_inferido":"Los estudiantes lograrán...",'
                                '"tareas":["Tarea 1","Tarea 2","Tarea 3","Tarea 4"]}]'
                            )
                            _plan_keys = _get_keys()
                            _plan_client = genai.Client(api_key=_plan_keys[0])
                            _plan_resp = _plan_client.models.generate_content(
                                model=GEMINI_MODEL, contents=_plan_prompt
                            )
                            _plan_clean = _plan_resp.text.replace("```json","").replace("```","").strip()
                            # Handle both array and object responses
                            _plan_data = json.loads(_plan_clean)
                            if isinstance(_plan_data, dict):
                                _plan_data = _plan_data.get("bloques", _plan_data.get("clases", [_plan_data]))
                            if not isinstance(_plan_data, list):
                                _plan_data = []

                            _tot_sessions = 0
                            _tot_tasks    = 0
                            for _blk in _plan_data:
                                _h_ini = str(_blk.get("hora_inicio", "00:00")).strip()
                                _h_fin = str(_blk.get("hora_fin",    "00:00")).strip()
                                _mat   = str(_blk.get("materia", "Sin materia")).strip()
                                _obj_raw = _blk.get("objetivo_inferido", _blk.get("objetivo", _blk.get("tema", "")))
                                _obj   = str(_obj_raw).strip() if _obj_raw else ""
                                _tasks_raw = _blk.get("tareas", _blk.get("tareas_progresivas", []))
                                _tasks = []
                                if isinstance(_tasks_raw, list):
                                    for t in _tasks_raw:
                                        if isinstance(t, str): _tasks.append(t.strip())
                                        elif isinstance(t, dict): _tasks.append(str(list(t.values())[0]).strip())
                
                                if not _obj and not _tasks:
                                    continue
                                if not _tasks and _obj:
                                    _tasks = [_obj]

                                if not _tasks:
                                    continue

                                _sid = crear_sesion(
                                    semana=int(_plan_semana),
                                    materia=_mat,
                                    grado=_plan_grado.strip(),
                                    tema=_obj or f"Plan Semana {_plan_semana} — {_mat}",
                                    hora_inicio=_h_ini,
                                    hora_fin=_h_fin,
                                )
                                _mo_list = []
                                if _obj:
                                    _mo_list.append({"texto": f"🎯 {_obj}", "depende_de": None})
                                for _i, _t in enumerate(_tasks, 1):
                                    _mo_list.append({"texto": f"{_i}· {_t}", "depende_de": None})
                                guardar_micro_objetivos(_sid, _mo_list, origen_semana=int(_plan_semana))
                                _tot_sessions += 1
                                _tot_tasks    += len(_tasks)

                            if _tot_sessions == 0:
                                st.warning(
                                    f"La IA no encontró bloques estructurados. "
                                    f"Respuesta (primeros 400 chars):\n\n`{_plan_clean[:400]}`"
                                )
                            else:
                                st.success(
                                    f"✅ {_tot_sessions} bloques importados · "
                                    f"{_tot_tasks} tareas progresivas · Semana {_plan_semana} · "
                                    f"{_plan_grado.strip()}."
                                )
                    except json.JSONDecodeError:
                        st.error(f"JSON inválido. Respuesta de la IA:\n\n`{_plan_resp.text[:500]}`")
                    except Exception as _plan_e:
                        st.error(f"Error: {_plan_e}")

        # ── GESTIÓN DE USUARIOS ───────────────────────────────────────────────
        with adm_tab_users:
            st.markdown("#### Crear nuevo usuario")
            hojas_disponibles = get_all_hojas()
            cu_col1, cu_col2 = st.columns(2)
            with cu_col1:
                cu_nombre = st.text_input("Nombre completo", key="cu_nombre", placeholder="Ej. Ruddy Ribera")
                cu_email  = st.text_input("Correo electrónico", key="cu_email", placeholder="ruddy@laspalmas.edu")
                cu_pwd    = st.text_input("Contraseña inicial", type="password", key="cu_pwd")
            with cu_col2:
                cu_hoja = st.selectbox(
                    "Hoja de horario (nombre en el Excel)",
                    options=hojas_disponibles if hojas_disponibles else ["— sin horario cargado —"],
                    key="cu_hoja"
                )
                cu_hoja_manual = st.text_input("O escribe el nombre manualmente", key="cu_hoja_manual",
                                               placeholder="Ej. RUDDY")
                cu_rol  = st.selectbox("Rol", ["docente", "admin"], key="cu_rol")

            if st.button("Crear usuario", type="primary", key="btn_crear_usuario"):
                hoja_final = cu_hoja_manual.strip().upper() or (cu_hoja if cu_hoja != "— sin horario cargado —" else "")
                if cu_nombre and cu_email and cu_pwd and hoja_final:
                    ok = crear_usuario(cu_email, cu_pwd, cu_nombre, hoja_final, cu_rol)
                    if ok:
                        st.success(f"✅ Usuario **{cu_nombre}** creado.")
                    else:
                        st.error("El correo ya está registrado.")
                else:
                    st.warning("Completa todos los campos.")

            st.divider()
            st.markdown("#### Usuarios registrados")
            usuarios = get_all_usuarios()
            if usuarios:
                for u in usuarios:
                    u_col1, u_col2, u_col3 = st.columns([3, 1, 1])
                    with u_col1:
                        estado = "✅" if u["activo"] else "⛔"
                        st.markdown(f"{estado} **{u['nombre']}** — {u['email']} · `{u['nombre_hoja']}` · _{u['rol']}_")
                    with u_col2:
                        lbl = "Desactivar" if u["activo"] else "Activar"
                        if st.button(lbl, key=f"toggle_{u['id']}"):
                            toggle_usuario_activo(u["id"], not u["activo"])
                            st.rerun()
                    with u_col3:
                        if st.button("Eliminar", key=f"del_{u['id']}"):
                            eliminar_usuario(u["id"])
                            st.rerun()
            else:
                st.info("No hay usuarios registrados aún.")

        # ── RESET DIARIO (ADMIN) ───────────────────────────────────────────────
        with adm_tab_tracker:
            st.markdown("#### Reiniciar Tracker Diario")
            st.caption("Limpia todos los registros del día para un docente, permitiéndole empezar de cero.")
            _admin_rst_date = st.date_input("Fecha a reiniciar:")
            _admin_rst_hojas = get_all_hojas()
            _admin_rst_hoja = st.selectbox("Docente", _admin_rst_hojas if _admin_rst_hojas else ["—"], key="adm_rst_hoja")
            if st.button("↺ Reiniciar Tracker", type="primary", key="btn_adm_rst_dia"):
                if _admin_rst_hoja and _admin_rst_hoja != "—":
                    reset_dia_docente(_admin_rst_date.strftime('%Y-%m-%d'), _admin_rst_hoja)
                    st.success(f"Día {_admin_rst_date} reiniciado con éxito para la hoja: {_admin_rst_hoja}.")
                else:
                    st.error("Selecciona un docente válido.")

# ─────────────────────────────────────────────────────────────────────────────
# SELECTOR TEMPORAL
# ─────────────────────────────────────────────────────────────────────────────
_zona_opciones = ["🌅  Diario", "📅  Semanal", "📆  Trimestral"]
zona = st.radio(
    "Zona de trabajo:",
    _zona_opciones,
    horizontal=True,
    key="zona_temporal",
    label_visibility="collapsed",
)
st.markdown("---")


# ═════════════════════════════════════════════════════════════════════════════
# ZONA DIARIA — Dashboard del Día + Deuda Académica + Comisiones
# ═════════════════════════════════════════════════════════════════════════════
if zona == "🌅  Diario":
    tab_dia, tab_deuda, tab_comisiones = st.tabs([
        "📅 Mi Día",
        "📊 Deuda Académica",
        "🏛️ Mis Comisiones",
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 1 — MI DÍA (dashboard diario automático)
    # ══════════════════════════════════════════════════════════════════════════
    _DIAS_HOY = {0:'lunes',1:'martes',2:'miercoles',3:'jueves',4:'viernes'}
    _DIAS_LABEL = {'lunes':'Lunes','martes':'Martes','miercoles':'Miércoles',
                   'jueves':'Jueves','viernes':'Viernes'}
    _MESES_LABEL = {1:'enero',2:'febrero',3:'marzo',4:'abril',5:'mayo',
                    6:'junio',7:'julio',8:'agosto',9:'septiembre',
                    10:'octubre',11:'noviembre',12:'diciembre'}
    _BLOQUE_ICONO = {
        'ingreso':          '🏫',
        'clase':            '📚',
        'vigilancia_recreo':'👁️',
        'atencion_ppff':    '👨‍👩‍👧',
        'planificacion':    '📋',
        'recreo_libre':     '☕',
    }
    _BLOQUE_COLOR = {
        'ingreso':          '#607D8B',
        'clase':            '#1B5E20',
        'vigilancia_recreo':'#E65100',
        'atencion_ppff':    '#1A237E',
        'planificacion':    '#00695C',
        'recreo_libre':     '#9E9E9E',
    }
    _EVENTO_ICONO = {'feriado':'🚫','acto_civico':'🎌','institucional':'📣','curricular':'📖'}

    with tab_dia:
        import pytz
        from datetime import timedelta as _td
        bolivia_tz = pytz.timezone('America/La_Paz')
        _hoy_local = _dt.now(bolivia_tz)
        
        _hora_actual = _hoy_local.hour + _hoy_local.minute / 60.0
        _is_primary = "primaria" in str(ss.get("nivel_grado", "")).lower() or "primaria" in str(ss.get("nivel", "")).lower()
        if (_is_primary and _hora_actual >= 13.0) or (not _is_primary and _hora_actual >= 13.66):
            _hoy_local += _td(days=1)
            if _hoy_local.weekday() >= 5:
                _hoy_local += _td(days=(7 - _hoy_local.weekday()))
                
        _hoy = _hoy_local
        _dow = _hoy.weekday()   # 0=lun … 6=dom
        _fecha_iso = _hoy.strftime('%Y-%m-%d')
        _nombre_hoja = ss.get("usuario_hoja", "")

        # ── AID Day Header ────────────────────────────────────────────────────
        # School year starts Feb 2, 2026 → derive school week & día escolar
        _SCHOOL_START = _dt(2026, 2, 2).date()
        _school_delta = max(0, (_hoy.date() - _SCHOOL_START).days)
        _school_week  = _school_delta // 7 + 1
        _dia_escolar  = min(66, _school_delta * 5 // 7 + 1)

        if _dow < 5:
            _dia_es = _DIAS_HOY[_dow]
            _dia_label = f"{_DIAS_LABEL[_dia_es]}, {_hoy.day} de {_MESES_LABEL[_hoy.month]}"
        else:
            _dia_es = 'lunes'
            _dia_label = f"Fin de semana · {_hoy.day} de {_MESES_LABEL[_hoy.month]}"

        # ── Cargar datos ─────────────────────────────────────────────────────
        _bloques_horario  = get_horario_dia(_nombre_hoja, _dia_es) if _nombre_hoja else []
        _eventos_cal      = get_eventos_fecha(_fecha_iso)
        _actividades_crono= get_actividades_fecha(_fecha_iso, _nombre_hoja)
        _logs_dia         = get_logs_dia(_fecha_iso, _nombre_hoja) if _nombre_hoja else {}
        _vigilancias      = get_vigilancias()
        _semana_num       = _school_week  # school week number (not ISO)

        # Build header % from already-loaded data
        _pre_cerrados = sum(
            1 for _bh in _bloques_horario
            if _logs_dia.get(f"{_bh['dia_semana']}|{_bh['hora_inicio']}|{_bh['tipo_bloque']}", {}).get('cerrado', 0)
        )
        _pre_total = len(_bloques_horario)
        _pre_pct   = int(_pre_cerrados / _pre_total * 100) if _pre_total else 0

        _hdr_left, _hdr_right = st.columns([5, 1])
        with _hdr_left:
            st.markdown(
                f"""<div class="aid-day-header">
                  <div class="aid-day-header__left">
                    <div class="aid-day-header__title">Tu Ruta del Día</div>
                    <div class="aid-day-header__sub">{_dia_label}</div>
                  </div>
                  <div class="aid-day-header__right">
                    <div class="aid-day-header__meta">
                      <div class="aid-day-header__cycle">Semana {_school_week}</div>
                      <div class="aid-day-header__daycount">Día {_dia_escolar} / 66</div>
                    </div>
                    <div class="aid-day-header__pct">{_pre_pct}%</div>
                  </div>
                </div>""",
                unsafe_allow_html=True
            )
        # (El control de reseteo para administradores ahora está en el panel de administración general)

        if _dow >= 5:
            st.info("Hoy no hay clases. Aquí verás el próximo lunes cuando sea día hábil.")

        # ── Alerta de feriado ─────────────────────────────────────────────────
        _feriados = [e for e in _eventos_cal if e['tipo'] == 'feriado']
        if _feriados:
            st.error(f"🚫 **FERIADO:** {_feriados[0]['nombre_evento']}", icon="🚫")

        # ── Eventos institucionales del día ───────────────────────────────────
        _eventos_no_feriado = [e for e in _eventos_cal if e['tipo'] != 'feriado' and e['tipo'] != 'curricular']
        if _eventos_no_feriado:
            for _ev in _eventos_no_feriado:
                _ico = _EVENTO_ICONO.get(_ev['tipo'], '📣')
                st.info(f"{_ico} **{_ev['nombre_evento']}** — {_ev.get('descripcion','')}", icon=_ico)

        # ── Actividades de cronograma relevantes para este docente ────────────
        if _actividades_crono:
            with st.expander(f"⚡ {len(_actividades_crono)} actividad(es) institucional(es) hoy"):
                for _act in _actividades_crono:
                    _h = f"{_act.get('hora_inicio','')}–{_act.get('hora_fin','')}" if _act.get('hora_inicio') else ""
                    st.markdown(f"- **{_h}** {_act['actividad']} _{_act.get('a_cargo_de','')}_")

        # ── Sin horario cargado ───────────────────────────────────────────────
        if not _bloques_horario and not _nombre_hoja:
            st.warning("Tu horario aún no ha sido cargado por el administrador. Consulta a Ruddy.")
        elif not _bloques_horario:
            st.info(f"No hay bloques de horario para {_DIAS_LABEL.get(_dia_es,'hoy')}.")
        else:
            # ── Barra de progreso del día ─────────────────────────────────────
            _total_bloques   = len(_bloques_horario)
            _bloques_cerrados    = sum(
                1 for _b in _bloques_horario
                if _logs_dia.get(f"{_b['dia_semana']}|{_b['hora_inicio']}|{_b['tipo_bloque']}", {}).get('cerrado', 0)
            )
            _bloques_completados = sum(
                1 for _b in _bloques_horario
                if _logs_dia.get(f"{_b['dia_semana']}|{_b['hora_inicio']}|{_b['tipo_bloque']}", {}).get('completado', 0)
            )
            _pct = int(_bloques_cerrados / _total_bloques * 100) if _total_bloques else 0
            _progreso_txt = (
                f"**{_bloques_cerrados} cerrados · {_bloques_completados} marcados · "
                f"{_total_bloques} total — {_pct}% del día cerrado**"
            )
            st.progress(_pct / 100, text=_progreso_txt)
            st.divider()

            # ── Renderizar cada bloque ────────────────────────────────────────
            _now_hhmm = _hoy.strftime('%H:%M')

            for _b in _bloques_horario:
                _tipo      = _b['tipo_bloque']
                _hora_ini  = _b['hora_inicio']
                _hora_fin  = _b.get('hora_fin', '')
                _hora      = f"{_hora_ini} – {_hora_fin}"
                _color     = _BLOQUE_COLOR.get(_tipo, '#333')
                _icono     = _BLOQUE_ICONO.get(_tipo, '📌')
                _bkey      = f"{_b['dia_semana']}|{_hora_ini}|{_tipo}"
                _log       = _logs_dia.get(_bkey, {})
                _done      = bool(_log.get('completado', 0))
                _cerrado   = bool(_log.get('cerrado', 0))
                _notas_log = _log.get('notas') or ''
                _is_past   = bool(_hora_fin) and _hora_fin <= _now_hhmm
                _is_active = (bool(_hora_ini) and bool(_hora_fin)
                              and _hora_ini <= _now_hhmm < _hora_fin)

                # Build human label
                if _tipo == 'clase':
                    _nivel = f" — {_b['nivel_grado']}" if _b.get('nivel_grado') else ""
                    _label = f"{(_b.get('materia') or '').title()}{_nivel}"
                elif _tipo == 'vigilancia_recreo':
                    _zona = None
                    _nh_upp = _nombre_hoja.upper()
                    for k_name, v_zona in _vigilancias.items():
                        if k_name in _nh_upp or _nh_upp in k_name:
                            _zona = v_zona
                            break
                    _label = f"Guardia de Recreo · 📍 {_zona}" if _zona else f"Guardia de Recreo · {_b.get('ubicacion') or 'Patio'}"
                else:
                    _label = {
                        'ingreso':       'Horario de Ingreso',
                        'atencion_ppff': 'Atención a Padres de Familia',
                        'planificacion': 'Planificación / API',
                        'recreo_libre':  'Recreo',
                    }.get(_tipo, (_b.get('valor_original') or _tipo.replace('_', ' ').title()))

                # ── LOCKED STATE: compact read-only row (no HTML div) ─────────
                if _cerrado:
                    _rc1, _rc2, _rc3 = st.columns([1.8, 5.5, 2.2], gap="small")
                    with _rc1:
                        st.markdown(
                            f"<span class='aid-time' style='color:#3E4252'>{_hora}</span>",
                            unsafe_allow_html=True
                        )
                    with _rc2:
                        _note_suffix = f" · {_notas_log}" if _notas_log else ""
                        st.markdown(
                            f"<span style='color:#3E4252;text-decoration:line-through'>"
                            f"{_icono} {_label}</span>"
                            f"<span style='color:#2E313C;font-size:0.72rem'>{_note_suffix}</span>",
                            unsafe_allow_html=True
                        )
                    with _rc3:
                        st.markdown(
                            "<span class='aid-badge aid-badge--closed'>✓ Cerrado</span>",
                            unsafe_allow_html=True
                        )
                        if st.button("Reabrir", key=f"reab_{_fecha_iso}_{_bkey.replace('|','_')}", use_container_width=True):
                            reabrir_bloque(_fecha_iso, _nombre_hoja, _b)
                            st.rerun()
                    st.markdown(
                        "<div style='height:1px;background:#1A1D23;margin:1px 0'></div>",
                        unsafe_allow_html=True
                    )
                    continue

                # ── OPEN STATE ─────────────────────────────────────────────
                _badge_html = ""
                if _is_active:
                    _badge_html = '<span class="aid-badge aid-badge--active">● EN CURSO</span>'
                elif _is_past and not _done:
                    _badge_html = '<span class="aid-badge aid-badge--pending">⚠ PENDIENTE</span>'

                # ── Top-down objective sync (class blocks) ──────────────────
                _sesion_id = None
                _objetivos = []
                if _tipo == 'clase':
                    _mat_db = (_b.get('materia') or '').strip()
                    _niv_db = (_b.get('nivel_grado') or ss.grado_nivel).strip()
                    _sesion_id = get_or_create_sesion_diaria(
                        fecha=_fecha_iso, semana=_semana_num,
                        materia=_mat_db, grado=_niv_db,
                        hora_inicio=_hora_ini, hora_fin=_hora_fin
                    )
                    _objetivos = get_micro_objetivos(_sesion_id)
                    # Auto-inject from WeeklyPlan if none exist yet
                    if not _objetivos:
                        _plan_objs = get_objetivos_semana_materia(_semana_num, _mat_db, _niv_db)
                        if _plan_objs:
                            guardar_micro_objetivos(
                                _sesion_id,
                                [{"texto": o["texto"], "depende_de": o.get("depende_de")}
                                 for o in _plan_objs],
                                origen_semana=_semana_num
                            )
                            _objetivos = get_micro_objetivos(_sesion_id)

                _notas_key = f"notas_{_fecha_iso}_{_bkey.replace('|','_')}"
                # Only checkable tasks (not the 🎯 display-only objective line) count
                _checkable = [o for o in _objetivos if not o['texto'].startswith("🎯")]
                _has_checked = (
                    any(o['completado'] for o in _checkable)
                    if _checkable else True
                )

                # ── 3-column compact row layout ─────────────────────────────
                with st.container(border=True):
                    _c1, _c2, _c3 = st.columns([1.8, 5.5, 2.2], gap="small")

                    # Col 1 — Monospace time
                    with _c1:
                        st.markdown(
                            f"<div class='aid-time' style='padding-top:3px'>{_hora}</div>",
                            unsafe_allow_html=True
                        )

                    # Col 2 — Subject · badge · objectives · observations
                    with _c2:
                        st.markdown(
                            f"<span style='font-weight:700;color:#E8EAF0'>{_icono} {_label}</span>"
                            f"&nbsp;{_badge_html}",
                            unsafe_allow_html=True
                        )
                        if _tipo == 'clase':
                            if _objetivos:
                                for _obj in _objetivos:
                                    _txt = _obj['texto']
                                    # Inferred objective line — display-only, no checkbox
                                    if _txt.startswith("🎯"):
                                        st.markdown(
                                            f"<div style='font-size:0.75rem;color:#6B7185;"
                                            f"font-style:italic;padding:2px 0 3px 0'>{_txt}</div>",
                                            unsafe_allow_html=True
                                        )
                                    else:
                                        _chk_obj = st.checkbox(
                                            _txt,
                                            value=bool(_obj['completado']),
                                            key=f"obj_{_obj['id']}_{_fecha_iso}"
                                        )
                                        if _chk_obj != bool(_obj['completado']):
                                            marcar_objetivo(_obj['id'], _chk_obj)
                                            st.rerun()
                            else:
                                st.markdown(
                                    '<div class="aid-sync-state">📋 Sin objetivos en Plan Semanal</div>',
                                    unsafe_allow_html=True
                                )
                        # Compact single-line observations field
                        st.text_input(
                            "obs", value=_notas_log,
                            placeholder="Observación…",
                            key=_notas_key, label_visibility="collapsed"
                        )

                    # Col 3 — Actions
                    with _c3:
                        _btn_lbl = "✅ Finalizar" if _tipo == 'clase' else "✅ Cerrar"
                        if st.button(
                            _btn_lbl,
                            key=f"cerrar_{_fecha_iso}_{_bkey.replace('|','_')}",
                            type="primary" if _is_past else "secondary",
                            use_container_width=True,
                            disabled=not _nombre_hoja or not _has_checked,
                        ):
                            _notas_save = ss.get(_notas_key, _notas_log)
                            cerrar_bloque(_fecha_iso, _nombre_hoja, _b, _notas_save)
                            st.rerun()
                        _new_done = st.checkbox(
                            "✓ Listo", value=_done,
                            key=f"bloque_{_fecha_iso}_{_bkey.replace('|','_')}",
                        )
                        if _new_done != _done and _nombre_hoja:
                            marcar_bloque_diario(_fecha_iso, _nombre_hoja, _b, _new_done)
                            st.rerun()

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 2 — DEUDA ACADÉMICA + PLAN DE AMORTIGUACIÓN (mantenido del sistema anterior)
    # ══════════════════════════════════════════════════════════════════════════
    with tab_deuda:
        bucle_tabs = st.tabs([
            "📅 Registrar Clase",
            "✅ Lista de Clase Activa",
            "💸 Deuda Académica",
            "🔄 Plan de Amortiguación",
        ])

        with bucle_tabs[0]:
            st.markdown("#### Abre una nueva sesión de clase y extrae sus micro-objetivos")
            col_r1, col_r2, col_r3 = st.columns(3)
            with col_r1:
                reg_semana  = st.number_input("Semana N°", min_value=1, max_value=40, value=1, key="reg_semana")
                reg_materia = st.text_input("Materia", value=ss.get("materia", "Lenguaje"), key="reg_materia")
                reg_grado   = st.text_input("Grado", value=ss.grado_nivel, key="reg_grado")
            with col_r2:
                reg_tema        = st.text_input("Tema de la sesión", placeholder="Ej. Las modalidades oracionales", key="reg_tema")
                reg_hora_inicio = st.text_input("Hora inicio (HH:MM)", placeholder="08:00", key="reg_hora_inicio")
                reg_hora_fin    = st.text_input("Hora fin   (HH:MM)", placeholder="08:45", key="reg_hora_fin")
            with col_r3:
                st.markdown("**Plan de clase base (opcional)**")
                st.caption("Pega el plan generado en zona Semanal para extraer micro-objetivos automáticamente.")
                reg_plan_texto = st.text_area("Plan de clase", height=150, key="reg_plan_texto",
                                              placeholder="Pega aquí el contenido del plan de clase…")

            if st.button("📌 Registrar sesión y extraer micro-objetivos", type="primary",
                         disabled=not (reg_tema and reg_materia), key="btn_registrar_sesion"):
                with st.spinner("Registrando sesión y extrayendo micro-objetivos…"):
                    sesion_id = crear_sesion(
                        semana=int(reg_semana), materia=reg_materia, grado=reg_grado,
                        tema=reg_tema, hora_inicio=reg_hora_inicio, hora_fin=reg_hora_fin,
                    )
                    if reg_plan_texto.strip():
                        resultado_mo = generar_con_gemini(
                            prompt_filename="Motor_MicroObjetivos.txt",
                            variables={
                                "subject": reg_materia, "grado_nivel": reg_grado,
                                "tema": reg_tema, "plan_clase": reg_plan_texto[:3000],
                            },
                            expect_json=True, required_fields=["micro_objetivos"],
                        )
                        if resultado_mo and "micro_objetivos" in resultado_mo:
                            guardar_micro_objetivos(sesion_id, resultado_mo["micro_objetivos"],
                                                    origen_semana=int(reg_semana))
                            ss["bucle_sesion_activa"] = sesion_id
                            st.success(f"✅ Sesión registrada con {len(resultado_mo['micro_objetivos'])} micro-objetivos. "
                                       "Ve a **Lista de Clase Activa** para marcarlos.")
                        else:
                            ss["bucle_sesion_activa"] = sesion_id
                            st.info("Sesión registrada. Agrega micro-objetivos manualmente en **Lista de Clase Activa**.")
                    else:
                        ss["bucle_sesion_activa"] = sesion_id
                        st.success("✅ Sesión registrada. Añade micro-objetivos manualmente.")

            st.divider()
            st.markdown("#### Historial de sesiones")
            sesiones_hist = get_sesiones()
            if sesiones_hist:
                for s in sesiones_hist[:10]:
                    label = f"Sem {s['semana']} — {s['materia']} — {s['tema']} ({s['fecha']})"
                    if st.button(label, key=f"sel_sesion_{s['id']}"):
                        ss["bucle_sesion_activa"] = s["id"]
                        st.toast(f"Sesión activa: {s['tema']}", icon="✅")
            else:
                st.info("Aún no hay sesiones registradas.")

        with bucle_tabs[1]:
            sesion_activa_id = ss.get("bucle_sesion_activa")
            if not sesion_activa_id:
                st.info("👈 Ve a **Registrar Clase** y selecciona o crea una sesión.")
            else:
                sesion = get_sesion(sesion_activa_id)
                if not sesion:
                    st.warning("Sesión no encontrada. Selecciona otra.")
                else:
                    mins = minutos_para_fin_clase(sesion.get("hora_fin", ""))
                    if mins is not None:
                        if 0 < mins <= 5:
                            st.warning(f"⏰ **{mins} minutos restantes** para el final de la clase.", icon="⏰")
                        elif mins <= 0:
                            st.error("🔔 El bloque de clase ha terminado. Marca los objetivos y cierra la sesión.", icon="🔔")
                        else:
                            st.info(f"🕐 Tiempo restante: **{mins} minutos**", icon="🕐")

                    st.markdown(f"### ✅ {sesion['tema']}")
                    st.caption(f"Sem {sesion['semana']} · {sesion['materia']} · {sesion['grado']} · "
                               f"{sesion.get('hora_inicio','')}–{sesion.get('hora_fin','')}")
                    st.divider()

                    objetivos = get_micro_objetivos(sesion_activa_id)
                    if not objetivos:
                        st.info("Esta sesión no tiene micro-objetivos. Añade uno manualmente:")
                        nuevo_obj = st.text_input("Nuevo micro-objetivo", key="nuevo_obj_manual")
                        if st.button("➕ Agregar", key="btn_add_obj") and nuevo_obj.strip():
                            guardar_micro_objetivos(sesion_activa_id,
                                                    [{"texto": nuevo_obj, "depende_de": None}],
                                                    origen_semana=sesion.get("semana"))
                            st.rerun()
                    else:
                        completados = 0
                        for obj in objetivos:
                            col_chk, col_txt = st.columns([0.07, 0.93])
                            with col_chk:
                                checked = st.checkbox("", value=bool(obj["completado"]), key=f"obj_{obj['id']}")
                                if checked != bool(obj["completado"]):
                                    marcar_objetivo(obj["id"], checked)
                            with col_txt:
                                dep_txt = f" *(depende del objetivo #{obj['depende_de']})*" if obj.get("depende_de") else ""
                                if obj["completado"]:
                                    st.markdown(f"~~{obj['texto']}~~{dep_txt}")
                                    completados += 1
                                else:
                                    st.markdown(f"**{obj['texto']}**{dep_txt}")

                        pct = int(completados / len(objetivos) * 100)
                        st.progress(pct / 100, text=f"{completados}/{len(objetivos)} objetivos completados ({pct}%)")
                        col_all, col_none = st.columns(2)
                        with col_all:
                            if st.button("✅ Marcar todos", key="btn_all"):
                                marcar_multiples([o["id"] for o in objetivos], True); st.rerun()
                        with col_none:
                            if st.button("⬜ Desmarcar todos", key="btn_none"):
                                marcar_multiples([o["id"] for o in objetivos], False); st.rerun()

        with bucle_tabs[2]:
            st.markdown("#### 💸 Deuda Académica Acumulada")
            st.caption("Solo los objetivos NO completados.")
            filtro_mat = st.text_input("Filtrar por materia (vacío = todas)", key="deuda_filtro_mat")
            resumen = get_resumen_deuda(filtro_mat if filtro_mat.strip() else None)
            col_d1, col_d2, col_d3 = st.columns(3)
            col_d1.metric("Objetivos pendientes", resumen["total_pendientes"])
            col_d2.metric("Críticos (con dependencias)", resumen["bloqueados_criticos"],
                          delta="⚠️ Bloquean avance" if resumen["bloqueados_criticos"] > 0 else None,
                          delta_color="inverse")
            col_d3.metric("Semanas afectadas", len(resumen["semanas_afectadas"]))
            if resumen["deuda_detalle"]:
                import pandas as pd
                df_deuda = pd.DataFrame(resumen["deuda_detalle"])[
                    ["semana", "materia", "tema", "texto", "depende_de"]
                ]
                df_deuda.columns = ["Semana", "Materia", "Tema", "Objetivo Pendiente", "Depende de ID"]
                st.dataframe(df_deuda, use_container_width=True, hide_index=True)
            else:
                st.success("🎉 ¡Sin deuda académica! Todos los objetivos están completados.")

        with bucle_tabs[3]:
            st.markdown("#### 🔄 Generar Plan de Amortiguación (Día Buffer)")
            st.caption("PRIA tomará la deuda pendiente y generará un bloque de recuperación para la semana siguiente.")
            col_b1, col_b2 = st.columns(2)
            with col_b1:
                buf_semana_actual     = st.number_input("Semana actual", min_value=1, max_value=40, value=1, key="buf_semana_actual")
                buf_semanas_restantes = st.number_input("Semanas hábiles restantes en el trimestre",
                                                        min_value=0, max_value=20, value=8, key="buf_semanas_restantes")
            with col_b2:
                buf_materia  = st.text_input("Materia", value=ss.get("materia", ""), key="buf_materia")
                buf_objetivo = st.text_area("Macro-objetivo trimestral (copia del PDC)",
                                            height=80, key="buf_objetivo",
                                            placeholder="Ej. Desarrollar competencias comunicativas…")

            if st.button("⚡ Generar Plan de Amortiguación", type="primary", key="btn_gen_buffer"):
                deuda     = get_deuda_academica(buf_materia if buf_materia.strip() else None,
                                                semana_hasta=int(buf_semana_actual))
                bloqueados = get_dependencias_bloqueadas(buf_materia if buf_materia.strip() else None)
                if not deuda:
                    st.success("✅ No hay deuda académica. ¡No se necesita día buffer!")
                else:
                    deuda_texto = "\n".join(
                        [f"- Sem {d['semana']} | {d['materia']} | {d['tema']}: {d['texto']}" for d in deuda])
                    dep_texto = "\n".join(
                        [f"- Obj '{d['texto']}' depende del objetivo ID {d['depende_de']}" for d in bloqueados]
                    ) or "Ninguna dependencia crítica detectada."
                    with st.spinner("Motor_Recalibracion procesando el delta…"):
                        resultado_buf = generar_con_gemini(
                            prompt_filename="Motor_Recalibracion.txt",
                            variables={
                                "teacher_name":            ss.get("teacher_name", "Docente"),
                                "subject":                 buf_materia or "General",
                                "grado_nivel":             ss.grado_nivel,
                                "semana_actual":           str(buf_semana_actual),
                                "semana_siguiente":        str(int(buf_semana_actual) + 1),
                                "semanas_restantes":       str(buf_semanas_restantes),
                                "objetivo_trimestral":     buf_objetivo or "No especificado",
                                "deuda_academica":         deuda_texto,
                                "dependencias_bloqueadas": dep_texto,
                            },
                            expect_json=True, required_fields=["bloque_recuperacion"],
                        )
                    if resultado_buf:
                        guardar_plan_buffer(
                            semana_buffer=int(buf_semana_actual) + 1,
                            materia=buf_materia or "General",
                            grado=ss.grado_nivel,
                            contenido_json=resultado_buf,
                        )
                        if resultado_buf.get("alerta_condensacion"):
                            st.error(f"⚠️ **ALERTA DE CONDENSACIÓN:** {resultado_buf.get('mensaje_alerta','')}")
                        st.success(f"✅ Plan buffer generado para Semana {int(buf_semana_actual)+1}")
                        st.info(resultado_buf.get("resumen_ejecutivo", ""))
                        bloque = resultado_buf.get("bloque_recuperacion", {})
                        st.markdown(f"### 📦 {bloque.get('titulo','Bloque de Recuperación')}")
                        st.markdown(f"**Duración estimada:** {bloque.get('duracion_estimada_min','?')} min")
                        st.markdown(f"**Estrategia:** {bloque.get('estrategia','')}")
                        st.markdown("**Micro-objetivos de recuperación:**")
                        for mo in bloque.get("micro_objetivos_recuperacion", []):
                            prioridad = "🔴" if mo.get("prioridad") == "alta" else "🟡"
                            st.markdown(f"{prioridad} {mo['texto']}")
                        if resultado_buf.get("objetivos_diferidos"):
                            with st.expander("📌 Objetivos diferidos al próximo trimestre"):
                                for d in resultado_buf["objetivos_diferidos"]:
                                    st.markdown(f"- **{d['texto']}** — *{d.get('razon','')}*")

            st.divider()
            st.markdown("#### Historial de planes de amortiguación")
            _buf_mat_filter = ss.get("buf_materia", "").strip() or None
            planes = get_planes_buffer(_buf_mat_filter)
            if planes:
                for p in planes:
                    icono = "⚠️" if p["alerta_condensacion"] else "📦"
                    with st.expander(f"{icono} Semana {p['semana_buffer']} — {p['materia']} — {p['creado_en'][:10]}"):
                        st.markdown(p.get("resumen", ""))
                        bloque = p["contenido_json"].get("bloque_recuperacion", {})
                        for mo in bloque.get("micro_objetivos_recuperacion", []):
                            st.markdown(f"- {mo['texto']}")
            else:
                st.info("Aún no hay planes de amortiguación generados.")

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 3 — MIS COMISIONES
    # ══════════════════════════════════════════════════════════════════════════
    with tab_comisiones:
        _nombre_hoja_com = ss.get("usuario_hoja", "")
        _comisiones = get_comisiones_docente(_nombre_hoja_com) if _nombre_hoja_com else []
        if not _comisiones:
            st.info("No se encontraron comisiones asignadas. El administrador debe importar el archivo de comisiones.")
        else:
            for _com in _comisiones:
                with st.container(border=True):
                    st.markdown(f"### 🏛️ {_com['comision']}")
                    if _com.get('funciones'):
                        with st.expander("Ver funciones de esta comisión"):
                            for _fn in _com['funciones'].split('\n'):
                                if _fn.strip():
                                    st.markdown(f"- {_fn.strip()}")

        st.divider()
        st.markdown("#### Próximos actos cívicos y eventos institucionales")
        import pandas as pd
        _hoy_iso = _dt.now().strftime('%Y-%m-%d')
        _proximos = get_eventos_rango(_hoy_iso, f"{_dt.now().year}-12-31")
        _relevantes = [e for e in _proximos if e['tipo'] in ('acto_civico', 'institucional')][:10]
        if _relevantes:
            for _ev in _relevantes:
                _ico = _EVENTO_ICONO.get(_ev['tipo'], '📣')
                st.markdown(f"{_ico} **{_ev['fecha']}** — {_ev['nombre_evento']}")
                if _ev.get('descripcion'):
                    st.caption(_ev['descripcion'])
        else:
            st.info("No hay eventos próximos en el calendario.")


# ═════════════════════════════════════════════════════════════════════════════
# ZONA SEMANAL — Plan de Clase, Diapositivas, Ficha, Pop Quiz, Exportar
# ═════════════════════════════════════════════════════════════════════════════
elif zona == "📅  Semanal":
    tab_clase, tab_diap, tab_ficha, tab_quiz, tab_export = st.tabs([
        "🚀 Plan de Clase",
        "🖼️ Diapositivas",
        "🎮 Ficha",
        "📝 Pop Quiz",
        "📤 Exportar",
    ])

    # ── PLAN DE CLASE ─────────────────────────────────────────────────────────
    with tab_clase:
        st.markdown("Crea el plan de clase de 45 min.")
        with st.container(border=True):
            st.markdown("#### 🎯 Configuración de la Lección")
            _tema_opts = temas_reales if temas_reales else ["(carga el libro en el panel lateral)"]
            tema_final = st.selectbox("📍 Tema:", _tema_opts, key="sel_tema_m1")

            tema_hash_nuevo = _topic_hash(tema_final)
            if tema_hash_nuevo != ss.tema_hash:
                ss.res_m1a_prev              = ss.res_m1a
                ss.res_m1a                   = None
                ss.res_m1b                   = None
                ss.res_m1c                   = None
                ss.tema_hash                 = tema_hash_nuevo
                ss.tema_activo               = tema_final
                ss.mostrar_adaptaciones_prev = False

            pags_tb_auto  = paginas_temas.get(tema_final, "—")
            pags_sb_temas = datos_sb.get("paginas_temas", {}) if datos_sb else {}
            pags_sb_auto  = pags_sb_temas.get(tema_final, "—") if usar_sb == "Sí" else "N/A"

            c_pags1, c_pags2 = st.columns(2)
            with c_pags1:
                st.markdown("**📄 Páginas TextBook:**")
                st.markdown(
                    f"<div style='background:rgba(200,230,201,0.35);border:2px dashed #4CAF50;"
                    f"border-radius:6px;padding:8px 14px;color:#1B5E20;font-weight:800;'>📖 {pags_tb_auto}</div>",
                    unsafe_allow_html=True)
            with c_pags2:
                st.markdown("**📄 Páginas Student Book:**")
                st.markdown(
                    f"<div style='background:rgba(200,230,201,0.35);border:2px dashed #4CAF50;"
                    f"border-radius:6px;padding:8px 14px;color:#1B5E20;font-weight:800;'>📖 {pags_sb_auto}</div>",
                    unsafe_allow_html=True)

            c_m1a, c_m1b_col = st.columns(2)
            with c_m1a:
                objetivo_gen = st.text_input("🎯 Objetivo General (opcional):", value="",
                                             placeholder="Déjalo vacío para que la IA lo defina…", key="obj_m1")
            with c_m1b_col:
                personaje_genero = st.radio("🧒 Género del personaje visual:",
                                            ["masculino", "femenino", "ambos"],
                                            horizontal=True, key="genero_m1")
            user_sug_m1 = st.text_area("💬 Sugerencias:", value="",
                                        placeholder="Indicaciones especiales para esta lección…",
                                        height=68, key="sug_m1")

        if ss.res_m1a_prev:
            col_prev, col_gen = st.columns([1, 2])
            with col_prev:
                label_prev = ("🙈 Ocultar adaptaciones anteriores"
                              if ss.mostrar_adaptaciones_prev else "👁️ Ver adaptaciones anteriores")
                if st.button(label_prev, use_container_width=True):
                    ss.mostrar_adaptaciones_prev = not ss.mostrar_adaptaciones_prev
            if ss.mostrar_adaptaciones_prev:
                with st.container(border=True):
                    st.markdown("#### 📋 Adaptaciones de la lección anterior")
                    prev = ss.res_m1a_prev
                    adapt_prev = prev.get("tabla_adaptaciones_clase", [])
                    if adapt_prev:
                        st.dataframe(adapt_prev, use_container_width=True, hide_index=True)
                    dua_prev = prev.get("dua_neuroinclusion", [])
                    if dua_prev:
                        st.markdown("**Directrices DUA anteriores:**")
                        for d in forzar_lista(dua_prev):
                            st.markdown(f"- {d}")
                    st.caption("Revisa estas adaptaciones. Si ya son suficientes, no necesitas generar un nuevo plan.")
            with col_gen:
                btn_generar = st.button("✨ Generar Nuevo Plan de Clase", use_container_width=True)
        else:
            btn_generar = st.button("✨ Generar Plan de Clase", use_container_width=True)

        if btn_generar:
            conceptos_clave = []; inteligencias_sug = []; palabras_clave = []
            if ss.res_m0a and "unidad_sintetizada" in ss.res_m0a:
                for t in ss.res_m0a["unidad_sintetizada"].get("temas_desarrollados", []):
                    if t.get("nombre", "") == tema_final:
                        conceptos_clave   = forzar_lista(t.get("conceptos_clave") or t.get("conceptos") or [])
                        inteligencias_sug = forzar_lista(t.get("inteligencias_sugeridas") or t.get("inteligencias") or [])
                        palabras_clave    = conceptos_clave
                        break

            contenido_real_tema       = contenido_temas.get(tema_final, "")
            ss.tema_activo            = tema_final
            ss.conceptos_activos      = conceptos_clave
            ss.palabras_clave_activas = palabras_clave
            ss.contenido_tema_activo  = contenido_real_tema

            contexto_prev = ""
            if ss.res_m1a_prev:
                contexto_prev = (
                    "CONTEXTO DE MEJORA: Ya existe un plan para otra lección de esta unidad. "
                    "Estudia las adaptaciones previas y MEJÓRALAS o VARÍA sin repetirlas exactamente.\n"
                    "Adaptaciones previas:\n"
                    + json.dumps({
                        k: ss.res_m1a_prev.get(k)
                        for k in ["mapa_cognitivo", "dua_neuroinclusion", "tabla_adaptaciones_clase"]
                        if k in ss.res_m1a_prev
                    }, ensure_ascii=False)[:800]
                )

            with st.spinner("Diseñando la estrategia de clase…"):
                vars_m1a = {
                    "tema_clase":              tema_final,
                    "conceptos_clave":         str(conceptos_clave) or "Derivar del contenido del libro",
                    "palabras_clave":          str(palabras_clave)  or tema_final,
                    "inteligencias_sugeridas": str(inteligencias_sug) or "Lingüística, Visual-espacial",
                    "contenido_curricular":    contenido_real_tema[:1500] or "No disponible",
                    "diagnosticos":            DIAGNOSTICOS or "No especificado",
                    "objetivo_general":        objetivo_gen or "No especificado",
                    "PAG_TB":                  pags_tb_auto,
                    "PAG_SB":                  pags_sb_auto,
                    "user_suggestions":        user_sug_m1 or "Ninguna",
                    "contexto_leccion_previa": contexto_prev,
                }
                ss.res_m1a  = generar_con_gemini("Motor_M1a.txt", vars_m1a, expect_json=True)
                ss.res_m1b  = None
                ss.res_m1c  = None
                ss.leccion_index += 1
                ss.mostrar_adaptaciones_prev = False

        if ss.res_m1a:
            res     = ss.res_m1a
            lec_num = ss.leccion_index
            st.success(f"✅ Plan de clase #{lec_num} — **{ss.tema_activo}**")
            c_izq, c_der = st.columns([1.6, 1], gap="large")
            with c_izq:
                st.markdown("### 🎯 Verbos Operativos (Bloom)")
                for v in forzar_lista(res.get("mapa_cognitivo", {}).get("verbos", [])):
                    st.checkbox(v, value=True, key=f"bloom_{lec_num}_{v}")
                st.markdown("### 🧩 Inteligencias Múltiples")
                im_data = res.get("inteligencias_multiples", [])
                if im_data:
                    st.dataframe(im_data, use_container_width=True, hide_index=True)
                st.markdown("### ⏱️ Secuencia Didáctica *(45 min)*")
                for bloque in res.get("secuencia_didactica", {}).get("bloques", []):
                    nb = bloque.get("nombre", "Bloque"); dur = bloque.get("duracion", "?")
                    with st.expander(f"**⏩ {nb}** — {dur} min"):
                        st.write(bloque.get("objetivo", ""))
                        if "nota" in bloque: st.info(f"💡 {bloque['nota']}")
            with c_der:
                st.markdown("### 🛡️ Directrices DUA")
                for d in forzar_lista(res.get("dua_neuroinclusion", [])):
                    st.markdown(f"- {d}")
                st.markdown("### ♿ Adaptaciones por Diagnóstico")
                adapt = res.get("tabla_adaptaciones_clase", [])
                if adapt:
                    st.dataframe(adapt, use_container_width=True, hide_index=True)
                st.markdown("### 👥 Perfil del Aula")
                st.info(res.get("perfil_aula_resumido", "—"))
            st.divider()
            st.download_button(
                "⬇️ Descargar Plan de Clase (.html)",
                data=generar_html_plan_clase(ss.res_m1a, ss.tema_activo, DIAGNOSTICOS),
                file_name=f"LasP_PlanClase_{_dt.now().strftime('%Y%m%d')}.html",
                mime="text/html", use_container_width=True,
            )

    # ── DIAPOSITIVAS ──────────────────────────────────────────────────────────
    with tab_diap:
        if not ss.res_m1a:
            st.warning("⚠️ Primero genera un **Plan de Clase** en la pestaña anterior.")
        else:
            tema_diap    = ss.tema_activo
            pags_tb_d    = paginas_temas.get(tema_diap, "—")
            pags_sb_d    = (datos_sb.get("paginas_temas", {}).get(tema_diap, "—")
                            if datos_sb and usar_sb == "Sí" else "N/A")
            pkw = str(ss.palabras_clave_activas) if ss.palabras_clave_activas else tema_diap

            personaje_diap = st.radio("🧒 Género del personaje:", ["masculino", "femenino", "ambos"],
                                       horizontal=True, key="genero_diap")
            user_sug_diap  = st.text_area("💬 Sugerencias:", value="", height=68, key="sug_diap",
                                           placeholder="Indicaciones para las diapositivas…")

            if st.button("🎨 Generar Diapositivas", use_container_width=True, type="primary", key="btn_gen_diap"):
                with st.spinner("M1b creando diapositivas…"):
                    vars_m1b = {
                        "plan_estrategico_json": json.dumps(ss.res_m1a, ensure_ascii=False),
                        "diagnosticos":          DIAGNOSTICOS or "No especificado",
                        "PAG_TB":                pags_tb_d,
                        "PAG_SB":                pags_sb_d,
                        "palabras_clave":        pkw,
                        "personaje_genero":      personaje_diap,
                        "user_suggestions":      user_sug_diap or "Ninguna",
                    }
                    ss.res_m1b = generar_con_gemini("Motor_M1b.txt", vars_m1b, expect_json=True)

            if ss.res_m1b:
                st.success("✅ Diapositivas generadas.")
                raw    = ss.res_m1b
                slides = (raw if isinstance(raw, list)
                          else raw.get("diapositivas") or raw.get("slides")
                          or (list(raw.values())[0] if isinstance(raw, dict) else []))
                for i, slide in enumerate(slides):
                    if isinstance(slide, dict):
                        titulo_s = slide.get("titulo") or slide.get("title") or f"Slide {i+1}"
                        with st.expander(f"**Slide {i+1}: {titulo_s}**"):
                            st.markdown(f"**📢 Guion:** {slide.get('guion_docente') or slide.get('guion','')}")
                            st.markdown(f"**🖥️ Pantalla:** {slide.get('texto_pantalla') or slide.get('texto','')}")
                            prompt_img = slide.get("prompt_imagen") or slide.get("prompt", "")
                            if prompt_img:
                                st.code(prompt_img, language=None)
                st.caption("Descarga el archivo .pptx desde la pestaña **📤 Exportar**.")

    # ── FICHA GAMIFICADA ──────────────────────────────────────────────────────
    with tab_ficha:
        if not ss.res_m1a:
            st.warning("⚠️ Primero genera un **Plan de Clase** en la pestaña anterior.")
        else:
            pkw_ficha     = str(ss.palabras_clave_activas) if ss.palabras_clave_activas else ss.tema_activo
            user_sug_fich = st.text_area("💬 Sugerencias:", value="", height=68, key="sug_ficha",
                                          placeholder="Indicaciones para la ficha gamificada…")

            if st.button("🎮 Generar Ficha Gamificada", use_container_width=True, type="primary", key="btn_gen_ficha"):
                with st.spinner("M1c creando ficha gamificada…"):
                    vars_m1c = {
                        "plan_estrategico_json": json.dumps(ss.res_m1a, ensure_ascii=False),
                        "diagnosticos":          DIAGNOSTICOS or "No especificado",
                        "conceptos_clave":       str(ss.conceptos_activos),
                        "palabras_clave":        pkw_ficha,
                        "user_suggestions":      user_sug_fich or "Ninguna",
                    }
                    ss.res_m1c = generar_con_gemini("Motor_M1c.txt", vars_m1c, expect_json=True)

            if ss.res_m1c:
                st.success("✅ Ficha generada.")
                raw   = ss.res_m1c
                ficha = raw.get("ficha_trabajo", raw) if isinstance(raw, dict) else {}
                st.markdown(f"## 🏆 {ficha.get('titulo_gancho','—')}")
                st.write(ficha.get("historia_gancho", ""))
                misiones = ficha.get("misiones", {})
                t_or, t_pu, t_so, t_pe, t_li = st.tabs(
                    ["🔮 Oráculo", "🌉 Puente", "🥣 Sopa", "📜 Pergamino", "🎨 Lienzo"])
                with t_or:
                    for q in misiones.get("oraculo", []):
                        st.markdown(f"**{q.get('pregunta','')}**")
                        for op in q.get("opciones", []): st.markdown(f"   • {op}")
                with t_pu:
                    for par in misiones.get("puente", []):
                        st.markdown(f"**{par.get('palabra','')}** → {par.get('significado','')}")
                with t_so:
                    for p in misiones.get("sopa", []): st.markdown(f"- {p}")
                with t_pe:
                    per = misiones.get("pergamino", {})
                    st.markdown(f"_{per.get('frase_con_espacios','')}_")
                    st.markdown(f"**Palabras secretas:** {', '.join(per.get('palabras_secretas',[]))}")
                with t_li:
                    st.markdown(misiones.get("lienzo", ""))
                adapt_m = ficha.get("adaptaciones_por_mision", [])
                if adapt_m:
                    st.divider()
                    st.markdown("**♿ Adaptaciones por Misión:**")
                    st.dataframe(adapt_m, use_container_width=True, hide_index=True)
                st.divider()
                st.download_button(
                    "⬇️ Descargar Ficha Gamificada (.html)",
                    data=generar_html_ficha(ss.res_m1c, ss.tema_activo, theme="aventura").encode("utf-8"),
                    file_name=f"LasP_Ficha_{_dt.now().strftime('%Y%m%d')}.html",
                    mime="text/html", use_container_width=True,
                )

    # ── POP QUIZ ──────────────────────────────────────────────────────────────
    with tab_quiz:
        if not ss.res_m1a:
            st.warning("⚠️ Primero genera un **Plan de Clase** en la pestaña anterior.")
        else:
            user_sug_m2 = st.text_area("💬 Sugerencias:", value="",
                                        placeholder="Indicaciones especiales para la evaluación…",
                                        height=68, key="sug_m2")
            c_btn1, c_btn2 = st.columns(2)
            with c_btn1:
                if st.button("🎲 Generar Pop Quiz DUA", use_container_width=True, key="btn_quiz"):
                    with st.spinner("Creando Pop Quiz…"):
                        vars_m2a = {
                            "plan_estrategico_json": json.dumps(ss.res_m1a, ensure_ascii=False),
                            "diagnosticos":          DIAGNOSTICOS or "No especificado",
                            "palabras_clave":        str(ss.palabras_clave_activas) or ss.tema_activo,
                            "proyecto_pbl":          ss.res_m0b or "No disponible",
                            "user_suggestions":      user_sug_m2 or "Ninguna",
                        }
                        ss.res_m2a = generar_con_gemini("Motor_M2a.txt", vars_m2a, expect_json=False)
            with c_btn2:
                if st.button("📊 Generar Guía del Tutor", use_container_width=True, key="btn_tutor",
                             disabled=not ss.res_m2a):
                    with st.spinner("Creando Guía del Tutor…"):
                        vars_m2b = {
                            "plan_estrategico_json": json.dumps(ss.res_m1a, ensure_ascii=False),
                            "pop_quiz":              ss.res_m2a,
                            "diagnosticos":          DIAGNOSTICOS or "No especificado",
                            "sintesis_unidad":       json.dumps(ss.res_m0a, ensure_ascii=False) if ss.res_m0a else "No disponible",
                            "proyecto_pbl":          ss.res_m0b or "No disponible",
                            "user_suggestions":      user_sug_m2 or "Ninguna",
                        }
                        ss.res_m2b = generar_con_gemini("Motor_M2b.txt", vars_m2b, expect_json=False)

            if ss.res_m2a:
                st.divider()
                st.markdown("### 🎲 Pop Quiz DUA")
                st.markdown(ss.res_m2a)
            if ss.res_m2b:
                st.divider()
                st.markdown("### 📊 Panel de Control del Tutor")
                st.markdown(ss.res_m2b)
            if ss.res_m2a:
                st.divider()
                st.download_button(
                    "⬇️ Descargar Pop Quiz + Guía Tutor (.html)",
                    data=generar_html_evaluaciones(
                        ss.res_m2a, ss.res_m2b or "", ss.tema_activo, theme="aventura"
                    ).encode("utf-8"),
                    file_name=f"LasP_Evaluaciones_{_dt.now().strftime('%Y%m%d')}.html",
                    mime="text/html", use_container_width=True,
                )

    # ── EXPORTAR (zona semanal) ───────────────────────────────────────────────
    with tab_export:
        render_panel_exportacion(st.session_state, DIAGNOSTICOS)


# ═════════════════════════════════════════════════════════════════════════════
# ZONA TRIMESTRAL — Plan de Unidad + ABP, PDC
# ═════════════════════════════════════════════════════════════════════════════
elif zona == "📆  Trimestral":
    tab_unidad, tab_pdc = st.tabs(["🏗️ Plan de Unidad y ABP", "📋 PDC Trimestral"])

    # ── PLAN DE UNIDAD Y ABP ──────────────────────────────────────────────────
    with tab_unidad:
        st.markdown("Genera la síntesis curricular completa de la unidad y el Proyecto ABP.")
        st.info("💡 Trabaja sobre la **unidad completa**. No se selecciona un tema específico aquí.", icon="ℹ️")

        c_cfg1, c_cfg2 = st.columns(2)
        with c_cfg1:
            recursos_aula   = st.text_input("🖥️ Recursos disponibles:", value="",
                                             placeholder="Ej: proyector, tablets, materiales de arte…")
            permitir_online = st.radio("🌐 ¿Investigación online en el ABP?", ["Sí", "No"],
                                        horizontal=True, key="online_m0")
        with c_cfg2:
            user_sug_m0 = st.text_area("💬 Sugerencias:", value="",
                                        placeholder="Escribe aquí cualquier indicación especial…",
                                        height=80, key="sug_m0")

        if st.button("🛠️ Generar Síntesis Curricular", use_container_width=True, key="btn_sintesis"):
            with st.spinner("Sintetizando unidad y conceptos clave…"):
                vars_m0a = {
                    "unidad_real":     datos_libro.get("unidad_real", ""),
                    "temas":           str(temas_reales),
                    "contenido_temas": json.dumps(contenido_temas, ensure_ascii=False),
                    "diagnosticos":    DIAGNOSTICOS or "No especificado",
                }
                resultado = generar_con_gemini("Motor_M0a.txt", vars_m0a, expect_json=True)
                ss.res_m0a = resultado
                ss.res_m0b = None
                ss.res_m0c = None

        if ss.res_m0a and "unidad_sintetizada" in ss.res_m0a:
            unidad = ss.res_m0a["unidad_sintetizada"]
            st.success("✅ ¡Síntesis curricular generada!")
            st.markdown(f"## 📘 {unidad.get('titulo','Sin Título')}")
            st.divider()

            c_izq, c_der = st.columns([1.6, 1], gap="large")
            with c_izq:
                st.markdown("### 📑 Temas de la Unidad")
                for tema in unidad.get("temas_desarrollados", []):
                    nombre_t = tema.get("nombre", "Tema")
                    with st.expander(f"**📍 {nombre_t}**", expanded=False):
                        cc, ci = st.columns(2)
                        with cc:
                            st.markdown("**🔑 Conceptos Clave:**")
                            for item in forzar_lista(tema.get("conceptos_clave") or tema.get("conceptos") or []):
                                st.markdown(f"- {item}")
                        with ci:
                            st.markdown("**🧠 Inteligencias Múltiples:**")
                            for item in forzar_lista(tema.get("inteligencias_sugeridas") or tema.get("inteligencias") or []):
                                st.markdown(f"- {item}")
                        texto_real = contenido_temas.get(nombre_t, "")
                        if texto_real:
                            st.markdown("**📖 Contenido del libro (extracto):**")
                            st.caption(texto_real[:500] + ("…" if len(texto_real) > 500 else ""))
            with c_der:
                st.markdown("### 👨‍🏫 Notas DUA del Docente")
                st.info(unidad.get("notas_docente", "—"))
                st.markdown("### 🚀 Proyecto ABP *(resumen)*")
                st.success(unidad.get("proyecto_pbl", "⚠️ La IA no generó el proyecto."))

            st.divider()
            if st.button("➡️ Diseñar Proyecto ABP Completo", use_container_width=True,
                         type="primary", key="btn_abp"):
                with st.spinner("M0b diseñando el Proyecto ABP…"):
                    vars_m0b = {
                        "unidad_json":                   json.dumps(unidad, ensure_ascii=False),
                        "diagnosticos":                  DIAGNOSTICOS or "No especificado",
                        "recursos_aula":                 recursos_aula or "No especificado",
                        "permitir_investigacion_online": permitir_online,
                        "user_suggestions":              user_sug_m0 or "Ninguna",
                    }
                    ss.res_m0b = generar_con_gemini("Motor_M0b.txt", vars_m0b, expect_json=False)
                if ss.res_m0b:
                    with st.spinner("M0c creando rúbricas y fichas…"):
                        vars_m0c = {
                            "proyecto_pbl":     ss.res_m0b,
                            "unidad_json":      json.dumps(unidad, ensure_ascii=False),
                            "diagnosticos":     DIAGNOSTICOS or "No especificado",
                            "user_suggestions": user_sug_m0 or "Ninguna",
                        }
                        ss.res_m0c = generar_con_gemini("Motor_M0c.txt", vars_m0c, expect_json=False)

        if ss.res_m0a and "unidad_sintetizada" in ss.res_m0a:
            st.divider()
            st.download_button(
                "⬇️ Descargar Síntesis Curricular (.html)",
                data=generar_html_sintesis(ss.res_m0a, DIAGNOSTICOS),
                file_name=f"LasP_Sintesis_{_dt.now().strftime('%Y%m%d')}.html",
                mime="text/html", use_container_width=True,
            )
        if ss.res_m0b:
            st.divider()
            st.markdown("### 📋 Proyecto ABP Detallado")
            st.markdown(ss.res_m0b)
        if ss.res_m0c:
            st.divider()
            st.markdown("### 📊 Rúbrica y Fichas de Proceso")
            st.markdown(ss.res_m0c)
        if ss.res_m0b:
            st.divider()
            _u_titulo = ss.res_m0a["unidad_sintetizada"].get("titulo", "Unidad") if ss.res_m0a else "Unidad"
            st.download_button(
                "⬇️ Descargar Proyecto ABP (.html)",
                data=generar_html_abp(ss.res_m0b, ss.res_m0c or "", _u_titulo),
                file_name=f"LasP_ABP_{_dt.now().strftime('%Y%m%d')}.html",
                mime="text/html", use_container_width=True,
            )

    # ── PDC TRIMESTRAL ────────────────────────────────────────────────────────
    with tab_pdc:
        st.markdown("Genera un **PDC Trimestral completo** adaptado a tu materia y calendario. "
                    "Sube tu PDC base y deja que Gemini distribuya las unidades automáticamente.")

        col_pdc1, col_pdc2 = st.columns(2)
        with col_pdc1:
            st.markdown("#### 👤 Contexto Docente")
            pdc_teacher = st.text_input("Nombre del docente",
                                        value=ss.get("pdc_teacher", ss.get("teacher_name", "")),
                                        placeholder="Ej. Ruddy Ribera", key="pdc_teacher")
            pdc_level   = st.selectbox("Nivel",
                                        ["Primaria (5to)", "Primaria (4to)", "Secundaria"],
                                        key="pdc_level")
            pdc_subject = st.selectbox("Materia",
                                        ["Lenguaje", "Matemáticas", "Ciencias Naturales",
                                         "Ciencias Sociales", "Inglés", "Educación Física", "Música", "Arte"],
                                        key="pdc_subject")
            pdc_semanas = st.number_input("Semanas hábiles del trimestre", min_value=6, max_value=20,
                                           value=12, key="pdc_semanas")
        with col_pdc2:
            st.markdown("#### 📂 Archivos Base")
            pdc_file   = st.file_uploader("PDC Base (.docx)", type=["docx"], key="pdc_file_upload",
                                           help="Tu PDC actual que será tomado como referencia.")
            pdc_libros = st.file_uploader("Libros de texto (PDF, opcional)", type=["pdf"],
                                           accept_multiple_files=True, key="pdc_libros_upload")

        st.divider()
        if st.button("⚡ Generar PDC Trimestral", type="primary", use_container_width=True,
                     disabled=not pdc_teacher, key="btn_pdc"):
            if not pdc_file:
                st.warning("⚠️ Sube tu PDC base (.docx) para continuar.")
            else:
                with st.spinner("Gemini está construyendo tu PDC trimestral…"):
                    pdc_texto     = _leer_docx_texto(pdc_file.read())
                    resultado_pdc = generar_pdc_trimestral(
                        teacher_name=pdc_teacher, level=pdc_level, subject=pdc_subject,
                        pdc_texto=pdc_texto, semanas_calendario=int(pdc_semanas),
                    )
                    if resultado_pdc:
                        ss["pdc_resultado"] = resultado_pdc
                        st.success("✅ PDC generado correctamente.")

        if ss.get("pdc_resultado"):
            resultado = ss["pdc_resultado"]
            st.divider()
            st.markdown("### 📄 PDC Generado")
            st.markdown(resultado)
            st.divider()
            st.download_button(
                "⬇️ Descargar PDC (.html)",
                data=generar_html_pdc(resultado, pdc_teacher, pdc_subject, pdc_level),
                file_name=f"PDC_{pdc_subject.replace(' ','_')}_{pdc_teacher.replace(' ','_')}.html",
                mime="text/html", use_container_width=True, key="pdc_download_btn"
            )

