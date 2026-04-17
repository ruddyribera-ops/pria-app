"""
exporters/docx.py - Word Document Export Functions
===================================================
DOCX helpers and all document generators (sintesis, abp, plan_clase, ficha, evaluaciones).
"""

import io
import os
import re
import random
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── LOGO PATH ────────────────────────────────────────────────────────────────
_DIR = Path(__file__).parent.parent
LOGO_PATH = str(_DIR / "logo_laspalmas.png")
if not os.path.exists(LOGO_PATH):
    LOGO_PATH = str(Path.cwd() / "logo_laspalmas.png")


# ─── COLORES ─────────────────────────────────────────────────────────────────
D_VERDE = RGBColor(0x1B, 0x5E, 0x20)
D_VERDE2 = RGBColor(0x2E, 0x7D, 0x32)
D_NEGRO = RGBColor(0x1A, 0x1A, 0x1A)
D_BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
D_ROJO = RGBColor(0xCC, 0x22, 0x00)
D_AZUL = RGBColor(0x1A, 0x3A, 0x6A)
D_TEAL = RGBColor(0x00, 0x7A, 0x6E)
D_MORADO = RGBColor(0x6A, 0x3D, 0x9A)
D_NARANJA = RGBColor(0xCC, 0x55, 0x00)
D_GRIS = RGBColor(0x44, 0x44, 0x44)
D_GRIS_CLR = RGBColor(0xAA, 0xAA, 0xAA)

FONT_TITLE = "Impact"
FONT_BODY = "Calibri"


# ─────────────────────────────────────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────────────


def _nuevo_doc() -> Document:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.90)
        s.right_margin = Inches(0.90)
    doc.styles["Normal"].font.name = FONT_BODY
    doc.styles["Normal"].font.size = Pt(11)
    return doc


def _cell_bg(cell, hex6: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    tcPr.append(shd)


def _row_h(row, cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), "exact")
    trPr.append(h)


def _runs_from_inline_md(para, text: str, size=11, color=None, italic=False):
    """Add runs to paragraph, converting **bold** inline markdown to actual bold."""
    col = color or D_NEGRO
    parts = re.split(r"\*\*(.+?)\*\*", str(text))
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.name = FONT_BODY
        run.font.size = Pt(size)
        run.font.color.rgb = col
        run.font.italic = italic
        if i % 2 == 1:  # odd = the bold part
            run.font.bold = True


def _worksheet_header(doc: Document, materia: str, tema: str, grado: str = "5to"):
    """Navy banner header — logo centred, no text clipping."""
    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    r0 = tbl.rows[0]
    _row_h(r0, 1.5)

    c0 = r0.cells[0]
    _cell_bg(c0, "1A3A6A")
    c0.width = Inches(3.2)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(6)
    run = p0.add_run("Fichas de Trabajo\n")
    run.font.name = FONT_BODY
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = D_BLANCO
    run2 = p0.add_run("Las Palmas School")
    run2.font.name = FONT_BODY
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0xC8, 0xE6, 0xC9)

    c1 = r0.cells[1]
    _cell_bg(c1, "1A3A6A")
    c1.width = Inches(2.6)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        p1.add_run().add_picture(LOGO_PATH, width=Inches(1.55))

    c2 = r0.cells[2]
    _cell_bg(c2, "1A3A6A")
    c2.width = Inches(1.7)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(8)
    rg = p2.add_run(f"{grado}\nPRIMARIA")
    rg.font.name = FONT_BODY
    rg.font.bold = True
    rg.font.size = Pt(16)
    rg.font.color.rgb = D_BLANCO

    r1 = tbl.rows[1]
    _row_h(r1, 0.6)
    r1.cells[0].merge(r1.cells[2])
    c_sub = r1.cells[0]
    _cell_bg(c_sub, "143060")
    ps = c_sub.paragraphs[0]
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = ps.add_run(materia.upper())
    rs.font.name = FONT_BODY
    rs.font.bold = True
    rs.font.size = Pt(12)
    rs.font.color.rgb = D_BLANCO

    doc.add_paragraph()

    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_before = Pt(4)
    p_t.paragraph_format.space_after = Pt(6)
    rt = p_t.add_run(tema.upper())
    rt.font.name = FONT_TITLE
    rt.font.bold = True
    rt.font.size = Pt(20)
    rt.font.color.rgb = D_ROJO
    pPr = p_t._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "12")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "CC2200")
    pBdr.append(bot)
    pPr.append(pBdr)
    doc.add_paragraph()


def _section_banner(doc: Document, texto: str, hex_color: str, size=12):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, hex_color)
    _row_h(tbl.rows[0], 0.55)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.font.name = FONT_BODY
    run.font.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = D_BLANCO
    doc.add_paragraph()


def _h(doc, text, lvl=2, color=None):
    clean = str(text).replace("**", "").replace("__", "")
    p = doc.add_heading(clean, level=lvl)
    for run in p.runs:
        run.font.color.rgb = color or D_AZUL
        run.font.name = FONT_BODY
        run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(lvl, 11))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)


def _p(doc, text, bold=False, italic=False, color=None, size=11):
    clean = str(text).replace("**", "").replace("__", "")
    p = doc.add_paragraph()
    run = p.add_run(clean)
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or D_NEGRO
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.space_before = Pt(3)


def _p_md(doc, text, size=11, color=None):
    """Paragraph with inline **bold** markdown support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.space_before = Pt(3)
    _runs_from_inline_md(p, text, size=size, color=color)


def _bul(doc, text, color=None, size=11):
    """Bullet item with inline **bold** support."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    _runs_from_inline_md(p, str(text), size=size, color=color)


def _bloom_badge_table(doc: Document, verbos: list):
    """Horizontal colored badge row — one cell per Bloom verb."""
    if not verbos:
        return
    colors = ["1B5E20", "007A6E", "1A3A6A", "6A3D9A", "CC5500", "2E6DA4", "8B0000"]
    tbl = doc.add_table(rows=1, cols=len(verbos))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _row_h(tbl.rows[0], 0.52)
    for i, verb in enumerate(verbos):
        cell = tbl.rows[0].cells[i]
        _cell_bg(cell, colors[i % len(colors)])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        run = p.add_run(str(verb).upper().replace("**", "")[:20])
        run.font.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = D_BLANCO
    doc.add_paragraph()


def _secuencia_timeline(doc: Document, bloques: list):
    """Visual timeline: coloured left block = phase/time, right = content."""
    block_colors = {"inicio": "1B5E20", "desarrollo": "1A3A6A", "cierre": "CC5500"}
    for bloque in bloques:
        nombre = str(bloque.get("nombre", "Bloque"))
        duracion = str(bloque.get("duracion", "?"))
        objetivo = str(bloque.get("objetivo", ""))
        nota = str(bloque.get("nota", ""))

        color = "1A3A6A"
        for key, c in block_colors.items():
            if key in nombre.lower():
                color = c
                break

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]
        left.width = Inches(1.1)
        right.width = Inches(5.9)

        _cell_bg(left, color)

        pl = left.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.paragraph_format.space_before = Pt(4)
        rl1 = pl.add_run(nombre.upper()[:10])
        rl1.font.bold = True
        rl1.font.size = Pt(9)
        rl1.font.color.rgb = D_BLANCO
        rl1.font.name = FONT_BODY
        rl2 = pl.add_run(f"\n{duracion} min")
        rl2.font.size = Pt(8)
        rl2.font.name = FONT_BODY
        rl2.font.color.rgb = RGBColor(0xC8, 0xE6, 0xC9)

        _cell_bg(right, "FAFAFA")
        pr = right.paragraphs[0]
        pr.paragraph_format.left_indent = Inches(0.10)
        pr.paragraph_format.space_before = Pt(4)
        pr.paragraph_format.space_after = Pt(3)
        _runs_from_inline_md(pr, objetivo, size=11)

        if nota:
            pn = right.add_paragraph()
            pn.paragraph_format.left_indent = Inches(0.10)
            pn.paragraph_format.space_after = Pt(4)
            rn = pn.add_run(f"💡 {nota.replace('**', '')}")
            rn.font.italic = True
            rn.font.size = Pt(10)
            rn.font.color.rgb = D_TEAL
            rn.font.name = FONT_BODY

        doc.add_paragraph()


def _dua_framework_header(doc: Document):
    """3-column DUA principle header banner."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _row_h(tbl.rows[0], 0.55)
    principles = [
        ("🧩 Representación", "6A3D9A"),
        ("✍️ Acción y Expresión", "1A3A6A"),
        ("💡 Motivación", "007A6E"),
    ]
    for i, (title, color) in enumerate(principles):
        cell = tbl.rows[0].cells[i]
        _cell_bg(cell, color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = D_BLANCO
        run.font.name = FONT_BODY
    doc.add_paragraph()


def _dua_note(doc: Document, icon: str, text: str, icon_hex: str = "007A6E"):
    """Margin-style DUA callout: colored icon badge on left, content on right."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]
    left.width = Inches(0.50)
    _cell_bg(left, icon_hex)
    _row_h(tbl.rows[0], 0.65)
    pl = left.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.paragraph_format.space_before = Pt(3)
    rl = pl.add_run(icon)
    rl.font.size = Pt(11)
    rl.font.color.rgb = D_BLANCO
    rl.font.bold = True
    _cell_bg(right, "F9FAFB")
    pr = right.paragraphs[0]
    pr.paragraph_format.left_indent = Inches(0.08)
    pr.paragraph_format.space_before = Pt(3)
    pr.paragraph_format.space_after = Pt(3)
    _runs_from_inline_md(pr, str(text).lstrip("- •"), size=10)
    doc.add_paragraph()


def _tbl(doc, filas: list, hdr_color="1A3A6A"):
    if not filas:
        return
    cols = list(filas[0].keys())
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, col in enumerate(cols):
        hdr[i].text = col
        _cell_bg(hdr[i], hdr_color)
        r = hdr[i].paragraphs[0].runs[0]
        r.font.color.rgb = D_BLANCO
        r.font.bold = True
        r.font.name = FONT_BODY
        r.font.size = Pt(10)
    for fila in filas:
        row = tbl.add_row().cells
        for i, col in enumerate(cols):
            row[i].text = str(fila.get(col, ""))
            _cell_bg(row[i], "F1F8E9")
            r = row[i].paragraphs[0].runs[0]
            r.font.color.rgb = D_NEGRO
            r.font.name = FONT_BODY
            r.font.size = Pt(10)
    doc.add_paragraph()


# ─── AI AUTOPRESENTACION FILTER ───────────────────────────────────────────────
_AI_INTRO_FRAGMENTS = [
    "soy el motor",
    "motor m0",
    "motor m1",
    "motor m2",
    "método palma-ribera",
    "palma-ribera",
    "he diseñado este",
    "me encanta acompañarte",
    "como experto en currículo",
    "bienvenidos, futuros",
    "a continuación te presento",
    "como motor",
    "he sido diseñado",
    "este pop quiz",
    "esta evaluación ha sido",
    "bienvenido a la",
    "estimado docente,",
    "querido estudiante,",
    "hola, soy",
    "aquí tienes",
    "a continuación encontrarás",
]


def _es_autopresentacion(linea: str) -> bool:
    ll = linea.lower()[:120]
    return any(f in ll for f in _AI_INTRO_FRAGMENTS)


def _md(doc, texto: str):
    """Convert markdown text to DOCX."""
    if not texto:
        return
    buf = []

    def flush():
        if not buf:
            return
        validas = [l for l in buf if not all(c in "-| " for c in l)]
        if not validas:
            buf.clear()
            return
        cols = [c.strip() for c in validas[0].split("|") if c.strip()]
        rows = []
        for rl in validas[1:]:
            vals = [c.strip() for c in rl.split("|") if c.strip()]
            if vals:
                rows.append(
                    dict(zip(cols, vals + [""] * (max(0, len(cols) - len(vals)))))
                )
        if rows:
            _tbl(doc, rows)
        buf.clear()

    for line in texto.split("\n"):
        s = line.strip()
        if s.startswith("|"):
            buf.append(s)
            continue
        else:
            flush()

        if _es_autopresentacion(s):
            continue
        if s in ("---", "___", "***") or (len(s) >= 3 and all(c == "-" for c in s)):
            continue
        if not s:
            doc.add_paragraph()
            continue

        if s.startswith("### "):
            _h(doc, s[4:], 3)
        elif s.startswith("## "):
            _h(doc, s[3:], 2)
        elif s.startswith("# "):
            _h(doc, s[2:], 1)
        elif s.startswith("> "):
            _p_md(doc, s[2:], color=D_TEAL)
        elif s.startswith("- ") or s.startswith("* "):
            _bul(doc, s[2:])
        else:
            _p_md(doc, s)
    flush()


def _start_two_columns(doc: Document):
    """Insert a continuous section break to begin 2-column newspaper layout."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "continuous")
    cols_el = OxmlElement("w:cols")
    cols_el.set(qn("w:num"), "2")
    cols_el.set(qn("w:space"), "720")
    sectPr.append(type_el)
    sectPr.append(cols_el)
    pPr.append(sectPr)


def _end_two_columns(doc: Document):
    """Insert a continuous section break to return to 1-column layout."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "continuous")
    cols_el = OxmlElement("w:cols")
    cols_el.set(qn("w:num"), "1")
    sectPr.append(type_el)
    sectPr.append(cols_el)
    pPr.append(sectPr)


# ─────────────────────────────────────────────────────────────────────────────
# FICHA HELPERS — Student worksheet components
# ─────────────────────────────────────────────────────────────────────────────


def _info_estudiante(doc: Document):
    """Name / Date / Section fields at top of student worksheet."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    campos = [("Nombre", 3.5), ("Fecha", 1.8), ("Sección", 1.8)]
    for i, (label, w_in) in enumerate(campos):
        cell = tbl.rows[0].cells[i]
        cell.width = Inches(w_in)
        _cell_bg(cell, "F1F8E9")
        _row_h(tbl.rows[0], 0.7)
        p = cell.paragraphs[0]
        r1 = p.add_run(f"{label}: ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.name = FONT_BODY
        r1.font.color.rgb = D_AZUL
        r2 = p.add_run("_" * (int(w_in * 8)))
        r2.font.size = Pt(11)
        r2.font.name = FONT_BODY
        r2.font.color.rgb = D_GRIS_CLR
    doc.add_paragraph()


def _historia_gancho_box(doc: Document, texto: str):
    """Italic boxed story hook paragraph."""
    if not texto:
        return
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, "EEF4FF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    run = p.add_run(str(texto))
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = D_NEGRO
    doc.add_paragraph()


def _opcion_multiple(doc: Document, pregunta: str, opciones: list, letra_colors: list):
    """Multiple choice question with circle markers."""
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(6)
    p_q.paragraph_format.space_after = Pt(2)
    _runs_from_inline_md(p_q, pregunta, size=12)
    for run in p_q.runs:
        run.font.bold = True

    letras = ["A)", "B)", "C)", "D)"]
    for j, opcion in enumerate(opciones):
        p_op = doc.add_paragraph()
        p_op.paragraph_format.left_indent = Inches(0.4)
        p_op.paragraph_format.space_before = Pt(2)
        p_op.paragraph_format.space_after = Pt(2)
        r_circ = p_op.add_run("○  ")
        r_circ.font.size = Pt(14)
        r_circ.font.color.rgb = D_GRIS
        r_letra = p_op.add_run(letras[j] + " ")
        r_letra.font.bold = True
        r_letra.font.size = Pt(11)
        r_letra.font.color.rgb = letra_colors[j % len(letra_colors)]
        r_text = p_op.add_run(str(opcion).replace("**", "").lstrip("abcdefgABCDEFG) "))
        r_text.font.size = Pt(11)
        r_text.font.name = FONT_BODY
        r_text.font.color.rgb = D_NEGRO
    doc.add_paragraph()


def _lineas_respuesta(doc: Document, n: int = 3, label: str = ""):
    """Dotted write lines for student responses."""
    if label:
        p_lbl = doc.add_paragraph()
        r = p_lbl.add_run(label)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = D_AZUL
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run("_" * 90)
        run.font.color.rgb = D_GRIS_CLR
        run.font.size = Pt(10)
    doc.add_paragraph()


def _caja_escritura(doc: Document, alto_cm: float = 5.0, instruccion: str = ""):
    """Large bordered box for drawing or extended writing."""
    if instruccion:
        p_i = doc.add_paragraph()
        r = p_i.add_run(instruccion)
        r.font.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = D_GRIS
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, "FFFFFF")
    _row_h(tbl.rows[0], alto_cm)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(" ")
    r.font.size = Pt(8)
    doc.add_paragraph()


def _palabra_badge(para, palabra: str, color: RGBColor):
    """Inline word badge for word bank."""
    r1 = para.add_run(f"  ▸ {palabra.upper()}  ")
    r1.font.name = FONT_BODY
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = color


def _banco_palabras(doc: Document, palabras: list):
    """Styled word bank row."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    colors = [D_AZUL, D_ROJO, D_VERDE2, D_MORADO, D_NARANJA]
    label = p.add_run("Banco de palabras:  ")
    label.font.bold = True
    label.font.size = Pt(11)
    label.font.color.rgb = D_NEGRO
    for i, pal in enumerate(palabras):
        _palabra_badge(p, str(pal), colors[i % len(colors)])


def _generar_sopa_letras(palabras: list, size: int = 14) -> list:
    """
    Word-search: ~1/3 horizontal, ~2/3 vertical, interleaved.
    """
    grid = [["" for _ in range(size)] for _ in range(size)]
    letras_relleno = "ABCDEFGHIJKLMNOPRSTUVWXZ"

    words = [p.upper().replace(" ", "").replace("Ñ", "N") for p in palabras]
    words = [w for w in words if 0 < len(w) <= size]
    if not words:
        for r in range(size):
            for c in range(size):
                grid[r][c] = random.choice(letras_relleno)
        return grid

    n = len(words)
    n_horiz = max(1, round(n / 3))
    directions = ["H"] * n_horiz + ["V"] * (n - n_horiz)
    random.shuffle(directions)

    def _find_positions(word, direction):
        ln = len(word)
        over, free = [], []
        if direction == "H":
            for r in range(size):
                for c in range(size - ln + 1):
                    if all(
                        grid[r][c + i] == "" or grid[r][c + i] == word[i]
                        for i in range(ln)
                    ):
                        hits = sum(1 for i in range(ln) if grid[r][c + i] == word[i])
                        (over if hits else free).append((r, c))
        else:
            for c in range(size):
                for r in range(size - ln + 1):
                    if all(
                        grid[r + i][c] == "" or grid[r + i][c] == word[i]
                        for i in range(ln)
                    ):
                        hits = sum(1 for i in range(ln) if grid[r + i][c] == word[i])
                        (over if hits else free).append((r, c))
        return over, free

    def _place(word, direction, r, c):
        ln = len(word)
        if direction == "H":
            for i in range(ln):
                grid[r][c + i] = word[i]
        else:
            for i in range(ln):
                grid[r + i][c] = word[i]

    for word, direction in zip(words, directions):
        over, free = _find_positions(word, direction)
        pool = over if over else free
        if pool:
            r, c = random.choice(pool)
            _place(word, direction, r, c)
        else:
            alt = "V" if direction == "H" else "H"
            over2, free2 = _find_positions(word, alt)
            pool2 = over2 if over2 else free2
            if pool2:
                r, c = random.choice(pool2)
                _place(word, alt, r, c)

    for r in range(size):
        for c in range(size):
            if not grid[r][c]:
                grid[r][c] = random.choice(letras_relleno)
    return grid


def _insertar_sopa_docx(doc: Document, grid: list):
    """Render word search grid as a table."""
    size = len(grid)
    tbl = doc.add_table(rows=size, cols=size)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_w = Inches(0.28)
    for r in range(size):
        row = tbl.rows[r]
        _row_h(row, 0.38)
        for c in range(size):
            cell = row.cells[c]
            cell.width = cell_w
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(grid[r][c]))
            run.font.name = FONT_BODY
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = D_NEGRO
    doc.add_paragraph()


def _frase_pergamino(doc: Document, frase: str, palabras_secretas: list):
    """Boxed fill-in-the-blank sentence."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, "FFF8E1")
    _row_h(tbl.rows[0], 1.4)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(frase))
    run.font.name = FONT_BODY
    run.font.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = D_NEGRO
    doc.add_paragraph()
    if palabras_secretas:
        _banco_palabras(doc, palabras_secretas)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX GENERATORS
# ─────────────────────────────────────────────────────────────────────────────


def generar_docx_sintesis(res_m0a: dict, diagnosticos: str) -> bytes:
    doc = _nuevo_doc()
    unidad = res_m0a.get("unidad_sintetizada", {})
    titulo = unidad.get("titulo", "Síntesis de Unidad")
    _worksheet_header(doc, "Planificación Curricular", titulo)

    _h(doc, "Perfil Neuro-Inclusivo del Aula", 2)
    _p(doc, diagnosticos)
    doc.add_paragraph()

    _h(doc, "Temas de la Unidad", 2)
    for tema in unidad.get("temas_desarrollados", []):
        _h(doc, tema.get("nombre", "Tema"), 3, D_VERDE)
        _p(doc, "Conceptos Clave:", bold=True, color=D_AZUL)
        for c in tema.get("conceptos_clave") or tema.get("conceptos") or []:
            _bul(doc, str(c), D_AZUL)
        doc.add_paragraph()
        _p(doc, "Inteligencias Múltiples sugeridas:", bold=True, color=D_VERDE)
        for i in tema.get("inteligencias_sugeridas") or tema.get("inteligencias") or []:
            _bul(doc, str(i))
        doc.add_paragraph()

    _h(doc, "Notas DUA del Docente", 2)
    _p_md(doc, unidad.get("notas_docente", "—"))
    doc.add_paragraph()

    _h(doc, "Proyecto ABP — Resumen", 2)
    _p_md(doc, unidad.get("proyecto_pbl", "—"))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_docx_abp(res_m0b: str, res_m0c: str, titulo_unidad: str) -> bytes:
    doc = _nuevo_doc()
    _worksheet_header(doc, "Proyecto ABP", titulo_unidad)
    _h(doc, "Diseño del Proyecto ABP", 2)
    _md(doc, res_m0b or "No generado.")
    doc.add_page_break()
    _h(doc, "Rúbrica y Fichas de Proceso", 2)
    _md(doc, res_m0c or "No generado.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_docx_plan_clase(res_m1a: dict, tema: str, diagnosticos: str) -> bytes:
    doc = _nuevo_doc()
    _worksheet_header(doc, "Plan de Clase · 45 minutos", tema)

    verbos = res_m1a.get("mapa_cognitivo", {}).get("verbos", [])
    if verbos:
        _h(doc, "Verbos Operativos — Taxonomía Bloom", 2)
        _bloom_badge_table(doc, verbos)

    im = res_m1a.get("inteligencias_multiples", [])
    if im:
        _h(doc, "Inteligencias Múltiples", 2)
        _tbl(doc, im)

    _section_banner(doc, "⏱️  Secuencia Didáctica — 45 minutos", "1A3A6A", size=13)
    _secuencia_timeline(doc, res_m1a.get("secuencia_didactica", {}).get("bloques", []))

    _section_banner(doc, "🧠  Directrices DUA — Neuro-Inclusión", "007A6E", size=13)
    _dua_framework_header(doc)
    for d in res_m1a.get("dua_neuroinclusion", []):
        dl = str(d).lower()
        if "tea" in dl:
            _dua_note(doc, "👁 TEA", str(d), "007A6E")
        elif "tdah" in dl or "adhd" in dl:
            _dua_note(doc, "⚡ TDAH", str(d), "CC5500")
        else:
            _bul(doc, str(d), size=11)
    doc.add_paragraph()

    adapt = res_m1a.get("tabla_adaptaciones_clase", [])
    if adapt:
        _section_banner(doc, "♿  Adaptaciones por Diagnóstico", "1A5276", size=13)
        _tbl(doc, adapt)

    _h(doc, "Perfil del Aula", 2)
    _p_md(doc, res_m1a.get("perfil_aula_resumido", diagnosticos), size=11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_docx_ficha(res_m1c: dict, tema: str) -> bytes:
    """
    Ficha Gamificada — diseño de hoja de trabajo real para imprimir.
    """
    doc = _nuevo_doc()
    raw = res_m1c
    ficha = raw.get("ficha_trabajo", raw) if isinstance(raw, dict) else {}

    _worksheet_header(doc, "Comunicación — Ficha Gamificada", tema)
    _info_estudiante(doc)
    _historia_gancho_box(doc, ficha.get("historia_gancho", ""))

    misiones = ficha.get("misiones", {})
    MC_COLORS = [D_AZUL, D_ROJO, D_VERDE2]

    _section_banner(doc, "🔮  Misión 1: El Oráculo", "2E6DA4", size=13)
    p_inst = doc.add_paragraph()
    r = p_inst.add_run(
        "Instrucciones: Lee cada pregunta y marca con ○ la respuesta correcta."
    )
    r.font.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = D_GRIS
    doc.add_paragraph()

    for idx, q in enumerate(misiones.get("oraculo", [])):
        _opcion_multiple(
            doc, f"{idx + 1}. {q.get('pregunta', '')}", q.get("opciones", []), MC_COLORS
        )

    _section_banner(doc, "🌉  Misión 2: El Puente", "007A6E", size=13)
    p_inst2 = doc.add_paragraph()
    r2 = p_inst2.add_run(
        "Instrucciones: Une con una línea cada palabra con su significado."
    )
    r2.font.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = D_GRIS
    doc.add_paragraph()

    puente_data = misiones.get("puente", [])
    if puente_data:
        import random as _rnd

        definiciones = [p.get("significado", "") for p in puente_data]
        definiciones_mezcladas = definiciones[:]
        _rnd.shuffle(definiciones_mezcladas)
        tbl_puente = doc.add_table(rows=len(puente_data) + 1, cols=2)
        tbl_puente.style = "Table Grid"
        tbl_puente.alignment = WD_TABLE_ALIGNMENT.CENTER
        _cell_bg(tbl_puente.rows[0].cells[0], "007A6E")
        _cell_bg(tbl_puente.rows[0].cells[1], "007A6E")
        for j, hdr_txt in enumerate(["PALABRA", "SIGNIFICADO"]):
            ph = tbl_puente.rows[0].cells[j].paragraphs[0]
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rh = ph.add_run(hdr_txt)
            rh.font.bold = True
            rh.font.color.rgb = D_BLANCO
            rh.font.name = FONT_BODY
            rh.font.size = Pt(11)
        for k, par in enumerate(puente_data):
            row = tbl_puente.rows[k + 1]
            _cell_bg(row.cells[0], "E8F5E9")
            _cell_bg(row.cells[1], "FFFFFF")
            _row_h(row, 0.7)
            p_pal = row.cells[0].paragraphs[0]
            p_pal.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_pal = p_pal.add_run(str(par.get("palabra", "")))
            r_pal.font.bold = True
            r_pal.font.color.rgb = D_VERDE
            r_pal.font.name = FONT_BODY
            r_pal.font.size = Pt(11)
            p_def = row.cells[1].paragraphs[0]
            r_def = p_def.add_run(str(definiciones_mezcladas[k]))
            r_def.font.color.rgb = D_NEGRO
            r_def.font.name = FONT_BODY
            r_def.font.size = Pt(10)
    doc.add_paragraph()

    _section_banner(doc, "🥣  Misión 3: La Sopa de Letras", "1A5276", size=13)
    p_inst3 = doc.add_paragraph()
    r3 = p_inst3.add_run(
        "Instrucciones: Encuentra y encierra las siguientes palabras en la sopa de letras."
    )
    r3.font.italic = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = D_GRIS
    doc.add_paragraph()

    palabras_sopa = misiones.get("sopa", [])
    if palabras_sopa:
        _banco_palabras(doc, palabras_sopa)
        grid = _generar_sopa_letras(palabras_sopa, size=12)
        _insertar_sopa_docx(doc, grid)

    _section_banner(doc, "📜  Misión 4: El Pergamino", "6A3D9A", size=13)
    p_inst4 = doc.add_paragraph()
    r4 = p_inst4.add_run(
        "Instrucciones: Completa la frase usando las palabras del banco."
    )
    r4.font.italic = True
    r4.font.size = Pt(10)
    r4.font.color.rgb = D_GRIS
    doc.add_paragraph()

    per = misiones.get("pergamino", {})
    _frase_pergamino(
        doc, per.get("frase_con_espacios", ""), per.get("palabras_secretas", [])
    )

    _section_banner(doc, "🎨  Misión 5: El Lienzo", "CC5500", size=13)
    lienzo_inst = str(
        misiones.get("lienzo", "Crea tu obra de arte basándote en lo aprendido.")
    )
    p_inst5 = doc.add_paragraph()
    r5 = p_inst5.add_run(lienzo_inst.replace("**", ""))
    r5.font.size = Pt(11)
    r5.font.color.rgb = D_NEGRO
    r5.font.name = FONT_BODY
    doc.add_paragraph()
    _caja_escritura(
        doc, alto_cm=6.0, instruccion="✏️ Dibuja o escribe tu respuesta en este espacio:"
    )

    adapt = ficha.get("adaptaciones_por_mision", [])
    if adapt:
        doc.add_page_break()
        _section_banner(
            doc, "📋  ANEXO DOCENTE — Adaptaciones por Misión", "1A3A6A", size=13
        )
        p_nota = doc.add_paragraph()
        r_nota = p_nota.add_run(
            "Esta página es solo para el docente. No imprimir para los estudiantes."
        )
        r_nota.font.italic = True
        r_nota.font.color.rgb = D_ROJO
        r_nota.font.size = Pt(10)
        doc.add_paragraph()
        _tbl(doc, adapt)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_docx_evaluaciones(res_m2a: str, res_m2b: str, tema: str) -> bytes:
    doc = _nuevo_doc()

    _worksheet_header(doc, "Evaluación Formativa DUA", tema)
    _info_estudiante(doc)
    _section_banner(doc, "🎲  Pop Quiz — ¡Misión de Cierre!", "2E6DA4", size=13)
    _md(doc, res_m2a or "No generado.")

    if res_m2b:
        doc.add_page_break()
        _section_banner(
            doc,
            "📊  PANEL DE CONTROL DEL TUTOR  —  Solo para el Docente",
            "1A5276",
            size=13,
        )
        p_nota = doc.add_paragraph()
        r_nota = p_nota.add_run(
            "Esta página es confidencial. No distribuir a los estudiantes."
        )
        r_nota.font.italic = True
        r_nota.font.color.rgb = D_ROJO
        r_nota.font.size = Pt(10)
        doc.add_paragraph()
        _start_two_columns(doc)
        _md(doc, res_m2b)
        _end_two_columns(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
