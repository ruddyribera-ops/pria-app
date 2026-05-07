"""DOCX import service for parsing curriculum documents."""

from typing import Optional, Dict, Any
from pathlib import Path


class ImportService:
    """Service for importing curriculum content from DOCX files."""

    SUPPORTED_FORMATS = {".docx", ".doc"}

    def __init__(self):
        """Initialize import service."""
        try:
            from docx import Document
            self.docx_available = True
        except ImportError:
            self.docx_available = False

    async def import_from_docx(
        self,
        file_path: str,
    ) -> Dict[str, Any]:
        """Import PDC content from a DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with extracted content structure

        Raises:
            ValueError: If python-docx is not installed or file cannot be parsed
        """
        if not self.docx_available:
            raise ValueError(
                "DOCX import requires 'python-docx' library. "
                "Install it with: pip install python-docx"
            )

        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                f"Supported formats: {', '.join(self.SUPPORTED_FORMATS)}"
            )

        try:
            from docx import Document
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")

        # Extract content structure
        content = {
            "title": self._extract_title(doc),
            "objetivo": self._extract_section(doc, "objetivo"),
            "contenidos": self._extract_section(doc, "contenidos"),
            "momentos": self._extract_section(doc, "momentos"),
            "recursos": self._extract_section(doc, "recursos"),
            "periodos": self._extract_section(doc, "periodos"),
            "criterios": self._extract_section(doc, "criterios"),
        }

        return content

    def _extract_title(self, doc: Any) -> str:
        """Extract title from document (first heading or first paragraph)."""
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                return para.text.strip()
            if para.text.strip():
                return para.text.strip()
        return "Imported PDC"

    def _extract_section(self, doc: Any, section_name: str) -> str:
        """Extract a section from document by heading."""
        content_lines = []
        found_section = False

        for para in doc.paragraphs:
            text = para.text.strip()

            # Check if this is the start of our section
            if section_name.lower() in text.lower():
                found_section = True
                continue

            # Stop if we hit another heading (new section)
            if found_section and para.style.name.startswith("Heading"):
                break

            # Collect content
            if found_section and text:
                content_lines.append(text)

        return "\n".join(content_lines) if content_lines else ""

    async def import_from_file(self, file_path: str) -> Dict[str, Any]:
        """Generic import method that detects file format.

        Args:
            file_path: Path to file

        Returns:
            Extracted content dictionary
        """
        file_ext = Path(file_path).suffix.lower()

        if file_ext in {".docx", ".doc"}:
            return await self.import_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
