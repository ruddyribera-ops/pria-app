"""
DocxExportService - Generates DOCX (Word) documents from PDC content in PRIA v7
"""
from typing import Optional
from datetime import datetime
from app.models.pdc import PDC
from app.schemas.export import BrandingConfig

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    """Convert hex color (#RRGGBB) to RGB tuple"""
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def shade_cell(cell, color: str) -> None:
    """Apply background color to a table cell"""
    try:
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color.lstrip("#"))
        cell._element.get_or_add_tcPr().append(shading_elm)
    except Exception:
        pass


class DocxExportService:
    """Service for generating DOCX exports from PDC content"""

    @staticmethod
    def export_pdc(
        pdc: PDC,
        branding: BrandingConfig,
        format_type: str = "detailed",
    ) -> bytes:
        """
        Generate a DOCX document from PDC content with school branding.

        Args:
            pdc: PDC model instance with content
            branding: BrandingConfig with school branding details
            format_type: "simple" or "detailed"

        Returns:
            DOCX file as bytes

        Raises:
            ImportError: If python-docx is not available
            ValueError: If PDC data is invalid
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is not installed. "
                "Install it with: pip install python-docx"
            )

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        style.font.name = branding.primary_font or "Arial"
        style.font.size = Pt(11)

        # Add header section
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Header with school info
        header_paragraph = doc.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_paragraph.add_run(branding.school_name)
        header_run.font.size = Pt(16)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(*hex_to_rgb(branding.header_color))

        # Add school logo if available
        if branding.logo_url:
            try:
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_paragraph.add_run().add_picture(branding.logo_url, width=Inches(1.5))
            except Exception:
                # Silently skip if logo cannot be loaded
                pass

        # Document title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(pdc.title or "Plan de Desarrollo Curricular")
        title_run.font.size = Pt(14)
        title_run.font.bold = True

        # Subject and grade
        info_paragraph = doc.add_paragraph()
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_paragraph.add_run(f"{pdc.subject} - {pdc.grade_level}")
        info_run.font.size = Pt(11)

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Fecha: {datetime.utcnow().strftime('%d/%m/%Y')}")
        date_run.font.size = Pt(10)

        doc.add_paragraph()  # Spacing

        # MESCP Table
        if format_type == "detailed":
            doc.add_heading("Planificación por Objetivos (MESCP)", level=2)
            table = doc.add_table(rows=1, cols=6)
            table.style = "Light Grid Accent 1"

            # Header row
            header_cells = table.rows[0].cells
            headers = ["Objetivo", "Contenidos", "Momentos", "Recursos", "Períodos", "Criterios"]
            for i, header_text in enumerate(headers):
                cell = header_cells[i]
                cell.text = header_text
                shade_cell(cell, branding.header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

            # Data row
            row_cells = table.add_row().cells
            row_cells[0].text = pdc.objetivo or "No definido"
            row_cells[1].text = pdc.contenidos or "No definido"
            row_cells[2].text = pdc.momentos or "No definido"
            row_cells[3].text = pdc.recursos or "No definido"
            row_cells[4].text = pdc.periodos or "No definido"
            row_cells[5].text = pdc.criterios or "No definido"

        # Adaptaciones section
        if hasattr(pdc, "adaptations") and pdc.adaptations:
            doc.add_heading("Adaptaciones Curriculares", level=2)
            for adaptation in pdc.adaptations[:5]:  # Limit to 5 adaptations
                p = doc.add_paragraph(
                    f"• {getattr(adaptation, 'description', 'Adaptación')}",
                    style="List Bullet"
                )
                p.paragraph_format.left_indent = Inches(0.5)

        # Inteligencias Múltiples section
        if hasattr(pdc, "inteligencias") and pdc.inteligencias:
            doc.add_heading("Inteligencias Múltiples", level=2)
            intelligences = [
                getattr(intel, "name", f"Inteligencia {i}")
                for i, intel in enumerate(pdc.inteligencias[:8])
            ]
            p = doc.add_paragraph(", ".join(intelligences))
            p.paragraph_format.left_indent = Inches(0.25)

        # Productos section
        if hasattr(pdc, "productos") and pdc.productos:
            doc.add_heading("Productos Esperados", level=2)
            for producto in pdc.productos[:5]:  # Limit to 5 products
                p = doc.add_paragraph(
                    f"• {getattr(producto, 'name', 'Producto')}",
                    style="List Bullet"
                )
                p.paragraph_format.left_indent = Inches(0.5)

        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(branding.footer_text or branding.school_name)
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(*hex_to_rgb(branding.footer_color))

        page_num_para = doc.add_paragraph()
        page_num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        page_num_run = page_num_para.add_run(f"Generado: {datetime.utcnow().strftime('%d/%m/%Y %H:%M:%S')}")
        page_num_run.font.size = Pt(8)

        # Return bytes
        from io import BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
