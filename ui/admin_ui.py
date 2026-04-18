"""
ui/admin_ui.py - Admin Panel
============================
Admin panel with file uploads, user management, and daily reset.
"""

import streamlit as st
import os
import io
import re
import tempfile
import json

from google import genai

from ui.helpers import (
    GEMINI_MODEL,
    _get_keys,
    _rotate_key,
    log_event,
    get_motor_stats,
    generar_pdc_trimestral,
    _leer_docx_texto,
)
from parser_archivos import (
    parse_horarios,
    parse_calendario,
    parse_cronograma,
    parse_comisiones,
    parse_vigilancia_pdf,
)
from db import (
    guardar_horario_docente,
    guardar_eventos_calendario,
    guardar_actividades_cronograma,
    guardar_comisiones,
    guardar_vigilancias,
    crear_sesion,
    guardar_micro_objetivos,
    get_all_hojas,
    crear_usuario,
    get_all_usuarios,
    toggle_usuario_activo,
    eliminar_usuario,
    actualizar_usuario_admin,
    reset_dia_docente,
    get_horario_docente_completo,
    guardar_bloque_horario_manual,
    eliminar_bloque_horario_manual,
)


def render_admin_panel():
    """Render admin panel tabs. Call from app_ui.py when usuario_rol == admin."""
    adm_tab_arch, adm_tab_users, adm_tab_tracker, adm_tab_bloques = st.tabs(
        ["📂 Archivos Fuente", "👥 Gestión de Usuarios", "🌅 Reset Diario", "✏️ Editar Bloques"]
    )

    # ═══════════════════════════════════════════════════════════════
    # TAB 1: ARCHIVOS FUENTE
    # ═══════════════════════════════════════════════════════════════
    with adm_tab_arch:
        st.markdown(
            "Sube los archivos oficiales del colegio. Al subir, los datos anteriores se reemplazan."
        )
        st.info(
            "💾 Todos los archivos importados aquí se guardan permanentemente en la base de datos.",
            icon="💾",
        )
        col_a1, col_a2 = st.columns(2)

        with col_a1:
            st.markdown("**📅 Horarios Oficiales (.xlsx)**")
            st.info(
                "💡 El nombre de la pestaña en el Excel (ej. RUDDY) debe coincidir con el código del docente.",
                icon="💡",
            )
            f_horario = st.file_uploader(
                "Horarios",
                type=["xlsx"],
                key="adm_horario",
                label_visibility="collapsed",
            )
            if f_horario and st.button("Importar horarios", key="btn_imp_horario"):
                with st.spinner("Parseando horarios..."):
                    try:
                        registros = parse_horarios(f_horario.read())
                        guardar_horario_docente(registros)
                        st.success(
                            f"✅ {len(registros)} bloques importados para {len(set(r['nombre_hoja'] for r in registros))} docentes."
                        )
                    except Exception as _e:
                        st.error(f"Error: {_e}")

            st.markdown("**📋 Cronograma Semanal Docente (.xlsx)**")
            f_crono = st.file_uploader(
                "Cronograma",
                type=["xlsx"],
                key="adm_crono",
                label_visibility="collapsed",
            )
            if f_crono and st.button("Importar cronograma", key="btn_imp_crono"):
                with st.spinner("Parseando cronograma..."):
                    try:
                        registros = parse_cronograma(f_crono.read())
                        guardar_actividades_cronograma(registros)
                        st.success(f"✅ {len(registros)} actividades importadas.")
                    except Exception as _e:
                        st.error(f"Error: {_e}")

        with col_a2:
            st.markdown("**📆 Calendario Interno (.xlsx)**")
            f_cal = st.file_uploader(
                "Calendario", type=["xlsx"], key="adm_cal", label_visibility="collapsed"
            )
            if f_cal and st.button("Importar calendario", key="btn_imp_cal"):
                with st.spinner("Parseando calendario..."):
                    try:
                        registros = parse_calendario(f_cal.read())
                        guardar_eventos_calendario(registros)
                        st.success(f"✅ {len(registros)} eventos importados.")
                    except Exception as _e:
                        st.error(f"Error: {_e}")

            st.markdown("**🏛️ Comisiones Docentes (.docx)**")
            f_com = st.file_uploader(
                "Comisiones", type=["docx"], key="adm_com", label_visibility="collapsed"
            )
            if f_com and st.button("Importar comisiones", key="btn_imp_com"):
                with st.spinner("Parseando comisiones..."):
                    try:
                        registros = parse_comisiones(f_com.read())
                        guardar_comisiones(registros)
                        st.success(f"✅ {len(registros)} asignaciones importadas.")
                    except Exception as _e:
                        st.error(f"Error: {_e}")

            st.markdown("**👁️ Roles de Vigilancia Recreos (.pdf)**")
            st.caption(
                "Parsea automáticamente las grillas de PRIMARIA y SECUNDARIA del PDF."
            )
            f_vig = st.file_uploader(
                "Vigilancias", type=["pdf"], key="adm_vig", label_visibility="collapsed"
            )
            if f_vig and st.button("Importar desde PDF", key="btn_imp_vig"):
                with st.spinner("Parseando grillas de vigilancia..."):
                    try:
                        from db.horario import (
                            guardar_vigilancia_asignaciones,
                            get_vigilancia_asignaciones,
                        )

                        registros = parse_vigilancia_pdf(f_vig.read())
                        guardar_vigilancia_asignaciones(registros)
                        # Show summary
                        existentes = get_vigilancia_asignaciones()
                        resumen = {}
                        for r in existentes:
                            key = r["nombre_hoja"]
                            resumen[key] = resumen.get(key, 0) + 1
                        top = sorted(resumen.items(), key=lambda x: -x[1])[:8]
                        st.success(
                            f"✅ {len(registros)} asignaciones importadas "
                            f"({len(existentes)} total). "
                            + " | ".join(f"{k}: {v}" for k, v in top)
                        )
                    except ImportError:
                        st.error("Instala pdfplumber: pip install pdfplumber")
                    except Exception as _e:
                        st.error(f"Error: {_e}")

        # PLAN SEMANAL DOCENTE
        st.divider()
        st.markdown("**📋 Plan Semanal Docente (.docx / .txt)**")
        st.caption("La IA infiere objetivos desde el campo CONTENIDO.")
        pa_r1c1, pa_r1c2 = st.columns([5, 2])
        with pa_r1c1:
            f_plan = st.file_uploader(
                "Plan semanal",
                type=["docx", "txt"],
                key="adm_plan",
                label_visibility="collapsed",
            )
        with pa_r1c2:
            from datetime import datetime as _dt

            sw_default = max(
                1, (_dt.now().date() - _dt(2026, 2, 2).date()).days // 7 + 1
            )
            plan_semana = st.number_input(
                "Semana escolar N°",
                min_value=1,
                max_value=40,
                value=sw_default,
                key="adm_plan_semana",
            )

        pa_r2c1, pa_r2c2 = st.columns([3, 3])
        with pa_r2c1:
            plan_grado = st.text_input(
                "Grado / Nivel", key="adm_plan_grado", placeholder="Ej. 5to Primaria"
            )
        with pa_r2c2:
            hojas_plan = get_all_hojas()
            plan_hoja = st.selectbox(
                "Sincronizar con el horario de:",
                hojas_plan if hojas_plan else ["— sin horarios —"],
                key="adm_plan_hoja",
            )

        can_import = bool(f_plan and plan_grado.strip())
        if st.button(
            "🧠 Inferir objetivos y tareas progresivas",
            type="primary",
            key="btn_imp_plan",
            disabled=not can_import,
        ):
            with st.spinner("Aplicando Método Palma-Ribera…"):
                try:
                    import docx as _dx

                    if f_plan.name.lower().endswith(".docx"):
                        doc = _dx.Document(io.BytesIO(f_plan.read()))
                        texts = [
                            p.text.strip() for p in doc.paragraphs if p.text.strip()
                        ]
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text.strip():
                                        texts.append(cell.text.strip())
                        plan_txt = "\n".join(texts)
                    else:
                        plan_txt = f_plan.read().decode("utf-8", errors="ignore")
                    if not plan_txt.strip():
                        st.error("El archivo está vacío.")
                    else:
                        plan_prompt = (
                            "Eres un experto en el Método Palma-Ribera para planificación pedagógica "
                            f"de Las Palmas School. Grado: {plan_grado.strip()}.\n\n"
                            "Se te proporciona un plan semanal docente. Tu tarea es:\n"
                            "1. Identificar TODOS los bloques de clase con su horario y materia.\n"
                            "2. Tomar el campo CONTENIDO de cada bloque como fuente primaria.\n"
                            "3. INFERIR el objetivo de aprendizaje usando un verbo cognitivo de la taxonomía de Bloom.\n"
                            "4. DESCOMPONER cada contenido en exactamente 3 o 4 tareas progresivas.\n\n"
                            f"PLAN SEMANAL:\n{plan_txt[:6000]}\n\n"
                            "Responde ÚNICAMENTE con un array JSON válido."
                            '[{"hora_inicio":"07:55","hora_fin":"08:40","materia":"LENGUAJE",'
                            '"objetivo_inferido":"Los estudiantes lograrán...","tareas":["Tarea 1","Tarea 2","Tarea 3"]}]'
                        )
                        plan_keys = _get_keys()
                        plan_client = genai.Client(api_key=plan_keys[0])
                        plan_resp = plan_client.models.generate_content(
                            model=GEMINI_MODEL, contents=plan_prompt
                        )
                        plan_clean = (
                            plan_resp.text.replace("```json", "")
                            .replace("```", "")
                            .strip()
                        )
                        plan_data = json.loads(plan_clean)
                        if isinstance(plan_data, dict):
                            plan_data = plan_data.get(
                                "bloques", plan_data.get("clases", [plan_data])
                            )
                        if not isinstance(plan_data, list):
                            plan_data = []
                        tot_sessions = 0
                        tot_tasks = 0
                        for blk in plan_data:
                            h_ini = str(blk.get("hora_inicio", "00:00")).strip()
                            h_fin = str(blk.get("hora_fin", "00:00")).strip()
                            mat = str(blk.get("materia", "Sin materia")).strip()
                            obj_raw = blk.get(
                                "objetivo_inferido",
                                blk.get("objetivo", blk.get("tema", "")),
                            )
                            obj = str(obj_raw).strip() if obj_raw else ""
                            tasks_raw = blk.get(
                                "tareas", blk.get("tareas_progresivas", [])
                            )
                            tasks = []
                            if isinstance(tasks_raw, list):
                                for t in tasks_raw:
                                    if isinstance(t, str):
                                        tasks.append(t.strip())
                                    elif isinstance(t, dict):
                                        tasks.append(str(list(t.values())[0]).strip())
                            if not obj and not tasks:
                                continue
                            if not tasks and obj:
                                tasks = [obj]
                            if not tasks:
                                continue
                            sid = crear_sesion(
                                semana=int(plan_semana),
                                materia=mat,
                                grado=plan_grado.strip(),
                                tema=obj or f"Plan Semana {plan_semana} — {mat}",
                                hora_inicio=h_ini,
                                hora_fin=h_fin,
                            )
                            mo_list = []
                            if obj:
                                mo_list.append(
                                    {"texto": f"🎯 {obj}", "depende_de": None}
                                )
                            for i, t in enumerate(tasks, 1):
                                mo_list.append(
                                    {"texto": f"{i}· {t}", "depende_de": None}
                                )
                            guardar_micro_objetivos(
                                sid, mo_list, origen_semana=int(plan_semana)
                            )
                            tot_sessions += 1
                            tot_tasks += len(tasks)
                        if tot_sessions == 0:
                            st.warning(f"La IA no encontró bloques estructurados.")
                        else:
                            st.success(
                                f"✅ {tot_sessions} bloques importados · {tot_tasks} tareas · Semana {plan_semana}."
                            )
                except json.JSONDecodeError:
                    st.error(f"JSON inválido.")
                except Exception as _plan_e:
                    st.error(f"Error: {_plan_e}")

    # ═══════════════════════════════════════════════════════════════
    # TAB 2: GESTIÓN DE USUARIOS
    # ═══════════════════════════════════════════════════════════════
    with adm_tab_users:
        st.markdown("#### Crear nuevo usuario")
        hojas_disponibles = get_all_hojas()
        cu_col1, cu_col2 = st.columns(2)
        with cu_col1:
            cu_nombre = st.text_input(
                "Nombre completo", key="cu_nombre", placeholder="Ej. Ruddy Ribera"
            )
            cu_email = st.text_input(
                "Correo electrónico", key="cu_email", placeholder="ruddy@laspalmas.edu"
            )
            cu_pwd = st.text_input("Contraseña inicial", type="password", key="cu_pwd")
        with cu_col2:
            cu_hoja = st.selectbox(
                "Hoja de horario",
                options=hojas_disponibles
                if hojas_disponibles
                else ["— sin horario cargado —"],
                key="cu_hoja",
            )
            cu_hoja_manual = st.text_input(
                "O escribe el nombre manualmente",
                key="cu_hoja_manual",
                placeholder="Ej. RUDDY",
            )
            cu_rol = st.selectbox("Rol", ["docente", "admin"], key="cu_rol")
        if st.button("Crear usuario", type="primary", key="btn_crear_usuario"):
            hoja_final = cu_hoja_manual.strip().upper() or (
                cu_hoja if cu_hoja != "— sin horario cargado —" else ""
            )
            if cu_nombre and cu_email and cu_pwd and hoja_final:
                ok = crear_usuario(cu_email, cu_pwd, cu_nombre, hoja_final, cu_rol)
                if ok:
                    st.success(f"✅ Usuario **{cu_nombre}** creado.")
                else:
                    st.error("El correo ya está registrado.")
            else:
                st.warning("Completa todos los campos.")
        st.divider()
        st.markdown("#### Usuarios registrados")
        usuarios = get_all_usuarios()
        if usuarios:
            # ── Batch actions header ──────────────────────────────────────
            ba_col1, ba_col2 = st.columns(2)
            with ba_col1:
                st.caption(f"{len(usuarios)} usuario(s) total")
            with ba_col2:
                if st.button("🔄 Refrescar lista", key="btn_refresh_users"):
                    st.rerun()

            # ── Per-user rows ───────────────────────────────────────────────
            for u in usuarios:
                is_expanded = st.session_state.get(f"expand_user_{u['id']}", False)

                with st.container(border=True):
                    # ── Row header ─────────────────────────────────────────
                    u_col1, u_col2, u_col3, u_col4 = st.columns([3, 1, 1, 1])
                    estado = "✅" if u["activo"] else "⛔"
                    pwd_flag = " 🔑" if u.get("must_change_password") else ""
                    with u_col1:
                        st.markdown(
                            f"{estado} **{u['nombre']}**{pwd_flag}  \n"
                            f"`{u['email']}` · `{u['nombre_hoja']}` · _{u['rol']}_"
                        )
                    with u_col2:
                        lbl = "Desactivar" if u["activo"] else "Activar"
                        if st.button(lbl, key=f"toggle_{u['id']}"):
                            toggle_usuario_activo(u["id"], not u["activo"])
                            st.rerun()
                    with u_col3:
                        if st.button("Editar", key=f"edit_{u['id']}"):
                            st.session_state[f"expand_user_{u['id']}"] = not is_expanded
                            st.rerun()
                    with u_col4:
                        if st.button("🗑 Eliminar", key=f"del_{u['id']}"):
                            eliminar_usuario(u["id"])
                            st.rerun()

                    # ── Edit form (expanded) ─────────────────────────────────
                    if is_expanded:
                        st.markdown("---")
                        st.caption(
                            "Editar usuario (los campos vacíos no se actualizan)"
                        )
                        e_col1, e_col2, e_col3 = st.columns(3)
                        with e_col1:
                            e_nombre = st.text_input(
                                "Nombre completo",
                                value=u.get("nombre", ""),
                                key=f"e_nombre_{u['id']}",
                            )
                            e_email = st.text_input(
                                "Correo electrónico",
                                value=u.get("email", ""),
                                key=f"e_email_{u['id']}",
                            )
                        with e_col2:
                            e_hoja = st.text_input(
                                "Nombre hoja",
                                value=u.get("nombre_hoja", ""),
                                key=f"e_hoja_{u['id']}",
                            )
                            e_rol = st.selectbox(
                                "Rol",
                                ["docente", "admin"],
                                index=0 if u.get("rol") == "docente" else 1,
                                key=f"e_rol_{u['id']}",
                            )
                        with e_col3:
                            e_pwd = st.text_input(
                                "Nueva contraseña (dejar vacío para no cambiar)",
                                type="password",
                                key=f"e_pwd_{u['id']}",
                                placeholder="Mínimo 8 caracteres",
                            )
                            e_forzar_pwd = st.checkbox(
                                "Forzar cambio de password",
                                value=bool(u.get("must_change_password")),
                                key=f"e_forzar_{u['id']}",
                            )

                        save_clicked = st.button(
                            "💾 Guardar cambios",
                            type="primary",
                            key=f"save_{u['id']}",
                        )
                        if save_clicked:
                            ok, msg = actualizar_usuario_admin(
                                usuario_id=int(u["id"]),
                                nuevo_email=e_email.strip() or None,
                                nuevo_nombre=e_nombre.strip() or None,
                                nueva_password=e_pwd.strip() or None,
                            )
                            if ok:
                                # Also update must_change_password flag if changed
                                if e_forzar_pwd != bool(u.get("must_change_password")):
                                    from db._base import _conn, _USE_PG

                                    with _conn() as c:
                                        flag_val = (
                                            "TRUE"
                                            if e_forzar_pwd
                                            else "FALSE"
                                            if _USE_PG
                                            else "0"
                                        )
                                        c.execute(
                                            f"UPDATE usuarios SET must_change_password = {flag_val} WHERE id = %s"
                                            if _USE_PG
                                            else f"UPDATE usuarios SET must_change_password = {flag_val} WHERE id = ?",
                                            (int(u["id"]),),
                                        )
                                st.success(f"✅ Usuario actualizado: {msg}")
                                st.session_state[f"expand_user_{u['id']}"] = False
                                st.rerun()
                            else:
                                st.error(f"❌ {msg}")
        else:
            st.info("No hay usuarios registrados aún.")

    # ═══════════════════════════════════════════════════════════════
    # TAB 3: RESET DIARIO
    # ═══════════════════════════════════════════════════════════════
    with adm_tab_tracker:
        st.markdown("#### Reiniciar Tracker Diario")
        st.caption("Limpia todos los registros del día para un docente.")
        admin_rst_date = st.date_input("Fecha a reiniciar:")
        admin_rst_hojas = get_all_hojas()
        admin_rst_hoja = st.selectbox(
            "Docente", admin_rst_hojas if admin_rst_hojas else ["—"], key="adm_rst_hoja"
        )
        if st.button("↺ Reiniciar Tracker", type="primary", key="btn_adm_rst_dia"):
            if admin_rst_hoja and admin_rst_hoja != "—":
                reset_dia_docente(admin_rst_date.strftime("%Y-%m-%d"), admin_rst_hoja)
                st.success(f"Día {admin_rst_date} reiniciado para {admin_rst_hoja}.")
            else:
                st.error("Selecciona un docente válido.")

    # ═══════════════════════════════════════════════════════════════
    # TAB 4: EDITAR BLOQUES
    # ═══════════════════════════════════════════════════════════════
    with adm_tab_bloques:
        st.markdown("#### ✏️ Editar Bloques del Horario")
        st.caption("Edita, elimina o agrega bloques del horario de un docente.")

        # Teacher selector
        hojas = get_all_hojas()
        edit_hoja = st.selectbox(
            "Docente",
            options=hojas if hojas else ["— sin horarios —"],
            key="edit_bloque_hoja",
        )

        if edit_hoja and edit_hoja != "— sin horarios —":
            bloques = get_horario_docente_completo(edit_hoja)

            # Group by day
            dias_orden = ["lunes", "martes", "miercoles", "jueves", "viernes"]
            bloques_por_dia = {d: [] for d in dias_orden}
            for b in bloques:
                dia = (b.get("dia_semana") or "").lower()
                if dia in bloques_por_dia:
                    bloques_por_dia[dia].append(b)

            st.markdown("---")
            st.markdown(f"**Total de bloques:** {len(bloques)}")

            for dia in dias_orden:
                bloques_dia = bloques_por_dia[dia]
                if not bloques_dia:
                    continue

                with st.expander(f"**📅 {dia.upper()}** ({len(bloques_dia)} bloques)", expanded=True):
                    for blk in bloques_dia:
                        col1, col2, col3 = st.columns([3, 1, 1])
                        with col1:
                            tipo_icon = {"clase": "📚", "ingreso": "🚪", "vigilancia_recreo": "👁️", "recreo_libre": "☕", "planificacion": "📝", "atencion_ppff": "💬"}.get(blk.get("tipo_bloque", ""), "📌")
                            st.markdown(
                                f"**{tipo_icon} {blk.get('hora_inicio', '')} - {blk.get('hora_fin', '')}** "
                                f"| {blk.get('tipo_bloque', '')} | {blk.get('materia', '—')} | {blk.get('nivel_grado', '—')} | {blk.get('ubicacion', '—')}"
                            )
                        with col2:
                            if st.button("✏️", key=f"edit_btn_{blk['id']}"):
                                st.session_state[f"edit_bloque_{blk['id']}"] = not st.session_state.get(f"edit_bloque_{blk['id']}", False)
                        with col3:
                            if st.button("🗑️", key=f"del_btn_{blk['id']}"):
                                eliminar_bloque_horario_manual(int(blk["id"]))
                                st.success(f"Bloque eliminado.")
                                st.rerun()

                        # Edit form
                        if st.session_state.get(f"edit_bloque_{blk['id']}", False):
                            with st.form(key=f"edit_form_{blk['id']}"):
                                st.markdown("**Editar bloque:**")
                                e_hora_inicio = st.text_input("Hora inicio (HH:MM)", value=blk.get("hora_inicio", ""), key=f"e_hi_{blk['id']}")
                                e_hora_fin = st.text_input("Hora fin (HH:MM)", value=blk.get("hora_fin", ""), key=f"e_hf_{blk['id']}")
                                e_tipo = st.selectbox(
                                    "Tipo",
                                    ["clase", "ingreso", "vigilancia_recreo", "recreo_libre", "planificacion", "atencion_ppff"],
                                    index=["clase", "ingreso", "vigilancia_recreo", "recreo_libre", "planificacion", "atencion_ppff"].index(blk.get("tipo_bloque", "clase")) if blk.get("tipo_bloque") in ["clase", "ingreso", "vigilancia_recreo", "recreo_libre", "planificacion", "atencion_ppff"] else 0,
                                    key=f"e_tipo_{blk['id']}"
                                )
                                e_materia = st.text_input("Materia", value=blk.get("materia") or "", key=f"e_mat_{blk['id']}")
                                e_nivel = st.text_input("Nivel/Grado", value=blk.get("nivel_grado") or "", key=f"e_niv_{blk['id']}")
                                e_ubicacion = st.text_input("Ubicación", value=blk.get("ubicacion") or "", key=f"e_ub_{blk['id']}")
                                e_orden = st.number_input("Orden", value=blk.get("orden") or 0, key=f"e_ord_{blk['id']}")

                                col_save, col_cancel = st.columns(2)
                                with col_save:
                                    if st.form_submit_button("💾 Guardar"):
                                        guardar_bloque_horario_manual(
                                            nombre_hoja=edit_hoja,
                                            dia_semana=dia,
                                            hora_inicio=e_hora_inicio,
                                            hora_fin=e_hora_fin,
                                            tipo_bloque=e_tipo,
                                            materia=e_materia if e_materia else None,
                                            nivel_grado=e_nivel if e_nivel else None,
                                            ubicacion=e_ubicacion if e_ubicacion else None,
                                            orden=int(e_orden) if e_orden else None,
                                            bloque_id=int(blk["id"]),
                                        )
                                        st.success("Bloque actualizado.")
                                        st.session_state[f"edit_bloque_{blk['id']}"] = False
                                        st.rerun()
                                with col_cancel:
                                    if st.form_submit_button("Cancelar"):
                                        st.session_state[f"edit_bloque_{blk['id']}"] = False
                                        st.rerun()
                            st.markdown("---")

            # Add new block
            st.markdown("---")
            st.markdown("#### ➕ Agregar Nuevo Bloque")
            with st.expander("Agregar bloque", expanded=False):
                with st.form(key="add_new_bloque_form"):
                    add_dia = st.selectbox(
                        "Día",
                        dias_orden,
                        key="add_bloque_dia"
                    )
                    add_hi = st.text_input("Hora inicio (HH:MM)", placeholder="07:30", key="add_bloque_hi")
                    add_hf = st.text_input("Hora fin (HH:MM)", placeholder="08:30", key="add_bloque_hf")
                    add_tipo = st.selectbox(
                        "Tipo",
                        ["clase", "ingreso", "vigilancia_recreo", "recreo_libre", "planificacion", "atencion_ppff"],
                        key="add_bloque_tipo"
                    )
                    add_mat = st.text_input("Materia (opcional)", key="add_bloque_mat")
                    add_niv = st.text_input("Nivel/Grado (opcional)", key="add_bloque_niv")
                    add_ub = st.text_input("Ubicación (opcional)", key="add_bloque_ub")

                    if st.form_submit_button("➕ Agregar Bloque", type="primary"):
                        if add_hi and add_hf:
                            guardar_bloque_horario_manual(
                                nombre_hoja=edit_hoja,
                                dia_semana=add_dia,
                                hora_inicio=add_hi,
                                hora_fin=add_hf,
                                tipo_bloque=add_tipo,
                                materia=add_mat if add_mat else None,
                                nivel_grado=add_niv if add_niv else None,
                                ubicacion=add_ub if add_ub else None,
                            )
                            st.success("Bloque agregado.")
                            st.rerun()
                        else:
                            st.error("Hora inicio y fin son requeridas.")
