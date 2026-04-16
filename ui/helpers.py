"""
ui/helpers.py - CSS, Session State, and Helper Functions
========================================================
Contains:
- CSS design system
- Session state initialization
- Helper functions (forzar_lista, _bytes_hash, etc.)
- PDF analysis with Gemini
- Generic Gemini generator
- Diagnostic reader
- Schedule reader from XLSX
- PDC generator
"""

import os
import io
import json
import hashlib
import shutil
import time
import uuid
import tempfile
from datetime import datetime as _dt
from typing import Callable, Any

import streamlit as st
import openpyxl
from google import genai


# ═══════════════════════════════════════════════════════════════════════════════
# CSS — PRIA Design System v2
# ═══════════════════════════════════════════════════════════════════════════════

CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');

:root {
    --bg:       #0D0F12;
    --surface:  #16191D;
    --surface2: #1E2128;
    --border:   #2A2D35;
    --border2:  #3A3D48;
    --blue:     #00C2FF;
    --emerald:  #34C759;
    --amber:    #FF9500;
    --crimson:  #FF3B30;
    --text-1:   #E8EAF0;
    --text-2:   #9DA3B4;
    --text-3:   #6B7185;
    --g900: #1B5E20; --g700: #2E7D32; --g500: #43A047;
    --g300: #81C784; --g100: #C8E6C9; --g50: #E8F5E9;
    --shadow-xs: 0 1px 3px rgba(0,0,0,0.3);
    --shadow-sm: 0 2px 8px rgba(0,0,0,0.4);
    --shadow-md: 0 4px 20px rgba(0,0,0,0.5);
    --shadow-lg: 0 8px 40px rgba(0,0,0,0.6);
    --radius-sm: 8px;
    --radius-md: 12px;
    --radius-lg: 18px;
    --transition: all 0.18s cubic-bezier(0.4,0,0.2,1);
}

html, body, .stApp {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
    background: var(--bg) !important;
    min-height: 100vh;
}
.main .block-container {
    padding-top: 2rem !important;
    max-width: 1200px;
}

h1 { font-size: 2rem !important; font-weight: 800 !important; color: var(--text-1) !important; letter-spacing: -0.5px; }
h2 { color: var(--text-1) !important; font-weight: 700 !important; font-size: 1.3rem !important;
     border-left: 3px solid var(--blue); padding-left: 0.6rem; margin-top: 1.2rem !important; }
h3 { color: var(--text-1) !important; font-weight: 600 !important; font-size: 1.1rem !important; }
h4, h5, h6 { color: var(--text-2) !important; font-weight: 600 !important; }
body, p, span, li, div, td, th, caption, label, small, em, strong, blockquote,
.stMarkdown, .stText, .stCaption, div[data-testid="stMarkdownContainer"],
div[data-testid="stMarkdownContainer"] *, div[data-testid="stWidgetLabel"] p,
div[data-testid="stWidgetLabel"] span { color: var(--text-1) !important; }

div[data-testid="stTabs"] [role="tablist"] { gap: 4px !important; background: var(--surface) !important;
    border-radius: var(--radius-md) !important; padding: 5px !important; border: 1px solid var(--border) !important; }
button[data-baseweb="tab"] { border-radius: var(--radius-sm) !important; padding: 0.45rem 1.1rem !important;
    font-weight: 600 !important; font-size: 0.85rem !important; transition: var(--transition) !important;
    border: none !important; color: var(--text-2) !important; background: transparent !important; }
button[data-baseweb="tab"][aria-selected="true"] { background: var(--surface2) !important;
    color: var(--text-1) !important; box-shadow: var(--shadow-sm) !important; }

.stButton > button { background: var(--surface2) !important; color: var(--text-1) !important;
    font-weight: 600 !important; font-size: 0.88rem !important; border: 1px solid var(--border2) !important;
    border-radius: var(--radius-sm) !important; padding: 0.55rem 1.4rem !important;
    height: auto !important; min-height: 2.6rem !important; transition: var(--transition) !important; }
.stButton > button:hover { background: var(--border) !important; border-color: var(--blue) !important;
    color: var(--blue) !important; transform: translateY(-1px) !important; }
.stButton > button[kind="primary"] { background: var(--blue) !important; color: #000 !important;
    border: none !important; font-weight: 700 !important; }
.stButton > button[kind="primary"]:hover { background: #33CFFF !important; color: #000 !important;
    border: none !important; transform: translateY(-1px) !important; }

input[type="text"], input[type="number"], textarea { background-color: var(--surface) !important;
    color: var(--text-1) !important; border: 1px solid var(--border) !important;
    border-radius: var(--radius-sm) !important; transition: var(--transition) !important; }
input[type="text"]:focus, input[type="number"]:focus, textarea:focus { border-color: var(--blue) !important;
    box-shadow: 0 0 0 2px rgba(0,194,255,0.15) !important; outline: none !important; }
div[data-baseweb="select"] > div, div[data-baseweb="select"] span { background-color: var(--surface) !important;
    color: var(--text-1) !important; border: 1px solid var(--border) !important;
    border-radius: var(--radius-sm) !important; }
ul[data-baseweb="menu"] { background-color: var(--surface2) !important; border-radius: var(--radius-sm) !important;
    box-shadow: var(--shadow-md) !important; border: 1px solid var(--border) !important; }
ul[data-baseweb="menu"] li, ul[data-baseweb="menu"] li span { color: var(--text-1) !important; }
ul[data-baseweb="menu"] li:hover { background-color: var(--border) !important; }

div[data-testid="stExpander"] { background: var(--surface) !important; border: 1px solid var(--border) !important;
    border-radius: var(--radius-md) !important; margin-bottom: 0.6rem !important; overflow: hidden !important;
    border-left: 3px solid var(--border2) !important; transition: var(--transition) !important; }
details summary p, details summary span, div[data-testid="stExpander"] summary p { color: var(--text-2) !important;
    font-weight: 600 !important; font-size: 0.9rem !important; }
div[data-testid="stVerticalBlockBorderWrapper"] { background: var(--surface) !important;
    border: 1px solid var(--border) !important; border-radius: var(--radius-md) !important; }

div[data-testid="stAlert"] { border-radius: var(--radius-sm) !important; border-width: 0 !important;
    border-left-width: 3px !important; border-left-style: solid !important; }
div[data-testid="stAlert"][data-baseweb="notification"][kind="info"] { border-left-color: var(--blue) !important;
    background: rgba(0,194,255,0.07) !important; }
div[data-testid="stAlert"][data-baseweb="notification"][kind="success"] { border-left-color: var(--emerald) !important;
    background: rgba(52,199,89,0.07) !important; }
div[data-testid="stAlert"][data-baseweb="notification"][kind="warning"] { border-left-color: var(--amber) !important;
    background: rgba(255,149,0,0.07) !important; }
div[data-testid="stAlert"][data-baseweb="notification"][kind="error"] { border-left-color: var(--crimson) !important;
    background: rgba(255,59,48,0.07) !important; }

div[data-testid="stMetric"] { background: var(--surface) !important; border-radius: var(--radius-md) !important;
    padding: 1rem 1.2rem !important; border: 1px solid var(--border) !important;
    border-top: 2px solid var(--blue) !important; }
div[data-testid="stMetricValue"] { font-size: 1.8rem !important; font-weight: 800 !important; color: var(--text-1) !important; }
div[data-testid="stMetricLabel"] { font-weight: 500 !important; color: var(--text-2) !important;
    font-size: 0.78rem !important; text-transform: uppercase !important; letter-spacing: 0.06em !important; }

.stDataFrame { border-radius: var(--radius-md) !important; border: 1px solid var(--border) !important; overflow: hidden !important; }
.stDataFrame th { background: var(--surface2) !important; font-weight: 700 !important; font-size: 0.8rem !important;
    text-transform: uppercase !important; letter-spacing: 0.04em !important; color: var(--text-2) !important; }
.stDataFrame tr:hover td { background: var(--surface2) !important; }

section[data-testid="stSidebar"] { background: var(--surface) !important; border-right: 1px solid var(--border) !important; }
section[data-testid="stSidebar"] * { color: var(--text-1) !important; }

div[data-testid="stDownloadButton"] > button { background: var(--emerald) !important; color: #000 !important;
    font-weight: 700 !important; border-radius: var(--radius-sm) !important; border: none !important; }

.aid-day-header { background: var(--surface); border: 1px solid var(--border); border-radius: 8px;
    padding: 12px 18px; margin-bottom: 14px; display: flex; align-items: center; justify-content: space-between; }
.aid-day-header__left { display: flex; flex-direction: column; gap: 1px; }
.aid-day-header__title { font-size: 1.05rem; font-weight: 800; color: #E8EAF0; letter-spacing: -0.3px; }
.aid-day-header__sub { font-size: 0.75rem; font-weight: 500; color: #9DA3B4; }
.aid-day-header__right { display: flex; align-items: center; gap: 20px; }
.aid-day-header__cycle { font-size: 0.65rem; font-weight: 700; color: #00C2FF;
    letter-spacing: 0.1em; text-transform: uppercase; }
.aid-day-header__daycount { font-size: 0.68rem; font-weight: 600; color: #6B7185; letter-spacing: 0.04em; }
.aid-day-header__pct { font-size: 1.5rem; font-weight: 800; color: #E8EAF0; line-height: 1; }

.aid-badge { display: inline-block; padding: 1px 6px; border-radius: 4px; font-size: 0.65rem; font-weight: 700;
    letter-spacing: 0.07em; text-transform: uppercase; vertical-align: middle; flex-shrink: 0; }
.aid-badge--closed { background: rgba(52,199,89,0.12); color: #34C759; }
.aid-badge--active { background: rgba(0,194,255,0.12); color: #00C2FF; }
.aid-badge--pending { background: rgba(255,149,0,0.12); color: #FF9500; }

.aid-time { font-family: 'SF Mono', 'Fira Code', 'Consolas', monospace; font-size: 0.8rem; font-weight: 600;
    color: #9DA3B4; letter-spacing: 0.02em; }
</style>
"""


# ═══════════════════════════════════════════════════════════════════════════════
# SYSTEM PATHS
# ═══════════════════════════════════════════════════════════════════════════════

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROMPTS_DIR = os.path.join(BASE_DIR, "prompts_maestros")


def _secret(key: str, default=None):
    try:
        v = st.secrets.get(key, None)
        if v is not None:
            return v
    except Exception:
        pass
    return os.environ.get(key, default)


CACHE_DIR = _secret("CACHE_DIR", os.path.join(BASE_DIR, "cache_libros"))
LOG_DIR = _secret("LOG_DIR", os.path.join(BASE_DIR, "logs"))
SESSION_BASE_DIR = os.path.join(tempfile.gettempdir(), "pria_sessions")
for _d in (CACHE_DIR, LOG_DIR, SESSION_BASE_DIR):
    os.makedirs(_d, exist_ok=True)

GEMINI_MODEL = "gemini-2.5-flash"


# ═══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ═══════════════════════════════════════════════════════════════════════════════

SESSION_DEFAULTS = {
    "session_id": None,
    "autenticado": False,
    "grado_nivel": "5to primaria",
    "key_index": 0,
    "last_error": None,
    "last_generar_fn": None,
    "last_generar_vars": None,
    "last_generar_json": False,
    "uploaded_tb_bytes": None,
    "uploaded_tb_name": None,
    "uploaded_tb_hash": None,
    "tb_extracted": None,
    "uploaded_sb_bytes": None,
    "uploaded_sb_name": None,
    "uploaded_sb_hash": None,
    "sb_extracted": None,
    "uploaded_diag_files": [],
    "diagnosticos_texto": None,
    "diagnosticos_tabla": [],
    "res_m0a": None,
    "res_m0b": None,
    "res_m0c": None,
    "tema_activo": "",
    "tema_hash": "",
    "leccion_index": 0,
    "conceptos_activos": [],
    "palabras_clave_activas": [],
    "contenido_tema_activo": "",
    "res_m1a": None,
    "res_m1a_prev": None,
    "res_m1b": None,
    "res_m1c": None,
    "res_m2a": None,
    "res_m2b": None,
    "mostrar_adaptaciones_prev": False,
    "_pptx_cache": None,
    "teacher_name": "",
    "school_name": "",
}


def init_session_state():
    """Initialize session state with defaults if not present."""
    for key, val in SESSION_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = val


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════


def forzar_lista(valor) -> list:
    if isinstance(valor, list):
        return [str(v) for v in valor if v]
    if isinstance(valor, str):
        return [v.strip() for v in valor.split(",") if v.strip()]
    return []


def _bytes_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _get_keys() -> list:
    import json as _json

    raw = _secret("GEMINI_API_KEYS", "[]")
    if isinstance(raw, list):
        return [str(k).strip() for k in raw if str(k).strip()]
    try:
        parsed = _json.loads(raw)
        if isinstance(parsed, list):
            keys = [str(k).strip() for k in parsed if str(k).strip()]
        elif isinstance(parsed, str):
            keys = [parsed.strip()] if parsed.strip() else []
        else:
            keys = []
        if keys:
            return keys
    except Exception:
        pass
    single = str(_secret("GEMINI_API_KEY", "") or "").strip()
    if single:
        return [single]
    raw_s = str(raw or "").strip()
    return [raw_s] if raw_s and raw_s != "[]" else []


def _rotate_key():
    keys = _get_keys()
    if keys:
        st.session_state.key_index = (st.session_state.key_index + 1) % len(keys)


def _topic_hash(tema: str) -> str:
    if not tema:
        return ""
    return hashlib.md5(str(tema).strip().lower().encode()).hexdigest()


def _cleanup_old_sessions():
    try:
        ahora = time.time()
        for nombre in os.listdir(SESSION_BASE_DIR):
            ruta = os.path.join(SESSION_BASE_DIR, nombre)
            if os.path.isdir(ruta) and ahora - os.path.getmtime(ruta) > 86400:
                shutil.rmtree(ruta, ignore_errors=True)
    except Exception:
        pass


def _cleanup_old_cache():
    try:
        ahora = time.time()
        for nombre in os.listdir(CACHE_DIR):
            ruta = os.path.join(CACHE_DIR, nombre)
            if os.path.isfile(ruta) and ahora - os.path.getmtime(ruta) > 86400 * 30:
                os.remove(ruta)
    except Exception:
        pass


def _log_event(action: str, success: bool, error_msg: str = ""):
    import datetime

    entry = {
        "ts": datetime.datetime.utcnow().isoformat(),
        "session": (st.session_state.session_id or "")[:8],
        "action": action,
        "ok": success,
        "err": error_msg[:300] if error_msg else "",
    }
    try:
        log_path = os.path.join(LOG_DIR, "pilot.log")
        with open(log_path, "a", encoding="utf-8") as fh:
            fh.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception:
        pass


def get_session_temp_dir() -> str:
    if not st.session_state.session_id:
        st.session_state.session_id = str(uuid.uuid4())
    session_dir = os.path.join(SESSION_BASE_DIR, st.session_state.session_id)
    os.makedirs(session_dir, exist_ok=True)
    return session_dir


def _cargar_cache_hash(h: str) -> dict | None:
    path = os.path.join(CACHE_DIR, f"{h}.json")
    if not os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as fh:
            return json.load(fh)
    except Exception:
        return None


def _guardar_cache_hash(h: str, data: dict):
    path = os.path.join(CACHE_DIR, f"{h}.json")
    try:
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(data, fh, ensure_ascii=False, indent=2)
    except Exception:
        pass


# ── Motor output cache (7-day TTL, keyed on prompt + variables) ───────────────

_MOTOR_CACHE_TTL = 86400 * 7  # 7 days


def _motor_cache_key(prompt_filename: str, variables: dict) -> str:
    """Stable hash key for a motor call: prompt name + sorted variable contents."""
    payload = json.dumps(
        {"motor": prompt_filename, "vars": variables}, sort_keys=True, ensure_ascii=False
    )
    return "motor_" + hashlib.sha256(payload.encode()).hexdigest()


def _cargar_motor_cache(key: str):
    """Return cached motor result or None if missing/expired."""
    path = os.path.join(CACHE_DIR, f"{key}.json")
    if not os.path.exists(path):
        return None
    try:
        if time.time() - os.path.getmtime(path) > _MOTOR_CACHE_TTL:
            os.remove(path)
            return None
        with open(path, "r", encoding="utf-8") as fh:
            entry = json.load(fh)
        return entry.get("result")
    except Exception:
        return None


def _guardar_motor_cache(key: str, result, motor: str):
    """Persist a motor result to disk cache."""
    path = os.path.join(CACHE_DIR, f"{key}.json")
    try:
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(
                {"motor": motor, "result": result, "ts": time.time()},
                fh,
                ensure_ascii=False,
                indent=2,
            )
    except Exception:
        pass


def limpiar_motor_cache():
    """Delete all motor_*.json cache files. Called from sidebar clear button."""
    try:
        for nombre in os.listdir(CACHE_DIR):
            if nombre.startswith("motor_") and nombre.endswith(".json"):
                os.remove(os.path.join(CACHE_DIR, nombre))
    except Exception:
        pass


# Cleanup on import
_cleanup_old_sessions()
_cleanup_old_cache()


# ═══════════════════════════════════════════════════════════════════════════════
# DIAGNOSTIC READER
# ═══════════════════════════════════════════════════════════════════════════════


def leer_diagnosticos(archivos_subidos: list) -> tuple:
    if not archivos_subidos:
        return None, []
    textos_raw = []
    keys = _get_keys()
    for nombre_archivo, contenido_bytes in archivos_subidos:
        ext = nombre_archivo.lower().rsplit(".", 1)[-1]
        file_hash = _bytes_hash(contenido_bytes)
        cache_key = f"diag_{file_hash}"
        cached = _cargar_cache_hash(cache_key)
        if cached and "texto" in cached:
            textos_raw.append(f"[{nombre_archivo}]\n{cached['texto']}")
            continue
        if ext == "txt":
            try:
                c = contenido_bytes.decode("utf-8", errors="ignore").strip()
                if c:
                    textos_raw.append(f"[{nombre_archivo}]\n{c}")
                    _guardar_cache_hash(cache_key, {"texto": c})
            except Exception as e:
                textos_raw.append(f"[{nombre_archivo}] Error TXT: {e}")
        elif ext == "docx":
            try:
                import docx as _docx

                doc = _docx.Document(io.BytesIO(contenido_bytes))
                p = "\n".join(x.text for x in doc.paragraphs if x.text.strip())
                if p:
                    textos_raw.append(f"[{nombre_archivo}]\n{p}")
                    _guardar_cache_hash(cache_key, {"texto": p})
            except ImportError:
                textos_raw.append(f"[{nombre_archivo}] Instala python-docx.")
            except Exception as e:
                textos_raw.append(f"[{nombre_archivo}] Error DOCX: {e}")
        elif ext == "pdf" and keys:
            temp_path = os.path.join(
                get_session_temp_dir(), f"diag_{file_hash[:8]}_{nombre_archivo}"
            )
            with open(temp_path, "wb") as fh:
                fh.write(contenido_bytes)
            intentos_pdf = 0
            while intentos_pdf < len(keys):
                idx = st.session_state.key_index % len(keys)
                try:
                    client = genai.Client(api_key=keys[idx])
                    uploaded = client.files.upload(file=temp_path)
                    resp = client.models.generate_content(
                        model=GEMINI_MODEL,
                        contents=[
                            uploaded,
                            "Extrae todo el texto relevante de este documento de diagnóstico escolar. "
                            "Incluye nombre del estudiante (si aparece), diagnóstico principal, necesidades específicas "
                            "y recomendaciones pedagógicas. Responde SOLO con el texto extraído, sin comentarios adicionales.",
                        ],
                    )
                    client.files.delete(name=uploaded.name)
                    texto = resp.text.strip()
                    textos_raw.append(f"[{nombre_archivo}]\n{texto}")
                    _guardar_cache_hash(cache_key, {"texto": texto})
                    break
                except Exception as e:
                    err = str(e).upper()
                    if any(x in err for x in ["429", "QUOTA", "EXHAUSTED"]):
                        _rotate_key()
                        intentos_pdf += 1
                    else:
                        textos_raw.append(
                            f"[{nombre_archivo}] Error PDF: {str(e)[:200]}"
                        )
                        break
    if not textos_raw:
        return None, []
    texto_combinado = "\n\n".join(textos_raw)
    prompt_sintesis = (
        "Eres especialista en necesidades educativas especiales.\n"
        "Documentos de diagnóstico del aula:\n\n"
        f"{texto_combinado}\n\n"
        "Devuelve ÚNICAMENTE un objeto JSON con:\n"
        '  "perfil_texto": string compacto (ej. "TEA (2), TDAH (1)")\n'
        '  "tabla": lista de objetos con "Diagnóstico", "Estudiantes" (int), "Notas pedagógicas"\n'
        "Sin texto adicional ni markdown."
    )
    intentos_sint = 0
    while intentos_sint < len(keys):
        idx = st.session_state.key_index % len(keys)
        try:
            client = genai.Client(api_key=keys[idx])
            resp = client.models.generate_content(
                model=GEMINI_MODEL, contents=prompt_sintesis
            )
            clean = resp.text.replace("```json", "").replace("```", "").strip()
            data = json.loads(clean)
            return data.get("perfil_texto", None), data.get("tabla", [])
        except Exception as e:
            err = str(e).upper()
            if any(x in err for x in ["429", "QUOTA", "EXHAUSTED"]):
                _rotate_key()
                intentos_sint += 1
            else:
                break
    perfil_crudo = "; ".join(
        t.split("\n")[0].replace("[", "").replace("]", "") for t in textos_raw[:5]
    )
    return perfil_crudo, []


# ═══════════════════════════════════════════════════════════════════════════════
# MOTOR ALPHA — PDF Analysis
# ═══════════════════════════════════════════════════════════════════════════════


def analizar_pdf_ocr(pdf_bytes: bytes, filename: str) -> dict:
    VACIO = {
        "unidad_real": "Sin datos",
        "temas": [],
        "contenido_temas": {},
        "paginas_temas": {},
    }
    h = _bytes_hash(pdf_bytes)
    datos_cache = _cargar_cache_hash(h)
    if datos_cache:
        return datos_cache
    ruta_pdf = os.path.join(get_session_temp_dir(), filename)
    with open(ruta_pdf, "wb") as fh:
        fh.write(pdf_bytes)
    ruta_prompt = os.path.join(PROMPTS_DIR, "Motor_Alpha-2.txt")
    if not os.path.exists(ruta_prompt):
        return {**VACIO, "unidad_real": "Falta Motor_Alpha-2.txt"}
    with open(ruta_prompt, "r", encoding="utf-8") as fh:
        prompt_base = fh.read()
    prompt_completo = (
        prompt_base
        + """
═══════════════════════════════════════════════════════
INSTRUCCIÓN ADICIONAL — CONTENIDO CURRICULAR Y PÁGINAS
═══════════════════════════════════════════════════════
Además del índice, examina el cuerpo de cada tema y extrae:
1. "contenido_temas": dict donde cada clave es el nombre exacto del tema
   y el valor es el texto LITERAL del libro para ese tema (mínimo 80 palabras por tema).
2. "paginas_temas": dict donde cada clave es el nombre exacto del tema
   y el valor es el rango de páginas como aparece en el índice.
"""
    )
    keys = _get_keys()
    if not keys:
        return {**VACIO, "unidad_real": "Falta GEMINI_API_KEYS"}
    intentos = 0
    while intentos < len(keys):
        idx = st.session_state.key_index % len(keys)
        client = genai.Client(api_key=keys[idx])
        try:
            uploaded = client.files.upload(file=ruta_pdf)
            response = client.models.generate_content(
                model=GEMINI_MODEL, contents=[uploaded, prompt_completo]
            )
            client.files.delete(name=uploaded.name)
            clean = response.text.replace("```json", "").replace("```", "").strip()
            data = json.loads(clean)
            if "contenido_temas" not in data:
                data["contenido_temas"] = {}
            if "paginas_temas" not in data:
                data["paginas_temas"] = {}
            _guardar_cache_hash(h, data)
            return data
        except json.JSONDecodeError:
            return {**VACIO, "unidad_real": "Error de parsing JSON"}
        except Exception as e:
            err = str(e).upper()
            if any(x in err for x in ["429", "RESOURCE_EXHAUSTED", "QUOTA"]):
                st.session_state.key_index = (idx + 1) % len(keys)
                intentos += 1
            else:
                return {**VACIO, "unidad_real": f"Error Gemini: {str(e)[:150]}"}
    return {**VACIO, "unidad_real": "Todas las llaves agotadas"}


# ═══════════════════════════════════════════════════════════════════════════════
# MOTOR GEMINI GENERIC
# ═══════════════════════════════════════════════════════════════════════════════


def load_motor_prompt(motor_name: str) -> str:
    ruta = os.path.join(PROMPTS_DIR, f"{motor_name}.txt")
    if os.path.exists(ruta):
        with open(ruta, "r", encoding="utf-8") as f:
            return f.read()
    return None


def generar_con_gemini(
    prompt_filename: str,
    variables: dict,
    expect_json: bool = False,
    required_fields: list | None = None,
    use_cache: bool = True,
):
    st.session_state.last_generar_fn = prompt_filename
    st.session_state.last_generar_vars = variables
    st.session_state.last_generar_json = expect_json
    motor_name = prompt_filename.replace(".txt", "")
    template = load_motor_prompt(motor_name)
    if not template:
        msg = f"Error cargando prompt: {prompt_filename}"
        st.error(f"❌ {msg}")
        _log_event(f"generar:{prompt_filename}", False, msg)
        return None
    if "grado_nivel" not in variables:
        variables = {
            **variables,
            "grado_nivel": st.session_state.get("grado_nivel", "5to primaria"),
        }

    # ── Cache check ──────────────────────────────────────────────────────────
    cache_key = _motor_cache_key(prompt_filename, variables)
    if use_cache:
        cached = _cargar_motor_cache(cache_key)
        if cached is not None:
            _log_event(f"cache_hit:{prompt_filename}", True)
            return cached

    prompt_final = template
    for key, value in variables.items():
        prompt_final = prompt_final.replace("{" + key + "}", str(value))
    keys = _get_keys()
    intentos = 0
    while intentos < len(keys):
        idx = st.session_state.key_index % len(keys)
        client = genai.Client(api_key=keys[idx])
        try:
            _log_event(f"generar:{prompt_filename}", True)
            response = client.models.generate_content(
                model=GEMINI_MODEL, contents=prompt_final
            )
            if expect_json:
                clean = response.text.replace("```json", "").replace("```", "").strip()
                try:
                    data = json.loads(clean)
                except json.JSONDecodeError as e:
                    raw_snippet = response.text[:600]
                    st.session_state.last_error = response.text
                    _log_event(f"json_error:{prompt_filename}", False, str(e))
                    st.error(
                        f"❌ **JSON inválido** en `{prompt_filename}` — {e}\n\n**Respuesta (primeros 600 chars):**\n```\n{raw_snippet}\n```"
                    )
                    return None
                if required_fields:
                    faltantes = [f for f in required_fields if f not in data]
                    if faltantes:
                        st.warning(
                            f"⚠️ `{prompt_filename}` sin campos: `{', '.join(faltantes)}`"
                        )
                st.session_state.last_error = None
                if use_cache:
                    _guardar_motor_cache(cache_key, data, prompt_filename)
                return data
            st.session_state.last_error = None
            if use_cache:
                _guardar_motor_cache(cache_key, response.text, prompt_filename)
            return response.text
        except Exception as e:
            err = str(e).upper()
            if any(x in err for x in ["429", "RESOURCE_EXHAUSTED", "QUOTA"]):
                st.session_state.key_index = (idx + 1) % len(keys)
                intentos += 1
            else:
                _log_event(f"generar:{prompt_filename}", False, str(e))
                st.error(f"❌ Error en **{prompt_filename}**: {str(e)[:300]}")
                return None
    msg = "Todas las llaves API han agotado su cuota."
    _log_event(f"generar:{prompt_filename}", False, msg)
    st.error(f"⚠️ {msg}")
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# SCHEDULE READER (XLSX)
# ═══════════════════════════════════════════════════════════════════════════════


def leer_horario_xlsx(xlsx_bytes: bytes, teacher_name: str) -> tuple:
    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    target_sheet = None
    for sname in wb.sheetnames:
        if teacher_name.upper() in sname.upper():
            target_sheet = wb[sname]
            break
    if target_sheet is None:
        return None, None
    data = []
    for row in target_sheet.iter_rows(values_only=True):
        data.append([str(c).strip() if c is not None else "" for c in row])
    days = ["LUNES", "MARTES", "MIERCOLES", "JUEVES", "VIERNES"]
    header_row_idx = -1
    for i, row in enumerate(data):
        if "LUNES" in row:
            header_row_idx = i
            break
    if header_row_idx == -1:
        return [], []
    headers = data[header_row_idx]
    lunes_idx = headers.index("LUNES")
    time_headers = [h for h in headers[:lunes_idx] if h and h != "None"] or ["HORARIO"]
    day_indices = {d: (headers.index(d) if d in headers else -1) for d in days}
    schedule_map = {}
    skip_words = ["VIGILANCIA", "SNACK", "INGRESO", "NONE"]
    for row in data[header_row_idx + 1 :]:
        row_times = []
        is_valid = False
        for c in range(lunes_idx):
            val = str(row[c] if c < len(row) else "").strip()
            row_times.append(val)
            if "-" in val and any(ch.isdigit() for ch in val):
                is_valid = True
        if not is_valid:
            continue
        key = tuple(row_times)
        if key not in schedule_map:
            schedule_map[key] = {"times": row_times, **{d: [] for d in days}}
        for day, col_idx in day_indices.items():
            if col_idx != -1 and col_idx < len(row):
                val = str(row[col_idx] or "").strip()
                if val and not any(w in val.upper() for w in skip_words):
                    schedule_map[key][day].append(val)
    schedule = []
    for entry in schedule_map.values():
        schedule.append(
            {
                "times": entry["times"],
                "LUNES": " / ".join(entry["LUNES"]) or "Libre",
                "MARTES": " / ".join(entry["MARTES"]) or "Libre",
                "MIERCOLES": " / ".join(entry["MIERCOLES"]) or "Libre",
                "JUEVES": " / ".join(entry["JUEVES"]) or "Libre",
                "VIERNES": " / ".join(entry["VIERNES"]) or "Libre",
            }
        )
    return schedule, time_headers


# ═══════════════════════════════════════════════════════════════════════════════
# PDC GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════


def _leer_docx_texto(docx_bytes: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def generar_pdc_trimestral(
    teacher_name: str,
    level: str,
    subject: str,
    pdc_texto: str,
    semanas_calendario: int = 12,
) -> str | None:
    max_units = (
        4 if "matemática" in subject.lower() or "matematica" in subject.lower() else 3
    )
    return generar_con_gemini(
        prompt_filename="Motor_PDC_Trimestral.txt",
        variables={
            "teacher_name": teacher_name,
            "level": level,
            "subject": subject,
            "max_units": str(max_units),
            "semanas_calendario": str(semanas_calendario),
            "pdc_texto": pdc_texto[:3000],
        },
        expect_json=False,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT WRAPPERS (imported from exportar.py)
# ═══════════════════════════════════════════════════════════════════════════════


def get_motor_stats() -> dict:
    return {
        "total_motors": 0,
        "total_uses": 0,
        "success_rate": 0.0,
        "avg_duration": 0.0,
        "motors": [],
    }


# Module-level helpers object for convenience
class Helpers:
    """Convenience class grouping helper functions."""

    forzar_lista = staticmethod(forzar_lista)
    bytes_hash = staticmethod(_bytes_hash)
    get_keys = staticmethod(_get_keys)
    rotate_key = staticmethod(_rotate_key)
    topic_hash = staticmethod(_topic_hash)
    log_event = staticmethod(_log_event)
    session_temp_dir = staticmethod(get_session_temp_dir)
    cargar_cache = staticmethod(_cargar_cache_hash)
    guardar_cache = staticmethod(_guardar_cache_hash)


helpers = Helpers()
